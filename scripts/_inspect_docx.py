# -*- coding: utf-8 -*-
import sys
from docx import Document

def inspect(path, out):
    lines = []
    def w(s=""):
        lines.append(str(s))
    w("="*100)
    w("FILE: " + path)
    w("="*100)
    doc = Document(path)
    w("\n--- SECTIONS ---")
    for i, s in enumerate(doc.sections):
        w(f"Section {i}: W={s.page_width} H={s.page_height} L={s.left_margin} R={s.right_margin} T={s.top_margin} B={s.bottom_margin}")
    w(f"\nTOTAL PARAGRAPHS: {len(doc.paragraphs)}")
    w(f"TOTAL TABLES: {len(doc.tables)}")
    w("\n--- PARAGRAPHS (idx | style | font | text) ---")
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        style = p.style.name if p.style else "?"
        finfo = ""
        if p.runs:
            r = p.runs[0]
            sz = r.font.size.pt if r.font.size else None
            bold = r.font.bold
            name = r.font.name
            align = p.alignment
            finfo = f"[f={name} sz={sz} b={bold} al={align}]"
        if txt or style not in ("Normal",):
            w(f"{idx:04d} <{style}> {finfo} {txt}")
    w("\n--- TABLES ---")
    for ti, t in enumerate(doc.tables):
        w(f"\nTABLE {ti}: rows={len(t.rows)} cols={len(t.columns)}")
        for ri, row in enumerate(t.rows):
            cells = [c.text.strip().replace('\n',' / ') for c in row.cells]
            w(f"  R{ri}: " + " | ".join(cells))
    with open(out, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print("WROTE", out, "lines:", len(lines))

if __name__ == "__main__":
    inspect(sys.argv[1], sys.argv[2])
