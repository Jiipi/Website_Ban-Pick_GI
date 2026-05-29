from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "BaoCao_TongHop_Website_BanPick_Genshin.docx"
OUTPUT = ROOT / "BaoCao_TongHop_Website_BanPick_Genshin_CaiThien.docx"


def set_run_font(run, size: float = 13, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_para_text(paragraph: Paragraph, text: str, *, size: float = 13, italic: bool = False) -> None:
    style = paragraph.style
    alignment = paragraph.alignment
    paragraph.clear()
    paragraph.style = style
    paragraph.alignment = alignment
    run = paragraph.add_run(text)
    set_run_font(run, size=size, italic=italic)


def insert_paragraph_after(
    paragraph: Paragraph,
    text: str = "",
    *,
    style: str | None = None,
    italic: bool = False,
    bold: bool = False,
    align: int | None = None,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if align is not None:
        new_para.alignment = align
    run = new_para.add_run(text)
    size = 12 if italic else 13
    if style == "Heading 1":
        size = 16
    elif style == "Heading 2":
        size = 14
    set_run_font(run, size=size, bold=bold or bool(style and style.startswith("Heading")), italic=italic)
    return new_para


def insert_paragraph_before(
    paragraph: Paragraph,
    text: str = "",
    *,
    style: str | None = None,
    italic: bool = False,
    bold: bool = False,
    align: int | None = None,
) -> Paragraph:
    new_para = paragraph.insert_paragraph_before(text)
    if style:
        new_para.style = style
    if align is not None:
        new_para.alignment = align
    for run in new_para.runs:
        size = 12 if italic else 13
        if style == "Heading 1":
            size = 16
        elif style == "Heading 2":
            size = 14
        set_run_font(run, size=size, bold=bold or bool(style and style.startswith("Heading")), italic=italic)
    return new_para


def append_paragraph(
    doc: Document,
    text: str = "",
    *,
    style: str | None = None,
    italic: bool = False,
    bold: bool = False,
    align: int | None = None,
) -> Paragraph:
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    size = 12 if italic else 13
    if style == "Heading 1":
        size = 16
    elif style == "Heading 2":
        size = 14
    set_run_font(run, size=size, bold=bold or bool(style and style.startswith("Heading")), italic=italic)
    return paragraph


def insert_table_after(
    doc: Document,
    paragraph: Paragraph,
    headers: list[str],
    rows: list[list[str]],
    caption: str,
) -> Paragraph:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = 1
    apply_borders(table)
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)

    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_para.add_run(caption)
    set_run_font(caption_run, size=12, italic=True)

    paragraph._p.addnext(table._tbl)
    table._tbl.addnext(caption_para._p)
    return caption_para


def set_cell(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=10.5, bold=bold)


def apply_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:color"), "666666")
        borders.append(tag)
    tbl_pr.append(borders)


def paragraph_text(paragraph: Paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph_text(paragraph) == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def replace_exact(doc: Document, old: str, new: str, *, italic: bool = False) -> int:
    count = 0
    for paragraph in doc.paragraphs:
        if paragraph_text(paragraph) == old:
            set_para_text(paragraph, new, italic=italic)
            count += 1
    return count


def replace_contains(doc: Document, needle: str, new: str, *, italic: bool = False) -> int:
    count = 0
    for paragraph in doc.paragraphs:
        if needle in paragraph_text(paragraph):
            set_para_text(paragraph, new, italic=italic)
            count += 1
    return count


def add_frontmatter_notes(doc: Document) -> None:
    replacements = {
        "HỌC PHẦN: [TÊN HỌC PHẦN]": "HỌC PHẦN: [Cần bổ sung: tên học phần]",
        "GVHD: [TÊN GIẢNG VIÊN HƯỚNG DẪN]": "GVHD: [Cần bổ sung: tên giảng viên hướng dẫn]",
        "SVTH: [HỌ VÀ TÊN SINH VIÊN]": "SVTH: [Cần bổ sung: họ tên sinh viên/nhóm sinh viên]",
        "MSSV: [MÃ SỐ SINH VIÊN]": "MSSV: [Cần bổ sung: mã số sinh viên]",
        "LỚP: [TÊN LỚP]": "LỚP: [Cần bổ sung: lớp]",
    }
    for old, new in replacements.items():
        replace_exact(doc, old, new)

    toc = find_paragraph(doc, "MỤC LỤC")
    insert_paragraph_after(
        toc,
        "[Cần bổ sung thủ công: trong Word, tạo/cập nhật mục lục tự động bằng References > Table of Contents trước khi nộp chính thức.]",
        italic=True,
    )

    list_fig = find_paragraph(doc, "DANH MỤC HÌNH ẢNH")
    insert_paragraph_after(
        list_fig,
        "[Cần bổ sung thủ công: chèn ảnh chụp thật hoặc sơ đồ thật, sau đó cập nhật danh mục hình tự động.]",
        italic=True,
    )

    list_table = find_paragraph(doc, "DANH MỤC BẢNG")
    list_table_anchor = list_table
    for line in [
        "Bảng 3.10. Thiết kế chi tiết một số bảng dữ liệu chính .............................................",
        "Bảng 3.11. Danh sách API tiêu biểu của hệ thống .............................................",
        "Bảng 4.3. Checklist triển khai và minh chứng cần bổ sung .............................................",
        "Bảng A.1. Minh chứng sử dụng công cụ AI trong quá trình phát triển .............................................",
    ]:
        list_table_anchor = insert_paragraph_after(list_table_anchor, line)

    chapter_1 = find_paragraph(doc, "CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI")
    heading = insert_paragraph_before(chapter_1, "DANH MỤC TỪ VIẾT TẮT", style="Heading 1")
    rows = [
        ["API", "Application Programming Interface", "Giao diện lập trình ứng dụng"],
        ["CRUD", "Create, Read, Update, Delete", "Nhóm thao tác tạo, đọc, cập nhật và xóa dữ liệu"],
        ["ORM", "Object Relational Mapping", "Kỹ thuật ánh xạ đối tượng với bảng dữ liệu"],
        ["RLS", "Row Level Security", "Cơ chế phân quyền theo từng dòng dữ liệu trong Supabase/PostgreSQL"],
        ["UID", "User Identifier", "Mã định danh người chơi Genshin Impact"],
        ["VPS", "Virtual Private Server", "Máy chủ riêng ảo dùng để triển khai ứng dụng"],
    ]
    insert_table_after(doc, heading, ["Từ viết tắt", "Tên đầy đủ", "Ý nghĩa trong báo cáo"], rows, "Bảng 0.1. Danh mục từ viết tắt")


def fix_section_numbering(doc: Document) -> None:
    replacements = {
        "3.12. Tính liên tục của dữ liệu trong trận đấu": "3.13. Tính liên tục của dữ liệu trong trận đấu",
        "3.13. Tách rõ chức năng lõi và chức năng mở rộng": "3.14. Tách rõ chức năng lõi và chức năng mở rộng",
        "3.14. Thiết kế phân quyền trong phòng đấu": "3.15. Thiết kế phân quyền trong phòng đấu",
        "3.15. Thiết kế trạng thái và vòng đời phòng": "3.16. Thiết kế trạng thái và vòng đời phòng",
        "3.16. Thiết kế khả năng mở rộng luật giải đấu": "3.17. Thiết kế khả năng mở rộng luật giải đấu",
        "3.17. Thiết kế trải nghiệm quan sát cho khán giả": "3.18. Thiết kế trải nghiệm quan sát cho khán giả",
        "3.18. Design System và Style Guide": "3.12. Design System và Style Guide",
        "2.7. Docker, Vercel và triển khai ứng dụng": "2.7. Docker, VPS và triển khai ứng dụng",
    }
    for old, new in replacements.items():
        for paragraph in doc.paragraphs:
            text = paragraph_text(paragraph)
            if text == old or text.startswith(old + " "):
                set_para_text(paragraph, text.replace(old, new, 1))

    for paragraph in doc.paragraphs:
        text = paragraph_text(paragraph)
        if text == "Docker/Vercel":
            set_para_text(paragraph, "Docker/VPS")
        elif text == "2.7. Docker, Vercel và triển khai ứng dụng ..................................":
            set_para_text(paragraph, "2.7. Docker, VPS và triển khai ứng dụng .....................................")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() == "Docker/Vercel":
                    set_cell(cell, "Docker/VPS")


def improve_technology_and_deploy(doc: Document) -> None:
    replace_contains(
        doc,
        "Vercel phù hợp với ứng dụng Next.js",
        "Trong phạm vi hiện tại, ứng dụng đã có Dockerfile multi-stage và docker-compose để đóng gói, chạy container và cấu hình biến môi trường. Nếu triển khai theo yêu cầu VPS, cần bổ sung reverse proxy, domain và SSL/TLS; các thông tin này không nên được mô tả như đã hoàn tất khi chưa có URL kiểm chứng.",
    )
    replace_exact(
        doc,
        "Docker compose trong dự án cho phép cấu hình port, environment và volume dữ liệu. Điều này thuận tiện khi cần chạy ứng dụng theo dạng container, đặc biệt trong trường hợp muốn triển khai ở VPS hoặc môi trường không phải Vercel.",
        "Docker compose trong dự án cho phép cấu hình port, environment và biến môi trường. Điều này thuận tiện khi cần chạy ứng dụng theo dạng container, đặc biệt trong trường hợp muốn triển khai trên VPS hoặc môi trường máy chủ riêng có reverse proxy.",
    )
    replace_exact(doc, "[8] Vercel Documentation.", "[8] Docker Compose Documentation.")

    anchor = find_paragraph(doc, "2.8. Công cụ AI trong quá trình hỗ trợ phát triển")
    insert_paragraph_after(
        anchor,
        "Để đáp ứng yêu cầu minh chứng, báo cáo cần có phụ lục liệt kê prompt đã sử dụng, mục đích của từng prompt và kết quả thu được. Các prompt phải phản ánh đúng quá trình phát triển thực tế, không nên ghi chung chung hoặc tạo lại sau khi hoàn thành dự án.",
    )

    anchor = find_paragraph(doc, "4.15. Chi tiết triển khai và cấu hình môi trường")
    insert_table_after(
        doc,
        anchor,
        ["Hạng mục", "Hiện trạng trong báo cáo", "Cần bổ sung trước khi nộp"],
        [
            ["Dockerfile", "Đã có Dockerfile multi-stage cho Next.js standalone", "Chèn đoạn cấu hình chính hoặc ảnh build thành công"],
            ["Docker Compose", "Đã có docker-compose expose cổng 8000 và cấu hình env", "Chèn lệnh chạy và kết quả `docker compose ps`"],
            ["Supabase", "Đã dùng PostgreSQL, Auth và Realtime", "Bổ sung ảnh cấu hình Supabase hoặc mô tả bảng/publication"],
            ["VPS/domain/SSL", "Chưa có URL thật trong báo cáo", "[Cần bổ sung: domain HTTPS, ảnh chứng chỉ SSL, sơ đồ reverse proxy nếu có]"],
            ["GitHub", "Chưa có phần minh chứng repository", "[Cần bổ sung: URL repository, ảnh commit history, trạng thái public/private]"],
        ],
        "Bảng 4.3. Checklist triển khai và minh chứng cần bổ sung",
    )


def add_database_and_api_details(doc: Document) -> None:
    db_anchor = find_paragraph(doc, "3.9. Thiết kế cơ sở dữ liệu theo vòng đời trận đấu")
    insert_table_after(
        doc,
        db_anchor,
        ["Bảng", "Trường tiêu biểu", "Khóa/ràng buộc", "Vai trò"],
        [
            ["User", "id, email, name, role, createdAt, updatedAt", "id là khóa chính; email unique", "Lưu tài khoản đồng bộ với Supabase Auth và phân quyền ADMIN/REFEREE"],
            ["Room", "id, code, status, hostUserId, blueClientId, redClientId, costPerPoint", "code unique; hostUserId liên kết User", "Lưu phiên thi đấu, trạng thái phòng và người tham gia"],
            ["DraftLog", "id, roomId, player, action, characterId, turnNumber", "roomId liên kết Room; unique roomId-characterId", "Ghi lại từng lượt ban/pick để dựng lại trạng thái draft"],
            ["CharacterBuild", "id, roomId, player, characterId, rarity, consLevel, weaponRarity, totalCost", "unique roomId-player-characterId", "Lưu build sau draft và phục vụ tính Cost/Handicap"],
            ["Tournament", "id, slug, name, format, status, maxTeams, organizerId", "slug unique", "Lưu thông tin giải đấu và cấu hình chung"],
            ["TournamentMatch", "id, tournamentId, round, matchNumber, roomCode, winnerParticipantId", "unique tournamentId-round-matchNumber", "Quản lý trận trong bracket và liên kết với phòng draft"],
        ],
        "Bảng 3.10. Thiết kế chi tiết một số bảng dữ liệu chính",
    )
    insert_paragraph_after(
        db_anchor,
        "Các bảng trên được rút ra từ Prisma schema của dự án. Trước khi nộp chính thức, nên bổ sung ERD dạng hình ảnh để thể hiện rõ quan hệ một-nhiều giữa Room, DraftLog, CharacterBuild, Tournament và TournamentMatch. [Cần bổ sung thủ công: ảnh ERD hoặc sơ đồ quan hệ dữ liệu.]",
    )

    api_anchor = find_paragraph(doc, "3.10. Thiết kế API theo nhóm nghiệp vụ")
    insert_table_after(
        doc,
        api_anchor,
        ["Method", "Endpoint", "Chức năng", "Quyền/điều kiện"],
        [
            ["POST", "/api/room", "Tạo phòng đấu mới", "Người dùng đã đăng nhập/referee"],
            ["GET", "/api/room/[code]", "Lấy snapshot phòng, logs, builds, messages", "Client thuộc phòng hoặc caster hợp lệ"],
            ["POST", "/api/room/[code]/join", "Tham gia slot Blue/Red", "ClientId hợp lệ, slot còn trống"],
            ["DELETE", "/api/room/[code]/join", "Rời khỏi phòng", "Client đang giữ slot"],
            ["POST", "/api/draft", "Gửi lượt ban/pick", "Đúng đội, đúng lượt, phòng không pause"],
            ["POST", "/api/draft/undo", "Hoàn tác lượt draft", "Host/referee hợp lệ"],
            ["GET/POST", "/api/cost-catalog", "Đọc hoặc cập nhật bảng cost", "GET public; POST cần user và quyền phù hợp"],
            ["GET/POST", "/api/tournaments", "Danh sách hoặc tạo giải đấu", "POST yêu cầu ADMIN"],
            ["POST/DELETE", "/api/tournaments/[slug]/participants", "Đăng ký hoặc hủy participant", "UID/thông tin đội hợp lệ"],
            ["POST", "/api/tournaments/[slug]/bracket", "Tạo/cập nhật bracket", "Admin/organizer theo service"],
            ["POST", "/api/tournaments/[slug]/matches/[matchId]/result", "Ghi nhận kết quả trận", "Quyền quản lý giải/trận"],
            ["GET/POST", "/api/chat", "Đọc/gửi tin nhắn phòng", "Thành viên phòng hợp lệ"],
        ],
        "Bảng 3.11. Danh sách API tiêu biểu của hệ thống",
    )
    insert_paragraph_after(
        api_anchor,
        "Khi hoàn thiện báo cáo cuối cùng, mỗi API trọng tâm nên có thêm ví dụ request body, response body và mã lỗi thường gặp. [Cần bổ sung thủ công: ảnh Postman/cURL hoặc log phản hồi API nếu giảng viên yêu cầu minh chứng.]",
    )


def fix_ui_placeholders(doc: Document) -> None:
    replacements = {
        "Trước hình 4.1. giao diện trang chủ và điều hướng chính, báo cáo cần làm rõ màn hình này nằm ở bước nào trong luồng sử dụng và người dùng nhận được thông tin gì từ màn hình đó.": "Trang chủ là điểm vào đầu tiên của hệ thống, giúp người dùng hiểu nhanh mục đích của website và chuyển đến các luồng chính như tạo phòng, vào lobby, xem giải đấu hoặc sử dụng công cụ hỗ trợ. [Cần bổ sung thủ công: chèn ảnh chụp trang chủ đang chạy thực tế.]",
        "Trước hình 4.2. giao diện phòng ban/pick, báo cáo cần làm rõ màn hình này nằm ở bước nào trong luồng sử dụng và người dùng nhận được thông tin gì từ màn hình đó.": "Giao diện phòng Ban/Pick là màn hình trung tâm trong quá trình thi đấu. Người chơi cần nhìn thấy đội của mình, lượt hiện tại, nhân vật còn khả dụng, nhân vật đã bị cấm/chọn và trạng thái thời gian. [Cần bổ sung thủ công: chèn ảnh phòng draft với dữ liệu mẫu.]",
        "Trước hình 4.3. giao diện khai báo build và tính cost, báo cáo cần làm rõ màn hình này nằm ở bước nào trong luồng sử dụng và người dùng nhận được thông tin gì từ màn hình đó.": "Màn hình khai báo build xuất hiện sau khi draft hoàn tất. Tại đây, mỗi đội nhập thông tin rarity, cung mệnh và vũ khí cho các nhân vật đã chọn để hệ thống tính Cost. [Cần bổ sung thủ công: chèn ảnh màn hình build phase.]",
        "Trước hình 4.4. giao diện kết quả và time handicap, báo cáo cần làm rõ màn hình này nằm ở bước nào trong luồng sử dụng và người dùng nhận được thông tin gì từ màn hình đó.": "Trang kết quả tổng hợp dữ liệu cuối trận, bao gồm đội hình hai bên, tổng Cost, chênh lệch Cost và Time Handicap. Màn hình này giúp trọng tài công bố kết quả minh bạch hơn. [Cần bổ sung thủ công: chèn ảnh trang result với dữ liệu mẫu.]",
        "Trước hình 4.5. giao diện giải đấu và bracket, báo cáo cần làm rõ màn hình này nằm ở bước nào trong luồng sử dụng và người dùng nhận được thông tin gì từ màn hình đó.": "Giao diện giải đấu mở rộng website từ một trận đơn lẻ thành công cụ quản lý sự kiện. Trang này hiển thị danh sách giải, participant, bracket và liên kết trận đấu với phòng Ban/Pick. [Cần bổ sung thủ công: chèn ảnh trang tournament/bracket.]",
        "Sau hình, có thể thấy giao diện không chỉ đóng vai trò minh họa mà còn là minh chứng cho việc các yêu cầu đã phân tích ở Chương 3 được hiện thực thành sản phẩm có thể sử dụng.": "Ảnh giao diện cần được chụp từ môi trường chạy thật để chứng minh yêu cầu đã phân tích ở Chương 3 được hiện thực thành chức năng có thể sử dụng.",
    }
    for old, new in replacements.items():
        replace_exact(doc, old, new)


def improve_testing_and_limitations(doc: Document) -> None:
    anchor = find_paragraph(doc, "4.8. Kiểm thử chức năng")
    insert_paragraph_after(
        anchor,
        "Ngoài bảng kiểm thử chức năng, bản nộp chính thức nên đính kèm minh chứng chạy lệnh `npm run build`, `npm run lint`, `docker compose ps` hoặc ảnh màn hình thao tác thực tế. Các kết quả kiểm thử trong bảng dưới đây đang ở mức mô tả; phần minh chứng thực nghiệm cần được bổ sung thủ công nếu giảng viên yêu cầu.",
    )

    anchor = find_paragraph(doc, "5.4. Hạn chế còn tồn tại")
    insert_paragraph_after(
        anchor,
        "So với checklist công nghệ của học phần, báo cáo cần trình bày rõ các phần chưa hoàn thiện hoặc chưa có minh chứng: Server Actions chưa được mô tả trong sản phẩm, RLS ở mức database chưa có policy SQL minh chứng, Supabase Storage/file upload chưa được triển khai thành chức năng chính, và deployment domain/SSL cần URL thật để kiểm chứng. Những nội dung này không nên ghi là đã hoàn thành nếu chưa có mã nguồn hoặc ảnh minh chứng kèm theo.",
    )


def add_appendices(doc: Document) -> None:
    doc.add_page_break()
    append_paragraph(doc, "PHỤ LỤC A: MINH CHỨNG SỬ DỤNG CÔNG CỤ AI", style="Heading 1")
    append_paragraph(
        doc,
        "Phụ lục này dùng để đáp ứng yêu cầu minh chứng việc sử dụng công cụ AI trong quá trình phát triển. Các dòng dưới đây cần được thay bằng prompt thật đã sử dụng, mục đích sử dụng và kết quả thực tế thu được.",
    )
    table = doc.add_table(rows=1, cols=5)
    table.alignment = 1
    apply_borders(table)
    for i, header in enumerate(["STT", "Công cụ", "Prompt đã sử dụng", "Mục đích", "Kết quả/ghi chú"]):
        set_cell(table.rows[0].cells[i], header, bold=True)
    for row in [
        ["1", "[Cần bổ sung]", "[Cần bổ sung: prompt thật]", "Phân tích yêu cầu/chức năng", "[Cần bổ sung]"],
        ["2", "[Cần bổ sung]", "[Cần bổ sung: prompt thật]", "Thiết kế CSDL/API", "[Cần bổ sung]"],
        ["3", "[Cần bổ sung]", "[Cần bổ sung: prompt thật]", "Rà soát lỗi giao diện", "[Cần bổ sung]"],
        ["4", "[Cần bổ sung]", "[Cần bổ sung: prompt thật]", "Hỗ trợ viết/tối ưu code", "[Cần bổ sung]"],
        ["5", "[Cần bổ sung]", "[Cần bổ sung: prompt thật]", "Viết hoặc chỉnh báo cáo", "[Cần bổ sung]"],
        ["6", "[Cần bổ sung]", "[Cần bổ sung: prompt thật]", "Kiểm tra checklist nộp bài", "[Cần bổ sung]"],
    ]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    append_paragraph(
        doc,
        "Bảng A.1. Minh chứng sử dụng công cụ AI trong quá trình phát triển",
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.add_page_break()
    append_paragraph(doc, "PHỤ LỤC B: MINH CHỨNG TRIỂN KHAI VÀ MÃ NGUỒN", style="Heading 1")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = 1
    apply_borders(table)
    for i, header in enumerate(["Nội dung", "Minh chứng cần chèn", "Trạng thái"]):
        set_cell(table.rows[0].cells[i], header, bold=True)
    for row in [
        ["GitHub repository", "[Cần bổ sung: URL repository và ảnh commit history]", "Chưa có trong báo cáo"],
        ["Docker build/run", "[Cần bổ sung: ảnh/lệnh docker compose up và docker compose ps]", "Cần bổ sung minh chứng"],
        ["Domain HTTPS", "[Cần bổ sung: URL domain thật và ảnh chứng chỉ SSL]", "Cần bổ sung nếu đã deploy"],
        ["Supabase", "[Cần bổ sung: ảnh project Supabase, Auth, Database hoặc Realtime publication]", "Cần bổ sung minh chứng"],
        ["Sơ đồ ERD", "[Cần bổ sung: ảnh ERD tạo từ Prisma schema hoặc công cụ thiết kế]", "Chưa có ảnh trong DOCX"],
    ]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    append_paragraph(
        doc,
        "Bảng B.1. Checklist minh chứng cần bổ sung thủ công",
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    shutil.copyfile(SOURCE, OUTPUT)
    doc = Document(OUTPUT)

    add_frontmatter_notes(doc)
    fix_section_numbering(doc)
    improve_technology_and_deploy(doc)
    add_database_and_api_details(doc)
    fix_ui_placeholders(doc)
    improve_testing_and_limitations(doc)
    add_appendices(doc)

    doc.save(OUTPUT)
    print(f"Improved report written to: {OUTPUT}")


if __name__ == "__main__":
    main()
