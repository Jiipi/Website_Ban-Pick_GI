from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "BaoCao_Website_BanPick_Genshin.docx"
TITLE = "PHÂN TÍCH, THIẾT KẾ VÀ XÂY DỰNG WEBSITE HỖ TRỢ BAN/PICK NHÂN VẬT GENSHIN IMPACT"

FIGURES = [
    "Hình 3.1. Sơ đồ Use Case tổng quan của hệ thống",
    "Hình 3.2. Kiến trúc phân tầng của website",
    "Hình 3.3. Mô hình dữ liệu chính theo vòng đời trận đấu",
    "Hình 3.4. Luồng xử lý một lượt Ban/Pick",
    "Hình 4.1. Giao diện trang chủ và điều hướng chính",
    "Hình 4.2. Giao diện phòng Ban/Pick",
    "Hình 4.3. Giao diện khai báo build và tính Cost",
    "Hình 4.4. Giao diện kết quả và Time Handicap",
    "Hình 4.5. Giao diện giải đấu và bracket",
]

TABLES = [
    "Bảng 1.1. Phạm vi thực hiện của đề tài",
    "Bảng 2.1. Vai trò của các công nghệ trong dự án",
    "Bảng 3.1. Tác nhân và nhu cầu sử dụng hệ thống",
    "Bảng 3.2. Yêu cầu chức năng chính",
    "Bảng 3.3. Yêu cầu phi chức năng",
    "Bảng 3.4. Use Case chính của hệ thống",
    "Bảng 3.5. Nhóm dữ liệu chính trong cơ sở dữ liệu",
    "Bảng 3.6. Nhóm API theo nghiệp vụ",
    "Bảng 4.1. Cấu trúc mã nguồn theo tầng",
    "Bảng 4.2. Kết quả kiểm thử chức năng",
    "Bảng 5.1. Đối chiếu mục tiêu ban đầu và kết quả đạt được",
]

SECTIONS = {
    "CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI": [
        "1.1. Bối cảnh cộng đồng Genshin Impact và nhu cầu tổ chức Ban/Pick",
        "1.2. Thực trạng tổ chức thủ công và các hạn chế",
        "1.3. Lý do chọn đề tài",
        "1.4. Mục tiêu đề tài",
        "1.5. Phạm vi đề tài",
        "1.6. Ý nghĩa thực tiễn",
        "1.7. Nhìn nhận đề tài như một sản phẩm cộng đồng",
        "1.8. Sự khác biệt giữa công cụ hỗ trợ và trò chơi",
        "1.9. Tổng hợp vấn đề cần giải quyết",
        "1.10. Định hướng phương pháp thực hiện",
    ],
    "CHƯƠNG 2: CÔNG NGHỆ VÀ CÔNG CỤ PHÁT TRIỂN": [
        "2.1. Next.js App Router và React",
        "2.2. TypeScript trong kiểm soát dữ liệu",
        "2.3. Tailwind CSS và định hướng giao diện thi đấu",
        "2.4. Supabase Realtime và PostgreSQL",
        "2.5. Prisma ORM và thiết kế dữ liệu",
        "2.6. Tích hợp dữ liệu ngoài Enka và Genshin.dev",
        "2.7. Docker, Vercel và triển khai ứng dụng",
        "2.8. Công cụ AI trong quá trình hỗ trợ phát triển",
        "2.9. Liên hệ công nghệ với mục tiêu đề tài",
        "2.10. Đánh giá sự phù hợp của stack công nghệ",
        "2.11. Vai trò của công cụ phát triển trong chất lượng sản phẩm",
    ],
    "CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG": [
        "3.1. Quy trình nghiệp vụ tổng quát",
        "3.2. Tác nhân hệ thống",
        "3.3. Yêu cầu chức năng",
        "3.4. Yêu cầu phi chức năng",
        "3.5. Use Case chính",
        "3.6. Thiết kế luật Ban/Pick và kiểm tra phía server",
        "3.7. Thiết kế Cost và Time Handicap",
        "3.8. Kiến trúc phân tầng",
        "3.9. Thiết kế cơ sở dữ liệu theo vòng đời trận đấu",
        "3.10. Thiết kế API theo nhóm nghiệp vụ",
        "3.11. Thiết kế giao diện và trải nghiệm người dùng",
        "3.12. Tính liên tục của dữ liệu trong trận đấu",
        "3.13. Tách rõ chức năng lõi và chức năng mở rộng",
        "3.14. Thiết kế phân quyền trong phòng đấu",
        "3.15. Thiết kế trạng thái và vòng đời phòng",
        "3.16. Thiết kế khả năng mở rộng luật giải đấu",
        "3.17. Thiết kế trải nghiệm quan sát cho khán giả",
    ],
    "CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG": [
        "4.1. Tổ chức mã nguồn và cách phân lớp được áp dụng",
        "4.2. Xây dựng nền tảng phòng đấu và phân quyền",
        "4.3. Xây dựng draft engine realtime",
        "4.4. Xây dựng khai báo build, Cost Catalog và Time Handicap",
        "4.5. Hoàn thiện trải nghiệm trận đấu",
        "4.6. Mở rộng cộng đồng và giải đấu",
        "4.7. Triển khai cơ sở dữ liệu và ứng dụng",
        "4.8. Kiểm thử chức năng",
        "4.9. Kết quả giao diện người dùng",
        "4.10. Đánh giá quá trình hiện thực chức năng lõi",
        "4.11. Đánh giá khả năng vận hành thực tế",
        "4.12. Chi tiết xây dựng giao diện Draft Board",
        "4.13. Chi tiết xây dựng trang kết quả",
        "4.14. Chi tiết xây dựng tournament",
        "4.15. Chi tiết triển khai và cấu hình môi trường",
    ],
    "CHƯƠNG 5: TỔNG KẾT, ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN": [
        "5.1. Kết quả đạt được theo luồng end-to-end",
        "5.2. Đối chiếu mục tiêu ban đầu và mức độ hoàn thành",
        "5.3. Ưu điểm của hệ thống",
        "5.4. Hạn chế còn tồn tại",
        "5.5. Hướng phát triển",
        "5.6. Bài học rút ra và kết luận",
        "5.7. Kết luận theo giá trị sản phẩm",
        "5.8. Đánh giá mức độ phù hợp với báo cáo tổng kết",
        "5.9. Định hướng kiểm thử với người dùng thật",
        "5.10. Khả năng áp dụng trong cộng đồng",
    ],
}


def chapter_no(chapter_title: str) -> int:
    return int(chapter_title.split(" ")[1].rstrip(":"))


def next_subsection_index(chapter_title: str, label: str) -> int:
    return SECTIONS[chapter_title].index(label) + 1


def set_font(run, size=13, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def setup(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)
    for name, size in [("Normal", 13), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        if name.startswith("Heading"):
            style.font.bold = True
    doc.styles["Normal"].paragraph_format.line_spacing = 1.65
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Trang ")
    set_font(r, 11)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def para(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, first=True, bold=False, italic=False, size=13):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    if first and text:
        p.paragraph_format.first_line_indent = Cm(1)
    r = p.add_run(text)
    set_font(r, size, bold, italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        set_font(r, {1: 16, 2: 14, 3: 13}.get(level, 13), True)
    return p


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    set_font(r, 10.5, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def borders(tbl):
    pr = tbl._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:color"), "666666")
        b.append(e)
    pr.append(b)


def table(doc, caption, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(t)
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], val)
    para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, italic=True, size=12)


def figure(doc, caption, lines):
    t = doc.add_table(rows=1, cols=1)
    borders(t)
    text = "\n" + "\n".join(lines) + "\n"
    set_cell(t.rows[0].cells[0], text)
    for p in t.rows[0].cells[0].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, italic=True, size=12)


def section(doc, title, paragraphs):
    heading(doc, title, 2)
    for x in paragraphs:
        para(doc, x)


def extra_section(doc, title, paragraphs):
    heading(doc, title, 2)
    for p in paragraphs:
        para(doc, p)


def cover(doc):
    for line in ["TRƯỜNG ĐẠI HỌC ĐÀ LẠT", "KHOA CÔNG NGHỆ THÔNG TIN"]:
        para(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, bold=True, size=14)
    para(doc, "\n\nBÁO CÁO TỔNG KẾT HỌC PHẦN", align=WD_ALIGN_PARAGRAPH.CENTER, first=False, bold=True, size=18)
    para(doc, "HỌC PHẦN: [TÊN HỌC PHẦN]", align=WD_ALIGN_PARAGRAPH.CENTER, first=False, bold=True, size=14)
    para(doc, "\nĐỀ TÀI", align=WD_ALIGN_PARAGRAPH.CENTER, first=False, bold=True, size=16)
    para(doc, TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, bold=True, size=16)
    for line in ["\n\nGVHD: [TÊN GIẢNG VIÊN HƯỚNG DẪN]", "SVTH: [HỌ VÀ TÊN SINH VIÊN]", "MSSV: [MÃ SỐ SINH VIÊN]", "LỚP: [TÊN LỚP]"]:
        para(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, first=False)
    para(doc, "\n\n\nLâm Đồng, tháng 05 năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER, first=False)
    doc.add_page_break()


def front(doc):
    heading(doc, "NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN", 1)
    for _ in range(22):
        para(doc, "." * 105, first=False)
    para(doc, "Lâm Đồng, ngày ..... tháng ..... năm 2026", align=WD_ALIGN_PARAGRAPH.RIGHT, first=False)
    para(doc, "Giảng viên hướng dẫn", align=WD_ALIGN_PARAGRAPH.RIGHT, first=False)
    doc.add_page_break()

    heading(doc, "LỜI CẢM ƠN", 1)
    for text in [
        "Em xin chân thành cảm ơn quý thầy cô Khoa Công nghệ Thông tin đã truyền đạt cho em những kiến thức nền tảng về phân tích thiết kế hệ thống, lập trình web và triển khai phần mềm. Những kiến thức đó là cơ sở quan trọng để em thực hiện đề tài và hoàn thành báo cáo tổng kết này.",
        "Em xin gửi lời cảm ơn đến giảng viên hướng dẫn đã định hướng cách tiếp cận đề tài, giúp em nhìn website không chỉ như một tập hợp giao diện, mà là một hệ thống có người dùng, quy trình nghiệp vụ, dữ liệu, phân quyền và khả năng vận hành thực tế.",
        "Trong quá trình thực hiện, em có cơ hội vận dụng nhiều công nghệ hiện đại như Next.js, TypeScript, Supabase, Prisma, Tailwind CSS và Docker. Quan trọng hơn, em học được cách chuyển một nhu cầu cụ thể của cộng đồng thành một sản phẩm có quy trình, có dữ liệu và có thể mở rộng.",
        "Do thời gian và kinh nghiệm còn hạn chế, sản phẩm cũng như báo cáo chắc chắn vẫn còn thiếu sót. Em rất mong nhận được góp ý của quý thầy cô để tiếp tục hoàn thiện đề tài trong các phiên bản tiếp theo. Em xin chân thành cảm ơn!",
    ]:
        para(doc, text)
    doc.add_page_break()

    heading(doc, "MỤC LỤC", 1)
    for chapter, subs in SECTIONS.items():
        para(doc, chapter + " " + "." * 42, first=False, bold=True, size=12)
        for sub in subs:
            para(doc, sub + " " + "." * max(12, 76 - len(sub)), first=False, size=12)
    para(doc, "TÀI LIỆU THAM KHẢO " + "." * 55, first=False, bold=True, size=12)
    doc.add_page_break()

    heading(doc, "DANH MỤC HÌNH ẢNH", 1)
    for item in FIGURES:
        para(doc, item + " " + "." * 45, first=False, size=12)
    doc.add_page_break()
    heading(doc, "DANH MỤC BẢNG", 1)
    for item in TABLES:
        para(doc, item + " " + "." * 45, first=False, size=12)
    doc.add_page_break()


def chapter1(doc):
    heading(doc, "CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI", 1)
    section(doc, SECTIONS["CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI"][0], [
        "Genshin Impact là một trò chơi có cộng đồng người chơi lớn, trong đó hoạt động thi đấu giao hữu và thi đấu La Hoàn Cảnh Giới thường được tổ chức bởi các nhóm cộng đồng. Khác với hình thức chơi cá nhân, một trận đấu cộng đồng cần có luật rõ ràng, người điều phối và cách ghi nhận kết quả minh bạch để người chơi hai bên cảm thấy công bằng.",
        "Trong các trận đấu dạng Ban/Pick, chiến thuật không chỉ nằm ở kỹ năng điều khiển nhân vật mà còn ở việc lựa chọn đội hình. Người chơi phải cân nhắc nhân vật nào nên bị cấm, nhân vật nào nên được chọn, đội hình nào tạo lợi thế và lợi thế đó có cần quy đổi thành thời gian bù trừ hay không.",
        "Điểm đặc thù của cộng đồng Genshin là độ mạnh tài khoản khác nhau khá lớn. Nhân vật 5 sao, cung mệnh cao hoặc vũ khí 5 sao có thể tạo lợi thế đáng kể. Vì vậy, nhiều giải đấu cộng đồng áp dụng luật Cost và Time Handicap để cân bằng trận đấu, thay vì chỉ để hai đội tự do chọn đội hình mạnh nhất.",
        "Một trận La Hoàn Cảnh Giới đối kháng giữa hai người chơi thường gồm hai bước: chọn đội hình và thi đấu hoàn thành thử thách. Bước thứ hai diễn ra trong game, còn bước đầu tiên gần như được tổ chức ngoài game. Đây là khoảng trống mà các công cụ cộng đồng đang cố gắng lấp đầy, và cũng là lý do website Ban/Pick có chỗ đứng riêng.",
        "Ngoài ra, hoạt động cộng đồng còn ngày càng có yếu tố livestream và khán giả. Khi trận đấu được phát trực tiếp, người xem cần nhìn thấy đội nào đang ban, đội nào đang pick, nhân vật nào đã xuất hiện và kết quả handicap ra sao. Một website chuyên dụng dễ đáp ứng nhu cầu này hơn so với việc trọng tài chia sẻ ảnh chụp màn hình hoặc bảng tính.",
        "Từ bối cảnh đó, nhu cầu về một website chuyên dụng xuất hiện khá rõ: trọng tài cần công cụ điều phối, người chơi cần giao diện thao tác đúng lượt, còn khán giả cần màn hình theo dõi dễ hiểu. Đây là nền tảng thực tế để hình thành đề tài website Ban/Pick Genshin Impact.",
    ])
    section(doc, SECTIONS["CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI"][1], [
        "Trước khi có một công cụ chuyên dụng, nhiều buổi Ban/Pick được tổ chức bằng cách kết hợp Discord, bảng tính, ảnh nhân vật và tin nhắn xác nhận. Cách làm này có ưu điểm là dễ bắt đầu, nhưng khi số người tham gia tăng hoặc trận đấu có khán giả theo dõi, nó nhanh chóng bộc lộ hạn chế.",
        "Hạn chế đầu tiên là việc ghi nhận lượt dễ sai. Trọng tài phải nhớ lượt hiện tại, kiểm tra nhân vật đã bị cấm hay chưa và cập nhật trạng thái cho cả hai đội. Chỉ cần bỏ sót một lượt hoặc nhập nhầm tên nhân vật, toàn bộ quá trình draft có thể phải dừng lại để đối chiếu.",
        "Hạn chế thứ hai là tính Cost và Handicap thường phải làm thủ công. Nếu mỗi đội có nhiều nhân vật 5 sao, nhiều cấp cung mệnh và vũ khí khác nhau, việc tính bằng tay hoặc bằng bảng tính dễ mất thời gian và tạo tranh luận khi kết quả không rõ nguồn gốc.",
        "Hạn chế thứ ba là khó lưu trữ và xem lại. Sau trận đấu, dữ liệu thường nằm rải rác trong tin nhắn, ảnh chụp hoặc bảng tính riêng. Khi muốn thống kê nhân vật được chọn nhiều, trận nào thuộc giải nào hoặc kết quả handicap ra sao, ban tổ chức phải tổng hợp lại thủ công.",
    ])
    section(doc, SECTIONS["CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI"][2], [
        "Xuất phát từ những vấn đề trên, đề tài được lựa chọn với mục tiêu xây dựng một website giúp chuẩn hóa quy trình Ban/Pick nhân vật Genshin Impact. Website không thay thế trò chơi, cũng không can thiệp vào máy chủ của nhà phát hành, mà đóng vai trò là công cụ hỗ trợ tổ chức trận đấu cộng đồng.",
        "Lý do chọn đề tài còn đến từ tính phù hợp với học phần công nghệ mới. Sản phẩm cần giao diện web hiện đại, xử lý realtime, lưu trữ dữ liệu, phân quyền người dùng, tích hợp dữ liệu ngoài và có khả năng triển khai thực tế. Đây là những nội dung gắn chặt với xu hướng phát triển phần mềm hiện nay.",
        "Ngoài giá trị kỹ thuật, đề tài có bối cảnh sử dụng cụ thể. Người thực hiện không chỉ xây dựng các trang CRUD thông thường, mà phải phân tích một quy trình thi đấu có trạng thái, có luật, có ràng buộc và có nhiều vai trò người dùng. Điều này giúp báo cáo có chiều sâu hơn so với một website giới thiệu đơn giản.",
    ])
    section(doc, SECTIONS["CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI"][3], [
        "Mục tiêu nghiệp vụ của đề tài là xây dựng được luồng thi đấu end-to-end: trọng tài tạo phòng, người chơi tham gia theo đội, hai đội thực hiện Ban/Pick đúng lượt, khai báo build, hệ thống tính Cost và Time Handicap, sau đó hiển thị kết quả rõ ràng.",
        "Mục tiêu kỹ thuật là áp dụng kiến trúc có tổ chức, tách phần giao diện, phần xử lý use case, phần luật nghiệp vụ và phần truy cập dữ liệu. Các luật quan trọng như kiểm tra lượt, kiểm tra quyền và tính Cost không được đặt rải rác trong giao diện mà cần có vị trí rõ ràng trong mã nguồn.",
        "Mục tiêu sản phẩm là tạo ra một website đủ trực quan để người chơi dễ sử dụng, đủ minh bạch để trọng tài tin tưởng và đủ mở rộng để phát triển thành nền tảng hỗ trợ giải đấu cộng đồng trong tương lai.",
        "Bên cạnh ba nhóm mục tiêu trên, đề tài cũng đặt mục tiêu học tập rõ ràng. Người thực hiện cần vận dụng được kiến thức về phân tích thiết kế, lập trình web full-stack, tổ chức cơ sở dữ liệu, làm việc với dịch vụ realtime và quy trình triển khai. Việc đặt mục tiêu học tập song song với mục tiêu sản phẩm giúp đánh giá kết quả không chỉ dựa trên tính năng mà còn dựa trên năng lực đã rèn luyện.",
        "Các mục tiêu này có quan hệ chặt chẽ. Nếu chỉ đạt mục tiêu kỹ thuật mà không có nghiệp vụ rõ ràng, sản phẩm sẽ chỉ là demo công nghệ. Ngược lại, nếu chỉ đạt mục tiêu nghiệp vụ mà mã nguồn rối, dự án khó được mở rộng. Báo cáo vì vậy luôn cố gắng đối chiếu cả hai khía cạnh khi đánh giá mức độ hoàn thành.",
    ])
    table(doc, TABLES[0], ["Nội dung", "Trong phạm vi", "Ngoài phạm vi hiện tại"], [
        ["Ban/Pick", "Tạo phòng, vào đội, cấm/chọn theo lượt", "Tự động lấy kết quả trong game"],
        ["Build/Cost", "Khai báo build, tính Cost, tính Handicap", "Xác minh anti-cheat hoàn chỉnh"],
        ["Cộng đồng", "Lobby, chat, profile, lịch sử, tournament cơ bản", "Vận hành giải lớn với tải cao"],
        ["Triển khai", "Có Docker, cấu hình Supabase/PostgreSQL", "Giám sát production chuyên sâu"],
    ])
    section(doc, SECTIONS["CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI"][4], [
        "Trong phạm vi thực hiện, đề tài tập trung vào website hỗ trợ tổ chức Ban/Pick và các chức năng liên quan trực tiếp đến trận đấu. Những phần như lobby, tournament, profile hoặc overlay được xem là phần mở rộng để sản phẩm có định hướng cộng đồng rõ hơn.",
        "Đề tài chưa đặt mục tiêu xác minh toàn bộ dữ liệu trong game như một hệ thống anti-cheat hoàn chỉnh. Việc khai báo build vẫn cần sự trung thực của người chơi hoặc sự kiểm tra của trọng tài. Đây là giới hạn hợp lý vì Genshin Impact không cung cấp đầy đủ API chính thức cho mọi thông tin cần xác minh.",
        "Website cũng chưa được đánh giá với quy mô hàng trăm phòng thi đấu đồng thời. Vì vậy, phần triển khai hiện tại phù hợp với đồ án và các buổi thi đấu cộng đồng quy mô nhỏ đến vừa, còn vận hành production lớn cần thêm kiểm thử tải, giám sát và tối ưu.",
    ])
    section(doc, SECTIONS["CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI"][5], [
        "Ý nghĩa thực tiễn của đề tài nằm ở việc chuyển một quy trình tổ chức thủ công sang nền tảng web có dữ liệu tập trung. Khi mỗi lượt Ban/Pick được lưu lại, kết quả không còn phụ thuộc hoàn toàn vào ghi chú rời rạc của trọng tài.",
        "Đối với người chơi, website giúp họ tập trung vào chiến thuật thay vì phải tự nhớ nhân vật nào đã được chọn. Đối với khán giả, giao diện phân đội, lượt hiện tại và trang kết quả giúp việc theo dõi trận đấu trở nên rõ ràng hơn.",
        "Từ các mục tiêu và phạm vi trên, chương tiếp theo trình bày các công nghệ được lựa chọn để hiện thực website. Mỗi công nghệ không được chọn chỉ vì phổ biến, mà vì nó giải quyết một phần yêu cầu cụ thể của sản phẩm.",
    ])
    add_extras = [
        (SECTIONS["CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI"][6], [
            "Nếu chỉ xem website là nơi bấm chọn nhân vật, đề tài sẽ trở nên khá hẹp. Tuy nhiên, khi đặt trong bối cảnh tổ chức giải đấu cộng đồng, sản phẩm phải xử lý nhiều yếu tố hơn: quyền điều phối, trạng thái phòng, lịch sử trận, dữ liệu đội hình và trải nghiệm người xem.",
            "Cách nhìn này giúp báo cáo có mạch rõ ràng hơn. Mỗi chức năng được nhắc đến đều phải trả lời câu hỏi: nó phục vụ ai trong trận đấu và giúp giảm vấn đề thủ công nào. Nhờ đó, nội dung không bị lan man sang những phần kỹ thuật không cần thiết.",
            "Khi đề tài được đặt trong bối cảnh cộng đồng, vai trò của trọng tài, người chơi và khán giả không còn là vai trò trừu tượng. Mỗi vai trò có một nhu cầu cụ thể, và website phải lần lượt giải quyết các nhu cầu đó để được xem là sản phẩm có giá trị thực tế.",
        ]),
        (SECTIONS["CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI"][7], [
            "Website không can thiệp vào cơ chế gameplay của Genshin Impact. Nó chỉ hỗ trợ phần tổ chức bên ngoài: tạo luật, ghi nhận lựa chọn, tính bù trừ thời gian và trình bày kết quả. Việc phân biệt này quan trọng để xác định đúng phạm vi đề tài.",
            "Nhờ giới hạn phạm vi rõ ràng, đề tài có thể tập trung vào trải nghiệm tổ chức trận đấu thay vì cố gắng mô phỏng toàn bộ trò chơi. Đây cũng là cách phù hợp với một đồ án web trong thời gian học phần.",
            "Sự khác biệt này còn giúp người đọc báo cáo dễ đánh giá kết quả hơn. Tiêu chí thành công không phải là tái tạo Genshin Impact, mà là tổ chức được một phiên Ban/Pick có luật, có dữ liệu và có trải nghiệm rõ ràng cho mọi vai trò tham gia.",
        ]),
        (SECTIONS["CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI"][8], [
            "Từ các phân tích ở trên, có thể tổng hợp vấn đề của đề tài thành ba nhóm chính. Nhóm thứ nhất là vấn đề điều phối: trọng tài phải quản lý người chơi, lượt chọn, trạng thái nhân vật và luật giải đấu trong cùng một thời điểm. Khi thao tác bằng công cụ rời rạc, áp lực này dễ tạo sai sót.",
            "Nhóm thứ hai là vấn đề minh bạch. Người chơi cần biết vì sao một thao tác bị từ chối, vì sao đội này bị phạt thời gian nhiều hơn đội kia, và dữ liệu nào được dùng để tính kết quả. Nếu mọi thứ chỉ nằm trong bảng tính của trọng tài, người chơi khó kiểm chứng.",
            "Nhóm thứ ba là vấn đề lưu trữ. Một trận đấu cộng đồng có thể kết thúc trong vài phút, nhưng dữ liệu của nó vẫn có giá trị cho lịch sử, thống kê và giải đấu sau này. Website giúp biến dữ liệu đó thành tài sản có thể tái sử dụng thay vì thông tin tạm thời.",
            "Ba nhóm vấn đề này liên kết trực tiếp với ba nhóm giải pháp của website: điều phối bằng phòng đấu realtime, minh bạch bằng validation và tính Cost tự động, lưu trữ bằng cơ sở dữ liệu có cấu trúc.",
        ]),
        (SECTIONS["CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI"][9], [
            "Phương pháp thực hiện đề tài bắt đầu bằng việc khảo sát quy trình tổ chức Ban/Pick trong cộng đồng, sau đó chuyển quy trình đó thành các vai trò và use case. Đây là bước quan trọng để tránh xây dựng giao diện trước khi hiểu rõ nghiệp vụ.",
            "Sau khi xác định quy trình, đề tài tiến hành thiết kế dữ liệu và phân lớp mã nguồn. Việc thiết kế trước giúp các chức năng như draft, build, result và tournament có cùng nền tảng dữ liệu, không bị phát triển rời rạc.",
            "Cuối cùng, hệ thống được hiện thực và kiểm thử theo luồng người dùng. Cách kiểm thử này phù hợp với website có trạng thái, vì chỉ kiểm tra từng API riêng lẻ chưa đủ để chứng minh toàn bộ trận đấu hoạt động đúng.",
            "Cách tiếp cận trên cũng phù hợp với hai tài liệu mẫu đã tham khảo: báo cáo cần đi từ bối cảnh thực tế, qua phân tích thiết kế, rồi mới đến xây dựng và đánh giá kết quả.",
        ]),
    ]
    for title, paras in add_extras:
        extra_section(doc, title, paras)


def chapter2(doc):
    heading(doc, "CHƯƠNG 2: CÔNG NGHỆ VÀ CÔNG CỤ PHÁT TRIỂN", 1)
    contents = {
        SECTIONS["CHƯƠNG 2: CÔNG NGHỆ VÀ CÔNG CỤ PHÁT TRIỂN"][0]: [
            "Next.js là framework phát triển ứng dụng web dựa trên React, hỗ trợ xây dựng giao diện, định tuyến và API trong cùng một dự án. Với App Router, các trang và route được tổ chức theo thư mục, phù hợp với website có nhiều màn hình như trang chủ, lobby, phòng draft, trang build, kết quả và tournament.",
            "React đóng vai trò xây dựng giao diện theo component. Những phần như bảng draft, thẻ nhân vật, thanh điều hướng, khu vực kết quả và form khai báo build có thể được chia thành các component nhỏ, giúp quá trình phát triển và chỉnh sửa dễ kiểm soát hơn.",
            "Trong dự án này, Next.js không chỉ là công cụ dựng giao diện. Các API route của Next.js còn là nơi tiếp nhận request từ client và chuyển tiếp đến các service xử lý nghiệp vụ. Cách tổ chức full-stack này phù hợp với phạm vi đề tài vì giảm độ phức tạp triển khai nhưng vẫn giữ được cấu trúc rõ ràng.",
            "Một điểm thuận lợi khác của Next.js là cách tổ chức route theo thư mục giúp người đọc mã nguồn dễ hình dung website. Khi nhìn vào các thư mục như room, lobby, tournaments, tools hoặc api, có thể đoán được ngay nhóm chức năng tương ứng. Điều này hỗ trợ cả quá trình phát triển lẫn quá trình trình bày báo cáo.",
            "So với việc tách frontend và backend thành hai dự án riêng, lựa chọn Next.js giúp đồ án gọn hơn. Người thực hiện không phải duy trì hai quy trình build và deploy khác nhau, nhưng vẫn có thể tách nghiệp vụ ra service để tránh việc route handler trở nên quá dài.",
        ],
        SECTIONS["CHƯƠNG 2: CÔNG NGHỆ VÀ CÔNG CỤ PHÁT TRIỂN"][1]: [
            "TypeScript là phần mở rộng của JavaScript, bổ sung hệ thống kiểu dữ liệu tĩnh. Với một website có nhiều kiểu dữ liệu như Room, DraftLog, CharacterBuild, Tournament hoặc Notification, TypeScript giúp giảm lỗi truyền sai cấu trúc giữa các tầng.",
            "Trong luồng Ban/Pick, một số giá trị cần được kiểm soát chặt chẽ, ví dụ đội chỉ có BLUE hoặc RED, hành động chỉ có BAN hoặc PICK. Khi các giá trị này được mô tả bằng kiểu dữ liệu rõ ràng, mã nguồn dễ đọc hơn và việc tái cấu trúc cũng an toàn hơn.",
            "Tuy nhiên, TypeScript không thay thế kiểm thử nghiệp vụ. Nó giúp phát hiện lỗi kiểu dữ liệu, còn các quy tắc như đúng lượt, không chọn trùng hoặc tính handicap vẫn cần được triển khai trong domain policy và kiểm tra bằng dữ liệu cụ thể.",
        ],
        SECTIONS["CHƯƠNG 2: CÔNG NGHỆ VÀ CÔNG CỤ PHÁT TRIỂN"][2]: [
            "Tailwind CSS là framework CSS theo hướng utility-first. Thay vì viết nhiều file CSS riêng cho từng component, người phát triển có thể dùng các lớp tiện ích để điều chỉnh màu sắc, khoảng cách, bố cục và trạng thái hiển thị ngay trong giao diện.",
            "Đối với website Ban/Pick, giao diện cần thể hiện rõ hai đội, trạng thái lượt hiện tại, nhân vật đã bị cấm, nhân vật đã được chọn và kết quả handicap. Vì vậy, Tailwind phù hợp với việc xây dựng nhanh các vùng giao diện có màu sắc tương phản và dễ quan sát.",
            "Giao diện được định hướng theo phong cách eSports: rõ vai trò đội xanh/đội đỏ, ưu tiên thông tin quan trọng và hạn chế làm người dùng mất tập trung trong lúc thao tác. Đây là yêu cầu khác với một website giới thiệu thông thường.",
        ],
        SECTIONS["CHƯƠNG 2: CÔNG NGHỆ VÀ CÔNG CỤ PHÁT TRIỂN"][3]: [
            "Supabase cung cấp các dịch vụ backend như PostgreSQL, xác thực và realtime. Trong đề tài này, realtime là nhu cầu quan trọng vì một lượt Ban/Pick từ người chơi phải được cập nhật cho host, đối thủ và khán giả gần như ngay lập tức.",
            "PostgreSQL được sử dụng làm hệ quản trị cơ sở dữ liệu chính. Các dữ liệu như phòng đấu, lượt draft, build nhân vật, tin nhắn, giải đấu và thông báo cần được lưu bền vững để người dùng có thể tải lại trang hoặc xem lịch sử sau trận.",
            "Supabase giúp rút ngắn thời gian thiết lập hạ tầng, đặc biệt phù hợp với đồ án học phần. Tuy nhiên, khi triển khai lớn hơn, hệ thống vẫn cần theo dõi hiệu năng realtime, số lượng kết nối và chính sách bảo vệ dữ liệu.",
            "Trong bối cảnh phòng Ban/Pick, realtime không chỉ là hiệu ứng tiện lợi. Nó ảnh hưởng trực tiếp đến cảm nhận công bằng của trận đấu. Nếu một bên đã chọn nhân vật nhưng bên còn lại chưa thấy cập nhật, người chơi có thể hiểu sai lượt hoặc tiếp tục trao đổi dựa trên trạng thái cũ.",
            "Chính vì vậy, dữ liệu realtime luôn cần đi kèm dữ liệu bền vững. Supabase giúp giao diện cập nhật nhanh, nhưng trạng thái chính thức vẫn phải được lưu trong PostgreSQL để khi người dùng tải lại trang, mất kết nối hoặc mở thêm thiết bị khác, hệ thống vẫn có thể khôi phục đúng.",
        ],
        SECTIONS["CHƯƠNG 2: CÔNG NGHỆ VÀ CÔNG CỤ PHÁT TRIỂN"][4]: [
            "Prisma ORM là tầng truy cập dữ liệu giúp làm việc với PostgreSQL thông qua các model được định nghĩa trong `schema.prisma`. Thay vì viết SQL thô ở nhiều nơi, ứng dụng có thể truy cập dữ liệu qua Prisma Client với kiểu dữ liệu rõ ràng.",
            "Trong dự án, Prisma mô tả các bảng quan trọng như User, Room, DraftLog, CharacterBuild, LobbyPlayer, ChatMessage, Tournament, TournamentParticipant và TournamentMatch. Những model này phản ánh trực tiếp vòng đời của một trận đấu và giải đấu.",
            "Việc dùng Prisma cũng giúp báo cáo thiết kế cơ sở dữ liệu có căn cứ rõ ràng. Khi mô tả dữ liệu, báo cáo không cần liệt kê toàn bộ field, mà tập trung vào nhóm bảng và quan hệ phục vụ nghiệp vụ.",
            "Một lợi ích thực tế của Prisma là schema đóng vai trò như tài liệu sống cho cơ sở dữ liệu. Khi thêm trường như spectatorDelay, discordWebhookUrl hoặc tournament match, người phát triển có thể nhìn thấy thay đổi ngay trong một nơi tập trung thay vì phải tìm nhiều file migration rời rạc.",
            "Đối với đề tài học phần, điều này giúp giảm lỗi khi giải thích hệ thống. Những gì được trình bày trong báo cáo có thể đối chiếu với schema, tránh tình trạng báo cáo mô tả một bảng hoặc quan hệ không tồn tại trong sản phẩm.",
        ],
        SECTIONS["CHƯƠNG 2: CÔNG NGHỆ VÀ CÔNG CỤ PHÁT TRIỂN"][5]: [
            "Website có sử dụng dữ liệu ngoài để tăng tính tiện lợi cho người dùng. Genshin.dev hỗ trợ dữ liệu nhân vật, còn Enka.network có thể cung cấp thông tin hồ sơ người chơi dựa trên UID nếu người chơi công khai dữ liệu phù hợp.",
            "Việc tích hợp dữ liệu ngoài giúp giảm thao tác nhập tay và làm giao diện sinh động hơn. Ví dụ, khi người chơi nhập UID, hệ thống có thể lấy nickname hoặc thông tin profile để hiển thị trong phòng đấu.",
            "Tuy nhiên, dữ liệu ngoài không phải lúc nào cũng ổn định. Vì vậy, báo cáo xác định rõ rằng việc xác minh build hoàn chỉnh vẫn cần trọng tài hoặc cơ chế bổ sung trong tương lai, thay vì khẳng định hệ thống đã tự động xác minh mọi thông tin trong game.",
        ],
        SECTIONS["CHƯƠNG 2: CÔNG NGHỆ VÀ CÔNG CỤ PHÁT TRIỂN"][6]: [
            "Docker được sử dụng để đóng gói ứng dụng, giúp giảm khác biệt giữa môi trường phát triển và môi trường chạy production. Dockerfile multi-stage cho phép tách bước cài đặt phụ thuộc, build ứng dụng và chạy bản production gọn hơn.",
            "Vercel phù hợp với ứng dụng Next.js vì hỗ trợ quy trình build và deploy thuận tiện. Trong trường hợp cần chạy container riêng, docker-compose cũng đã chuẩn bị môi trường cho ứng dụng kết nối database và dữ liệu cấu hình.",
            "Đối với một website có Supabase, PostgreSQL và nhiều biến môi trường, việc quản lý cấu hình triển khai là phần quan trọng. Các biến như DATABASE_URL, Supabase URL, Supabase key hoặc webhook không nên hard-code trong mã nguồn.",
        ],
        SECTIONS["CHƯƠNG 2: CÔNG NGHỆ VÀ CÔNG CỤ PHÁT TRIỂN"][7]: [
            "Trong quá trình thực hiện, công cụ AI được sử dụng như trợ lý hỗ trợ phân tích mã nguồn, gợi ý cấu trúc báo cáo và rà soát cách diễn đạt. Việc sử dụng AI phù hợp với tinh thần học phần công nghệ mới, nhưng nội dung cuối cùng vẫn cần được kiểm chứng theo sản phẩm thật.",
            "AI không thay thế người thực hiện đề tài. Các quyết định về phạm vi, nghiệp vụ, mức độ hoàn thiện và cách đánh giá sản phẩm vẫn phải dựa trên mã nguồn, tài liệu dự án và kết quả kiểm thử.",
            "Sau khi xác định công nghệ, bước tiếp theo là phân tích yêu cầu và thiết kế hệ thống. Chương 3 trình bày cách các vai trò người dùng, luồng nghiệp vụ và dữ liệu được chuyển thành mô hình thiết kế cụ thể.",
        ],
    }
    for title in SECTIONS["CHƯƠNG 2: CÔNG NGHỆ VÀ CÔNG CỤ PHÁT TRIỂN"][:8]:
        section(doc, title, contents[title])
    table(doc, TABLES[1], ["Công nghệ", "Vai trò", "Lý do sử dụng trong đề tài"], [
        ["Next.js/React", "Trang và API", "Phù hợp website full-stack nhiều màn hình"],
        ["TypeScript", "Kiểm soát kiểu dữ liệu", "Giảm lỗi giữa frontend, service và domain"],
        ["Tailwind CSS", "Giao diện", "Xây dựng nhanh giao diện responsive, rõ trạng thái"],
        ["Supabase/PostgreSQL", "Dữ liệu và realtime", "Lưu trạng thái trận và đồng bộ nhiều người dùng"],
        ["Prisma", "ORM", "Mô tả schema và truy cập dữ liệu có kiểu"],
        ["Docker/Vercel", "Triển khai", "Đưa ứng dụng ra môi trường chạy thực tế"],
    ])
    extra_section(doc, SECTIONS["CHƯƠNG 2: CÔNG NGHỆ VÀ CÔNG CỤ PHÁT TRIỂN"][8], [
        "Mỗi công nghệ được chọn đều gắn với một yêu cầu đã nêu ở Chương 1. Realtime phục vụ việc đồng bộ lượt chọn; PostgreSQL phục vụ lưu lịch sử; TypeScript và domain policy phục vụ tính đúng đắn; Tailwind phục vụ giao diện rõ ràng trong bối cảnh thi đấu.",
        "Cách trình bày này giúp chương công nghệ không trở thành danh sách công cụ rời rạc. Người đọc có thể thấy vì sao công nghệ đó cần thiết đối với sản phẩm, thay vì chỉ biết dự án đã sử dụng những thư viện nào.",
        "Sau khi xác định công nghệ và lý do sử dụng, bước tiếp theo là chuyển các nhu cầu nghiệp vụ thành mô hình thiết kế cụ thể. Đây là nội dung chính của Chương 3, nơi các quy trình, dữ liệu và giao diện được mô tả ở mức chi tiết hơn.",
    ])
    extra_section(doc, SECTIONS["CHƯƠNG 2: CÔNG NGHỆ VÀ CÔNG CỤ PHÁT TRIỂN"][9], [
        "Stack công nghệ của đề tài được lựa chọn theo tiêu chí vừa đủ cho một sản phẩm thật nhưng không vượt quá phạm vi học phần. Next.js giúp gom frontend và backend trong cùng dự án; Supabase giảm thời gian dựng hạ tầng; Prisma giúp dữ liệu rõ ràng; Tailwind hỗ trợ xây dựng giao diện nhanh.",
        "Nếu dùng một backend riêng như Express hoặc NestJS, dự án có thể linh hoạt hơn ở quy mô lớn, nhưng đổi lại sẽ tăng khối lượng thiết lập và triển khai. Với mục tiêu của đề tài, cách tiếp cận full-stack bằng Next.js là hợp lý hơn.",
        "Nếu chỉ dùng localStorage hoặc dữ liệu tĩnh, việc demo có thể nhanh hơn nhưng không phản ánh đúng nhu cầu lưu lịch sử và realtime. Vì vậy, PostgreSQL và Supabase là lựa chọn phù hợp để dữ liệu phòng đấu có tính bền vững.",
        "Điểm cần lưu ý là stack hiện tại vẫn phụ thuộc vào dịch vụ ngoài. Khi triển khai production lâu dài, cần theo dõi giới hạn của Supabase, chi phí realtime, bảo mật key và quy trình backup database.",
    ])
    extra_section(doc, SECTIONS["CHƯƠNG 2: CÔNG NGHỆ VÀ CÔNG CỤ PHÁT TRIỂN"][10], [
        "Bên cạnh framework chính, các công cụ như ESLint, TypeScript compiler, Prisma schema và Docker đóng vai trò hỗ trợ chất lượng. Chúng không trực tiếp tạo ra giao diện mới, nhưng giúp giảm lỗi và làm quá trình phát triển có kỷ luật hơn.",
        "Với một website có nhiều trạng thái như phòng đấu, việc sửa một phần có thể ảnh hưởng sang phần khác. TypeScript và cách tách service giúp phát hiện lỗi sớm hơn, trong khi Docker giúp kiểm tra ứng dụng trong môi trường gần với production.",
        "Trong quá trình viết báo cáo, việc có các file cấu hình, schema và service rõ ràng cũng giúp người thực hiện trình bày chính xác hơn. Báo cáo không cần suy đoán kiến trúc, mà có thể dựa vào chính mã nguồn đã triển khai.",
        "Điều này cho thấy công cụ phát triển không chỉ phục vụ lập trình, mà còn hỗ trợ quá trình tổng kết và đánh giá sản phẩm. Một dự án được tổ chức tốt sẽ dễ báo cáo, dễ bảo trì và dễ mở rộng hơn.",
    ])


def chapter3(doc):
    heading(doc, "CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    section(doc, SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][0], [
        "Quy trình nghiệp vụ của website bắt đầu từ trọng tài hoặc host tạo phòng đấu. Phòng có mã riêng để người chơi tham gia, đồng thời lưu các thiết lập như Cost per point, trạng thái phòng, người chơi hai đội và các ràng buộc liên quan đến giải đấu.",
        "Sau khi vào phòng, hai người chơi giữ slot đội Xanh và đội Đỏ. Khi draft bắt đầu, hệ thống điều khiển lượt Ban/Pick theo template. Mỗi hành động hợp lệ được lưu thành DraftLog, giúp giao diện có thể dựng lại trạng thái bất kỳ lúc nào.",
        "Khi draft hoàn tất, người chơi khai báo build cho các nhân vật đã chọn. Hệ thống dựa trên độ hiếm nhân vật, cung mệnh, vũ khí và Cost Catalog để tính tổng Cost của từng đội. Chênh lệch Cost được quy đổi thành Time Handicap.",
        "Một điểm cần lưu ý là quy trình này không tuyến tính tuyệt đối. Trong quá trình draft, có thể xảy ra tình huống pause khi mạng người chơi gặp sự cố hoặc trọng tài cần can thiệp. Hệ thống vì vậy thiết kế trạng thái phòng có các trường isPaused, pausedAt và pauseReason để hỗ trợ tình huống thực tế.",
        "Sau bước result, dữ liệu trận đấu không kết thúc vai trò. Nó còn được dùng trong lịch sử cá nhân, history theo phòng, hoặc trở thành một game thuộc series BO1/BO3/BO5/BO7 trong khuôn khổ tournament. Đây là cách website kết nối từng trận lẻ thành một hệ thống thi đấu có chiều dọc.",
        "Luồng tổng quát này cho thấy website không chỉ là danh sách nhân vật. Nó là một hệ thống có trạng thái, có quyền thao tác, có dữ liệu lưu trữ và có kết quả phục vụ thi đấu.",
    ])
    figure(doc, FIGURES[0], ["Host/Referee", "↓ tạo phòng và điều phối", "Player Blue / Player Red → Ban/Pick → Build", "Spectator/Caster → Theo dõi realtime", "Admin → Quản lý dữ liệu và cấu hình"])
    table(doc, TABLES[2], ["Tác nhân", "Vai trò", "Nhu cầu chính"], [
        ["Host/Referee", "Điều phối trận", "Tạo phòng, quản lý lượt, xử lý tình huống"],
        ["Player Blue", "Tuyển thủ đội xanh", "Tham gia phòng, Ban/Pick, khai báo build"],
        ["Player Red", "Tuyển thủ đội đỏ", "Tham gia phòng, Ban/Pick, khai báo build"],
        ["Spectator/Caster", "Theo dõi/trình chiếu", "Xem trạng thái realtime và kết quả"],
        ["Admin", "Quản trị", "Quản lý dữ liệu, cấu hình, người dùng"],
    ])
    section(doc, SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][1], [
        "Các tác nhân được xác định dựa trên cách một trận đấu cộng đồng diễn ra. Host hoặc referee là người có quyền cao nhất trong phòng, chịu trách nhiệm tạo phòng, mời người chơi và theo dõi tiến trình.",
        "Player Blue và Player Red là hai tác nhân trực tiếp thực hiện lượt Ban/Pick. Họ chỉ được thao tác khi đến lượt đội mình và chỉ được khai báo build cho các nhân vật thuộc đội mình.",
        "Spectator hoặc caster không trực tiếp thay đổi kết quả, nhưng cần giao diện quan sát rõ ràng. Admin quản lý các phần rộng hơn như cấu hình, dữ liệu, tournament hoặc cost catalog.",
    ])
    table(doc, TABLES[3], ["Nhóm", "Yêu cầu", "Màn hình/chức năng liên quan"], [
        ["Phòng đấu", "Tạo phòng, tham gia đội, xác định vai trò", "Home, Lobby, Room"],
        ["Draft", "Ban/Pick đúng lượt, không chọn trùng", "DraftBoard, Draft API"],
        ["Build", "Khai báo rarity, cung mệnh, vũ khí", "Build page, BuildService"],
        ["Kết quả", "Tính Cost và Time Handicap", "Result page, CostPolicy"],
        ["Giải đấu", "Tạo giải, participant, bracket, match", "Tournament pages"],
        ["Cộng đồng", "Chat, profile, history, notification", "Lobby, Profile, History"],
    ])
    section(doc, SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][2], [
        "Yêu cầu chức năng được gom theo hành trình người dùng thay vì liệt kê theo từng route. Cách trình bày này giúp người đọc thấy chức năng phục vụ bước nào trong quy trình thi đấu.",
        "Nhóm phòng đấu và draft là lõi của hệ thống. Nếu hai nhóm này hoạt động không đúng, các chức năng khác như build, kết quả hoặc tournament đều mất ý nghĩa. Vì vậy, các yêu cầu như đúng lượt, đúng đội và không chọn trùng được ưu tiên kiểm tra ở server.",
        "Các nhóm cộng đồng và giải đấu là phần mở rộng. Chúng làm sản phẩm giống một nền tảng tổ chức sự kiện hơn, nhưng trong báo cáo cần tách rõ phần đã triển khai ở mức nền tảng và phần cần tiếp tục hoàn thiện.",
    ])
    table(doc, TABLES[4], ["Nhóm", "Yêu cầu", "Cách đáp ứng"], [
        ["Realtime", "Cập nhật lượt nhanh", "Supabase Realtime và refresh trạng thái"],
        ["Bảo mật", "Chặn thao tác sai quyền", "Kiểm tra clientId, role, team ở backend"],
        ["Dễ dùng", "Người chơi hiểu lượt hiện tại", "Phân màu đội, trạng thái nhân vật, thông báo lỗi"],
        ["Bảo trì", "Dễ thay đổi luật", "Tách domain policy và application service"],
        ["Triển khai", "Chạy được ngoài máy cá nhân", "Docker, env vars, Supabase/PostgreSQL"],
    ])
    section(doc, SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][3], [
        "Yêu cầu phi chức năng tập trung vào trải nghiệm thi đấu và khả năng vận hành. Một lượt draft cần phản hồi nhanh, nhưng đồng thời vẫn phải được kiểm tra ở backend để tránh dữ liệu sai.",
        "Tính dễ dùng cũng là yêu cầu quan trọng. Giao diện phải cho người chơi biết họ đang ở đội nào, tới lượt ai, nhân vật nào đã bị cấm và nhân vật nào còn khả dụng. Nếu thông tin này không rõ, người chơi sẽ phải hỏi lại trọng tài, làm mất nhịp trận đấu.",
        "Về bảo trì, dự án cần tổ chức mã nguồn để khi luật cộng đồng thay đổi, người phát triển có thể chỉnh ở đúng nơi. Đây là lý do các rule như DraftPolicy, CostPolicy và RoomAccessPolicy được tách khỏi component giao diện.",
    ])
    table(doc, TABLES[5], ["Use Case", "Tác nhân", "Kết quả mong muốn"], [
        ["Tạo phòng", "Host", "Có mã phòng và cấu hình ban đầu"],
        ["Tham gia phòng", "Player", "Giữ đúng slot đội xanh/đỏ"],
        ["Thực hiện Ban/Pick", "Player", "Lượt hợp lệ được lưu và đồng bộ"],
        ["Khai báo build", "Player", "Có dữ liệu tính Cost"],
        ["Xem kết quả", "Tất cả", "Thấy tổng Cost và Handicap"],
        ["Quản lý giải đấu", "Host/Admin", "Có participant, bracket, match"],
    ])
    for title, paras in [
        (SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][4], [
            "Use Case tạo phòng là điểm bắt đầu của phiên đấu. Host nhập cấu hình cần thiết, hệ thống sinh mã phòng và lưu trạng thái ban đầu. Sau đó mã phòng được chia sẻ để người chơi tham gia.",
            "Use Case Ban/Pick là phần quan trọng nhất. Người chơi gửi lựa chọn, backend xác định lượt hiện tại dựa trên log đã có, kiểm tra quyền thao tác và chỉ lưu khi tất cả điều kiện hợp lệ.",
            "Use Case khai báo build và xem kết quả chuyển dữ liệu draft thành thông tin phục vụ thi đấu. Từ danh sách nhân vật đã chọn, mỗi đội nhập build, hệ thống tính Cost và hiển thị Time Handicap.",
        ]),
        (SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][5], [
            "Luật Ban/Pick được thiết kế theo template thay vì viết cứng trong giao diện. Template xác định thứ tự lượt, đội thao tác, hành động là ban hay pick và số lượng nhân vật cần chọn ở mỗi lượt.",
            "Khi nhận request draft, server không tin hoàn toàn vào dữ liệu client gửi lên. DraftService đọc trạng thái phòng, lấy log hiện tại, dùng DraftPolicy để xác định lượt hợp lệ và kiểm tra nhân vật có bị cấm hoặc đã được chọn hay chưa.",
            "Cách xử lý này giúp tăng tính công bằng. Ngay cả khi người dùng cố tình gửi request thủ công, backend vẫn có thể từ chối nếu người đó không đúng đội, phòng đang tạm dừng hoặc hành động không khớp lượt hiện tại.",
            "Bên cạnh template chuẩn, hệ thống cho phép cấu hình draftTemplate riêng cho từng phòng hoặc từng giải đấu. Điều này phù hợp với thực tế cộng đồng, nơi mỗi giải có thể áp dụng số lượt ban khác nhau hoặc bổ sung lượt swap.",
            "Một thiết kế quan trọng nữa là constraint giải đấu. Nếu giải cấm một số nhân vật theo phiên bản hoặc theo vòng đấu, hệ thống có thể đọc constraints của phòng và từ chối những lựa chọn vi phạm. Điều này giúp ban tổ chức kiểm soát luật mà không phải sửa code mỗi khi có thay đổi.",
        ]),
        (SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][6], [
            "Cost được thiết kế để quy đổi độ mạnh đội hình thành một con số có thể so sánh. Nhân vật 5 sao, cung mệnh cao và vũ khí 5 sao làm tổng Cost tăng lên; đội có Cost cao hơn sẽ chịu Time Handicap theo hệ số của phòng.",
            "Hệ thống không quyết định ai thắng trong La Hoàn. Nó chỉ tạo điều kiện thi đấu công bằng hơn bằng cách tính toán phần bù trừ thời gian một cách minh bạch và nhất quán.",
            "Việc có Cost Catalog giúp luật tính điểm có thể thay đổi theo giải hoặc theo meta. Đây là thiết kế mở rộng cần thiết vì luật cộng đồng có thể được điều chỉnh qua từng mùa.",
            "Trong quá trình thiết kế, một câu hỏi đặt ra là có nên hiển thị Cost ngay khi người chơi chọn nhân vật hay đợi đến bước khai báo build mới tính. Phương án hiện tại tách hai bước: pick chỉ ghi nhận lựa chọn, còn cost được tính ở bước build. Điều này phản ánh thực tế là độ mạnh phụ thuộc vào cung mệnh và vũ khí, không chỉ phụ thuộc tên nhân vật.",
            "Cách thiết kế này có ưu điểm là người chơi không bị áp lực ngay khi pick, đồng thời tạo dư địa cho luật giải thay đổi cách tính cost mà không phá vỡ luồng draft chính. Đây là yếu tố quan trọng đối với một sản phẩm muốn phục vụ nhiều giải khác nhau.",
        ]),
    ]:
        section(doc, title, paras)
    figure(doc, FIGURES[1], ["Presentation: src/app, src/components", "Application: services xử lý use case", "Domain: DraftPolicy, CostPolicy, RoomAccessPolicy", "Infrastructure: Prisma repositories, gateways", "Composition: wiring services"])
    section(doc, SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][7], [
        "Kiến trúc dự án được tổ chức theo hướng phân tầng. Tầng Presentation đảm nhận giao diện và API route, tầng Application xử lý các use case, tầng Domain chứa quy tắc nghiệp vụ, còn tầng Infrastructure kết nối database hoặc API ngoài.",
        "Ưu điểm của cách tổ chức này là các luật quan trọng không bị trộn trực tiếp vào component giao diện. Khi cần thay đổi công thức Cost hoặc thứ tự draft, người phát triển có thể tập trung vào domain policy thay vì sửa nhiều màn hình.",
        "Kiến trúc phân tầng cũng giúp báo cáo dễ trình bày hơn. Người đọc có thể hiểu hệ thống theo vai trò của từng lớp thay vì phải đọc toàn bộ cây thư mục mã nguồn.",
        "Composition root, ở vị trí `src/composition/services.ts`, là nơi gắn các adapter cụ thể với service. Nhờ đó, RoomService, DraftService hay TournamentService không phải biết chi tiết Prisma hay API ngoài. Đây là cách áp dụng nguyên lý phụ thuộc ngược ở mức vừa đủ cho phạm vi đồ án.",
        "Cách tổ chức này còn hỗ trợ quá trình kiểm thử. Domain policy có thể được kiểm thử bằng dữ liệu giả, không cần khởi tạo database. Service có thể được kiểm thử với repository giả nếu sau này cần. Đây là nền tảng để mở rộng test tự động trong các giai đoạn tiếp theo.",
    ])
    table(doc, TABLES[6], ["Nhóm dữ liệu", "Bảng tiêu biểu", "Vai trò"], [
        ["Trận đấu", "Room, DraftLog, CharacterBuild", "Lưu phòng, lượt draft và build"],
        ["Giao tiếp", "ChatMessage", "Lưu trao đổi trong phòng"],
        ["Giải đấu", "Tournament, Participant, Match", "Quản lý bracket và kết quả"],
        ["Cộng đồng", "LobbyPlayer, Friendship, Notification", "Online, bạn bè, thông báo"],
        ["Người dùng", "User, UserSettings", "Tài khoản và tùy chọn cá nhân"],
    ])
    figure(doc, FIGURES[2], ["Room 1-n DraftLog", "Room 1-n CharacterBuild", "Room 1-n ChatMessage", "Tournament 1-n Participant", "Tournament 1-n Match", "Lobby/Social liên kết theo UID/clientId"])
    section(doc, SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][8], [
        "Cơ sở dữ liệu được nhìn theo vòng đời trận đấu. Room là trung tâm, lưu trạng thái, người chơi, cấu hình và các thông tin điều phối. DraftLog lưu từng hành động ban hoặc pick để khôi phục lịch sử.",
        "CharacterBuild được tách khỏi DraftLog vì draft chỉ cho biết nhân vật nào được chọn, còn build cho biết nhân vật đó mạnh đến mức nào. Việc tách này làm dữ liệu rõ nghĩa hơn và thuận tiện cho việc tính Cost.",
        "Tournament và các bảng liên quan mở rộng một phòng đơn lẻ thành một phần của giải đấu. Đây là nền tảng để quản lý bracket, participant và kết quả từng match.",
        "Ngoài lõi trận đấu, cơ sở dữ liệu còn có nhóm cộng đồng gồm LobbyPlayer, ChatMessage, Friendship, Notification và ActivityEvent. Những bảng này giúp người chơi không chỉ thi đấu rồi rời đi, mà có thể giữ liên hệ với cộng đồng qua bạn bè, hoạt động và thông báo.",
        "Việc giữ trọn vẹn các nhóm dữ liệu trong cùng một schema giúp dự án dễ mở rộng. Khi muốn thêm chức năng như achievements hoặc missions, người phát triển có thể bổ sung bảng mới hoặc sử dụng lại các quan hệ đã có thay vì xây dựng lại từ đầu.",
        "Tuy nhiên, sự đầy đủ này cũng đặt ra trách nhiệm bảo trì. Mỗi bảng dữ liệu mới đòi hỏi quy trình quản lý, kiểm thử và phòng ngừa rò rỉ thông tin. Đây là điểm cần lưu ý khi phát triển sản phẩm vượt khỏi phạm vi đồ án.",
    ])
    table(doc, TABLES[7], ["Nhóm API", "Chức năng", "Ý nghĩa"], [
        ["Room", "Tạo phòng, join, snapshot", "Quản lý phiên đấu"],
        ["Draft", "Submit, undo lượt", "Bảo vệ luật Ban/Pick"],
        ["Build/Cost", "Lưu build, đọc catalog", "Tính kết quả"],
        ["Tournament", "Tạo giải, bracket, match", "Tổ chức sự kiện"],
        ["Community", "Chat, lobby, notification", "Tương tác người dùng"],
    ])
    section(doc, SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][9], [
        "API được thiết kế theo nhóm nghiệp vụ thay vì trình bày như danh sách endpoint dài. Nhóm Room phục vụ vòng đời phòng đấu, nhóm Draft xử lý lượt cấm/chọn, nhóm Build/Cost xử lý dữ liệu đội hình và nhóm Tournament phục vụ giải đấu.",
        "Cách chia này giúp route handler giữ vai trò tiếp nhận request, còn logic chính được chuyển vào service. Nhờ đó, báo cáo có thể giải thích hệ thống theo use case thay vì theo từng file route.ts.",
        "Mỗi request quan trọng đều cần đi qua kiểm tra phía server. Đây là điểm thiết kế quan trọng vì giao diện chỉ là lớp hiển thị, không thể là nơi duy nhất quyết định tính hợp lệ của thao tác.",
    ])
    figure(doc, FIGURES[3], ["Client chọn nhân vật", "↓", "API Draft nhận request", "↓", "DraftService kiểm tra phòng và quyền", "↓", "DraftPolicy kiểm tra lượt và nhân vật", "↓", "Repository lưu DraftLog", "↓", "Realtime cập nhật giao diện"])
    section(doc, SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][10], [
        "Giao diện được thiết kế theo hành trình người dùng. Trang chủ giúp bắt đầu nhanh; lobby hỗ trợ kết nối người chơi; phòng draft là nơi thao tác chính; trang build thu thập dữ liệu sức mạnh; trang result trình bày Cost và Handicap.",
        "Các màn hình không nên chỉ đẹp mà phải phục vụ tình huống thi đấu. Màu đội, lượt hiện tại, nhân vật đã dùng và thông báo lỗi cần nổi bật hơn các hiệu ứng trang trí.",
        "Từ những thiết kế trên, Chương 4 trình bày cách các ý tưởng được hiện thực thành chức năng cụ thể trong mã nguồn và giao diện website.",
    ])
    chapter3_extras = [
        (SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][11], [
            "Một ưu điểm của thiết kế theo log là trạng thái draft có thể được dựng lại từ dữ liệu đã lưu. Nếu người dùng tải lại trang giữa chừng, hệ thống không cần dựa vào trạng thái tạm trên trình duyệt mà có thể đọc lại Room và DraftLog.",
            "Điều này khác với cách làm thủ công, nơi một lỗi ghi chú có thể làm mất dấu toàn bộ tiến trình. Với dữ liệu có cấu trúc, lịch sử trận đấu có thể được dùng lại cho kết quả, thống kê hoặc giải đấu.",
            "Tính liên tục của dữ liệu cũng giúp việc bàn giao trận đấu giữa các trọng tài hoặc giữa các phiên thi đấu trở nên dễ hơn. Người tiếp nhận chỉ cần mở phòng theo mã, hệ thống tự dựng lại trạng thái mà không phải đối chiếu từng tin nhắn riêng lẻ.",
        ]),
        (SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][12], [
            "Chức năng lõi gồm tạo phòng, join, draft, build và result. Đây là phần cần đúng trước tiên. Các chức năng lobby, tournament, profile, notification hoặc overlay làm sản phẩm phong phú hơn nhưng không được làm mờ mục tiêu chính.",
            "Trong báo cáo, việc tách hai nhóm này giúp người đọc hiểu phạm vi thực hiện. Những gì đã ổn định được trình bày như kết quả, còn những phần cần hoàn thiện thêm được đưa vào hạn chế hoặc hướng phát triển.",
            "Cách phân loại này còn giúp định hướng phát triển sau đề tài. Nếu có thêm thời gian, ưu tiên đầu tiên là làm chắc lõi; còn các chức năng mở rộng sẽ được phát triển dần dần khi nhu cầu cộng đồng rõ ràng hơn.",
        ]),
        (SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][13], [
            "Phân quyền trong website không chỉ là đăng nhập hay chưa đăng nhập. Trong một phòng đấu, cùng một người dùng có thể là host, player của đội xanh, player của đội đỏ hoặc spectator. Mỗi vai trò có quyền khác nhau và ảnh hưởng trực tiếp đến kết quả trận.",
            "Thiết kế sử dụng clientId để nhận diện phiên tham gia phòng. Khi một clientId giữ slot đội xanh, các thao tác Ban/Pick với vai trò BLUE phải khớp clientId đó. Nếu một client khác gửi request thay đội xanh, backend sẽ từ chối.",
            "Host có vai trò điều phối, nhưng không nên thay thế mọi thao tác của player nếu luật giải yêu cầu người chơi tự chọn. Vì vậy, phân quyền cần đủ linh hoạt để hỗ trợ điều phối nhưng vẫn giữ tính công bằng của lượt chọn.",
            "Spectator và caster được phép xem trạng thái nhưng không được thay đổi draft. Điều này giúp website có thể mở cho khán giả hoặc livestream mà không tăng rủi ro làm sai dữ liệu trận đấu.",
        ]),
        (SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][14], [
            "Một phòng đấu không chỉ có trạng thái chờ và kết thúc. Trong quá trình vận hành, phòng có thể ở trạng thái WAITING, DRAFTING, BUILDING, FINISHED hoặc tạm dừng. Mỗi trạng thái quyết định giao diện nào được hiển thị và thao tác nào được chấp nhận.",
            "Trạng thái WAITING phù hợp khi người chơi chưa vào đủ hoặc host chưa bắt đầu. Khi lượt draft đầu tiên được thực hiện, phòng chuyển sang DRAFTING. Sau khi draft hoàn tất, dữ liệu build trở thành trọng tâm và cuối cùng là trang kết quả.",
            "Các trường như lastTurnStartedAt, blueBankTime và redBankTime giúp mô tả thời gian trong phòng. Nếu chỉ lưu danh sách nhân vật, website sẽ thiếu yếu tố thi đấu theo lượt có giới hạn thời gian.",
            "Thiết kế vòng đời phòng giúp giao diện tránh hiển thị sai bước. Ví dụ, khi trận đã FINISHED, người chơi không nên tiếp tục sửa build tùy ý nếu kết quả đã được công bố. Đây là phần cần tiếp tục kiểm soát chặt hơn trong các phiên bản sau.",
        ]),
        (SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][15], [
            "Luật cộng đồng có thể thay đổi theo từng mùa. Có giải muốn nhiều lượt ban hơn, có giải muốn áp dụng fearless draft, có giải giới hạn cung mệnh hoặc cấm một số nhân vật mới ra mắt. Nếu hệ thống viết cứng toàn bộ luật, mỗi giải mới sẽ cần sửa code.",
            "Draft template và constraints được thiết kế để giải quyết vấn đề này. Template mô tả thứ tự lượt, còn constraints mô tả các ràng buộc về nhân vật, vũ khí hoặc build. Nhờ vậy, phần lõi draft có thể dùng chung cho nhiều luật khác nhau.",
            "Fearless draft là ví dụ tiêu biểu. Trong một series nhiều game, nhân vật đã dùng ở game trước có thể không được dùng lại ở game sau. Để hỗ trợ luật này, hệ thống cần liên kết các room trong cùng series và đọc lịch sử lựa chọn trước đó.",
            "Ở phiên bản hiện tại, nền tảng dữ liệu đã chuẩn bị cho các hướng mở rộng này. Tuy nhiên, giao diện cấu hình luật cần được hoàn thiện thêm để ban tổ chức không phải chỉnh JSON hoặc cấu hình kỹ thuật thủ công.",
        ]),
        (SECTIONS["CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"][16], [
            "Khán giả không thao tác trực tiếp nhưng là nhóm người dùng quan trọng nếu trận đấu được livestream. Họ cần thấy diễn biến draft nhanh, rõ và ít bị nhiễu thông tin. Vì vậy, giao diện quan sát cần khác với giao diện thao tác của player.",
            "Spectator delay và overlay là hai ý tưởng phục vụ nhóm người dùng này. Delay giúp tránh lộ chiến thuật nếu giải yêu cầu; overlay giúp trình chiếu trạng thái draft lên stream gọn hơn so với chia sẻ toàn bộ màn hình trình duyệt.",
            "Thiết kế cho khán giả cũng ảnh hưởng đến cách đặt màu và kích thước chữ. Những con số như tổng Cost, thời gian handicap hoặc lượt hiện tại phải đủ nổi bật để người xem hiểu ngay dù không đọc toàn bộ thông tin.",
            "Phần này cho thấy website có thể phát triển theo hướng eSports cộng đồng. Sản phẩm không chỉ phục vụ người bấm nút, mà còn hỗ trợ cách trận đấu được trình bày trước người xem.",
        ]),
    ]
    for title, paras in chapter3_extras:
        extra_section(doc, title, paras)


def chapter4(doc):
    heading(doc, "CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG", 1)
    for title, paras in [
        (SECTIONS["CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG"][0], [
            "Từ thiết kế ở Chương 3, mã nguồn được tổ chức để phản ánh các tầng của hệ thống. Thư mục `src/app` chứa trang và API route, `src/components` chứa giao diện, `src/application/services` chứa use case, `src/domain` chứa luật nghiệp vụ và `src/infrastructure` chứa phần kết nối dữ liệu.",
            "Cách tổ chức này giúp quá trình xây dựng chức năng diễn ra có trật tự. Khi triển khai luồng draft, giao diện chỉ gửi yêu cầu và hiển thị kết quả; DraftService điều phối xử lý; DraftPolicy quyết định tính hợp lệ; repository lưu dữ liệu vào Prisma.",
            "Nhờ đó, báo cáo có thể mô tả việc xây dựng hệ thống theo các mốc chức năng thay vì liệt kê tất cả file. Đây cũng là cách trình bày gần với báo cáo đồ án website hơn.",
        ]),
        (SECTIONS["CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG"][1], [
            "Chức năng tạo phòng là mốc đầu tiên. Khi host tạo phòng, hệ thống sinh mã phòng, lưu cấu hình ban đầu và gắn thông tin host. Người chơi dùng mã phòng để tham gia đội xanh hoặc đội đỏ.",
            "Trong quá trình join phòng, hệ thống kiểm tra slot đội có còn trống không và client hiện tại có đang giữ đội còn lại hay không. Đây là bước quan trọng để tránh một người chiếm cả hai đội hoặc hai người cùng giữ một slot.",
            "Kết quả của mốc này là website có nền tảng phiên đấu riêng biệt. Mỗi phòng có mã, trạng thái, người chơi, cấu hình cost và dữ liệu liên quan để các bước draft, build và result hoạt động tiếp theo.",
            "Trong quá trình hiện thực, phần khó không nằm ở việc tạo một bản ghi phòng, mà nằm ở việc giữ cho vai trò trong phòng nhất quán. Host phải được nhận diện là người điều phối, player phải gắn với đúng đội, còn spectator không được gửi hành động thay người chơi.",
            "Vì vậy, service không chỉ trả về dữ liệu room mà còn trả về session gồm clientId, role và team. Giao diện dựa vào session để hiển thị nút phù hợp, nhưng quyền thật sự vẫn được kiểm tra lại ở các API quan trọng.",
        ]),
        (SECTIONS["CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG"][2], [
            "Draft engine là phần cốt lõi của sản phẩm. Khi người chơi chọn nhân vật, request được gửi đến backend, nơi hệ thống kiểm tra phòng có tồn tại không, người gửi có đúng đội không, phòng có đang tạm dừng không và lượt hiện tại có khớp hành động không.",
            "Nếu thao tác hợp lệ, hệ thống tạo DraftLog và cập nhật thời gian lượt. Nếu thao tác sai, API trả về lỗi để giao diện thông báo cho người dùng. Cách xử lý này giúp giảm phụ thuộc vào trọng tài và giảm nguy cơ sai lượt.",
            "Realtime giúp các client trong cùng phòng nhìn thấy trạng thái mới. Khi một bên ban hoặc pick, bên còn lại và spectator không cần tải lại trang thủ công. Đây là điểm làm website gần với trải nghiệm thi đấu trực tuyến hơn.",
            "Khi xây dựng draft engine, việc xử lý đồng thời cũng cần được quan tâm. Hai client có thể gửi request gần như cùng lúc nếu mạng chậm hoặc người dùng mở nhiều tab. Do đó, thao tác draft được đặt trong transaction và database có ràng buộc unique để giảm khả năng lưu trùng nhân vật.",
            "Bên cạnh đó, thời gian lượt và bank time được cập nhật sau mỗi hành động. Dữ liệu này giúp phòng đấu có cảm giác gần với một phiên thi đấu thật, nơi người chơi không thể kéo dài lượt chọn quá lâu mà không bị ảnh hưởng.",
        ]),
        (SECTIONS["CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG"][3], [
            "Sau khi draft hoàn tất, website chuyển trọng tâm sang khai báo build. Người chơi nhập thông tin nhân vật, cung mệnh và vũ khí cho các nhân vật đã pick. Hệ thống kiểm tra build có thuộc đúng đội và đúng danh sách pick hay không.",
            "Cost được tính dựa trên dữ liệu build và Cost Catalog. Công thức cơ bản giúp người dùng dễ hiểu, còn catalog tạo nền tảng để điều chỉnh luật theo từng giải. Khi hai đội có tổng Cost khác nhau, hệ thống tính Time Handicap bằng chênh lệch nhân với costPerPoint.",
            "Trang kết quả trình bày đội hình, tổng Cost, đội có Cost cao hơn và số giây bị phạt. Đây là đầu ra quan trọng nhất của quy trình vì nó được dùng ngay trong trận La Hoàn thực tế.",
            "Phần build cũng là nơi cần cân bằng giữa tự động và thủ công. Dữ liệu Enka có thể hỗ trợ lấy thông tin profile, nhưng không phải lúc nào build trong showcase cũng là build người chơi dùng khi thi đấu. Vì vậy, hệ thống vẫn cho phép khai báo và điều chỉnh dữ liệu build theo phiên trận.",
            "Khi kết quả được tính, người dùng không chỉ thấy con số cuối cùng. Báo cáo nên nhấn mạnh rằng từng build có thể được xem lại, tổng cost của đội được cộng từ nhiều nhân vật, và handicap là kết quả của một công thức minh bạch. Đây là điểm giúp giảm tranh luận trong quá trình tổ chức.",
        ]),
        (SECTIONS["CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG"][4], [
            "Ngoài draft và result, website còn có các chức năng giúp trận đấu vận hành mượt hơn. Timer và bank time tạo áp lực thời gian; pause/undo hỗ trợ xử lý sự cố; chat giúp host và player trao đổi ngay trong phòng.",
            "Những chức năng này không phải phần trang trí. Trong một trận đấu thật, người chơi có thể vào nhầm đội, mạng bị chậm hoặc chọn nhầm nhân vật. Website cần có cơ chế xử lý vừa đủ để trọng tài không phải hủy toàn bộ phòng.",
            "Ở phiên bản hiện tại, các chức năng này tạo nền tảng cho vận hành trận đấu quy mô nhỏ đến vừa. Nếu dùng cho giải lớn, cần bổ sung thêm audit log và quy trình xác nhận thao tác nhạy cảm.",
        ]),
        (SECTIONS["CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG"][5], [
            "Website đã có các phần mở rộng như lobby, profile, history, tournament và overlay. Lobby giúp người chơi nhìn thấy trạng thái online và lời mời; history giúp xem lại trận; tournament giúp liên kết nhiều trận thành một giải đấu có bracket.",
            "Cần phân biệt rõ phần đã triển khai và phần định hướng. Các bảng dữ liệu và giao diện cơ bản đã tạo nền tảng cho tournament, nhưng các tính năng nâng cao như seeding phức tạp, thống kê meta hoặc vận hành giải lớn vẫn cần phát triển thêm.",
            "Việc có nền tảng mở rộng cho thấy sản phẩm không chỉ là bản demo một màn hình. Tuy nhiên, báo cáo vẫn đặt chức năng lõi Ban/Pick, build và result làm trọng tâm đánh giá.",
            "Lobby và profile giúp sản phẩm có cảm giác cộng đồng hơn. Người chơi không chỉ vào một phòng duy nhất rồi rời đi, mà có thể có trạng thái online, avatar, nickname, bạn bè và thông báo. Những yếu tố này không bắt buộc cho một MVP, nhưng có ý nghĩa nếu website được dùng thường xuyên.",
            "Tournament là phần mở rộng có giá trị nhất về mặt nghiệp vụ. Khi một trận đấu đã có roomCode và result, việc liên kết nó vào match giúp ban tổ chức quản lý bracket, round và winner dễ hơn. Đây là bước chuyển từ công cụ tổ chức một trận sang nền tảng tổ chức giải.",
        ]),
        (SECTIONS["CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG"][6], [
            "Cơ sở dữ liệu được triển khai trên PostgreSQL thông qua Supabase. Prisma schema là nơi mô tả các bảng và quan hệ, sau đó ứng dụng sử dụng Prisma Client để đọc ghi dữ liệu.",
            "Ứng dụng có thể build bằng Next.js và đóng gói bằng Docker. Dockerfile multi-stage giúp tách giai đoạn build và giai đoạn chạy, còn docker-compose hỗ trợ cấu hình môi trường khi chạy container.",
            "Khi triển khai, các biến môi trường như DATABASE_URL, DIRECT_URL, Supabase URL, Supabase key và webhook cần được cấu hình ngoài mã nguồn. Đây là yêu cầu cơ bản để bảo vệ thông tin và giúp ứng dụng chạy ổn định ở môi trường khác nhau.",
        ]),
    ]:
        section(doc, title, paras)
    table(doc, TABLES[8], ["Tầng", "Thư mục/file tiêu biểu", "Vai trò"], [
        ["Presentation", "src/app, src/components", "Trang, API route, component giao diện"],
        ["Application", "src/application/services", "Điều phối use case"],
        ["Domain", "src/domain", "Luật draft, cost, phân quyền"],
        ["Infrastructure", "src/infrastructure", "Repository, gateway, Prisma"],
        ["Composition", "src/composition/services.ts", "Kết nối service và adapter"],
    ])
    section(doc, SECTIONS["CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG"][7], [
        "Kiểm thử được thực hiện theo các luồng quan trọng của sản phẩm. Trọng tâm là chứng minh hệ thống có thể chạy từ tạo phòng đến kết quả, đồng thời chặn được các thao tác sai quyền hoặc sai lượt.",
        "Bảng kiểm thử dưới đây trình bày dữ liệu thử, kết quả mong đợi và kết quả thực tế ở mức chức năng. Đây không thay thế kiểm thử tự động đầy đủ, nhưng giúp chứng minh các luồng chính đã được xem xét khi hoàn thiện đề tài.",
    ])
    table(doc, TABLES[9], ["Mã", "Kịch bản", "Dữ liệu thử", "Mong đợi", "Thực tế", "Đánh giá"], [
        ["TC01", "Tạo phòng", "Host + costPerPoint=10", "Có mã phòng", "Phòng được tạo, host có session", "Đạt"],
        ["TC02", "Join đội xanh", "Tên hợp lệ + clientId", "Giữ slot Blue", "blueClientId được lưu", "Đạt"],
        ["TC03", "Join trùng slot", "Client khác vào Blue", "Bị từ chối", "API trả 409", "Đạt"],
        ["TC04", "Pick sai lượt", "Red gửi khi lượt Blue", "Bị từ chối", "API trả lỗi validation", "Đạt"],
        ["TC05", "Pick trùng", "Chọn nhân vật đã dùng", "Không lưu", "Ràng buộc/validation chặn", "Đạt"],
        ["TC06", "Tính Handicap", "Blue 10 cost, Red 8 cost", "20 giây nếu costPerPoint=10", "Kết quả hiển thị đúng", "Đạt"],
        ["TC07", "Reload phòng", "Tải lại giữa draft", "Khôi phục trạng thái", "Đọc lại Room/DraftLog", "Đạt"],
    ])
    section(doc, SECTIONS["CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG"][8], [
        "Kết quả giao diện được trình bày theo hành trình sử dụng sản phẩm. Các hình dưới đây là vị trí minh họa cần được thay bằng ảnh chụp thật trước khi nộp chính thức nếu người dùng có môi trường chạy website đầy đủ.",
        "Mỗi giao diện được chọn vì thể hiện một phần quan trọng của sản phẩm: bắt đầu từ trang chủ, chuyển sang phòng draft, khai báo build, xem result và mở rộng sang tournament. Cách chọn hình này giúp báo cáo tập trung vào sản phẩm thay vì chèn quá nhiều ảnh rời rạc.",
    ])
    for caption, lines in [
        (FIGURES[4], ["Trang chủ giới thiệu mục đích website", "Nút tạo phòng/tham gia", "Điều hướng đến lobby, tournament, công cụ"]),
        (FIGURES[5], ["Hai khu vực đội Xanh/Đỏ", "Lượt hiện tại và timer", "Danh sách nhân vật với trạng thái ban/pick"]),
        (FIGURES[6], ["Form nhập rarity, cung mệnh, vũ khí", "Tính Cost theo từng nhân vật", "Dữ liệu chuẩn bị cho trang kết quả"]),
        (FIGURES[7], ["Tổng Cost hai đội", "Chênh lệch và Time Handicap", "Danh sách đội hình đã chọn"]),
        (FIGURES[8], ["Danh sách giải đấu", "Participant và bracket", "Liên kết match với phòng đấu"]),
    ]:
        para(doc, f"Trước {caption.lower()}, báo cáo cần làm rõ màn hình này nằm ở bước nào trong luồng sử dụng và người dùng nhận được thông tin gì từ màn hình đó.")
        figure(doc, caption, lines)
        para(doc, "Sau hình, có thể thấy giao diện không chỉ đóng vai trò minh họa mà còn là minh chứng cho việc các yêu cầu đã phân tích ở Chương 3 được hiện thực thành sản phẩm có thể sử dụng.")
    chapter4_extras = [
        (SECTIONS["CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG"][9], [
            "Trong quá trình xây dựng, phần cần ưu tiên nhất là luồng tạo phòng, draft, build và result. Khi luồng này chạy thông suốt, website đã giải quyết được bài toán chính của đề tài. Các chức năng mở rộng được phát triển xoay quanh luồng này, không thay thế trọng tâm của sản phẩm.",
            "Cách triển khai theo service và domain policy giúp quá trình sửa lỗi dễ hơn. Ví dụ, nếu kết quả handicap chưa đúng, người phát triển có thể tập trung vào phần cost thay vì dò toàn bộ giao diện.",
            "Việc đánh giá hiện thực theo từng mốc còn giúp xác định đâu là phần đã ổn định, đâu là phần đang ở dạng nền tảng. Đây là cơ sở để Chương 5 đối chiếu mục tiêu ban đầu với kết quả thực tế thay vì khẳng định toàn bộ chức năng đã hoàn thiện.",
        ]),
        (SECTIONS["CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG"][10], [
            "Với quy mô cộng đồng nhỏ hoặc buổi thi đấu thử nghiệm, hệ thống đã có đủ các thành phần cần thiết: phòng riêng, người chơi hai đội, trạng thái realtime, kết quả và lịch sử. Đây là mức phù hợp với phạm vi học phần.",
            "Nếu đưa vào vận hành giải lớn, hệ thống cần thêm logging, monitoring, kiểm thử tải và quy trình xử lý tranh chấp rõ ràng hơn. Những nội dung này được đặt ở Chương 5 như hạn chế và hướng phát triển, tránh trình bày nhầm thành chức năng đã hoàn thiện.",
            "Đánh giá vận hành thực tế cũng cho thấy giá trị của việc phân tách rõ giữa phần đã làm và phần định hướng. Người dùng cuối, đặc biệt là ban tổ chức giải đấu, cần biết chính xác họ đang sử dụng phiên bản nào và phần nào sẽ được nâng cấp trong tương lai.",
        ]),
        (SECTIONS["CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG"][11], [
            "Draft Board là màn hình cần được đầu tư nhiều nhất vì đây là nơi người chơi tương tác liên tục. Giao diện phải vừa hiển thị danh sách nhân vật, vừa hiển thị lượt hiện tại, vừa cho thấy lịch sử ban/pick của hai đội.",
            "Khi xây dựng Draft Board, trạng thái của nhân vật được chia thành nhiều nhóm: còn khả dụng, đã bị ban, đã được pick hoặc bị cấm bởi luật giải đấu. Cách phân trạng thái này giúp người chơi không phải tự đối chiếu danh sách bằng mắt.",
            "Các nút thao tác chỉ nên nổi bật khi người dùng có quyền hành động. Nếu spectator hoặc đội chưa đến lượt nhìn thấy nút tương tự player đang đến lượt, giao diện sẽ gây hiểu nhầm. Vì vậy, UI cần phối hợp với session và trạng thái lượt hiện tại.",
            "Draft Board cũng là nơi kiểm tra tốt nhất giữa frontend và backend. Nếu giao diện hiển thị đúng nhưng backend từ chối sai, người dùng sẽ bối rối; nếu backend đúng nhưng giao diện chậm cập nhật, trận đấu cũng bị gián đoạn. Hai phía phải hoạt động đồng bộ.",
        ]),
        (SECTIONS["CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG"][12], [
            "Trang kết quả là bước tổng hợp cuối cùng của luồng thi đấu. Người dùng không còn quan tâm đến thao tác từng lượt, mà cần nhìn thấy đội hình hai bên, tổng Cost và đội chịu Time Handicap. Vì vậy, trang này cần trình bày ít nhưng rõ.",
            "Khi xây dựng trang kết quả, dữ liệu được lấy từ build đã lưu và Cost Catalog hiện tại. Hệ thống tính cost từng nhân vật, sau đó cộng thành tổng đội. Sự tách bạch giữa cost từng nhân vật và tổng đội giúp trọng tài dễ giải thích nếu có người chơi thắc mắc.",
            "Trang kết quả cũng cần hỗ trợ chia sẻ hoặc trình chiếu. Nếu dùng cho livestream, giao diện nên tránh quá nhiều chi tiết nhỏ và ưu tiên con số chính. Đây là lý do các phần tổng cost và handicap nên có vị trí nổi bật.",
            "Kết quả không chỉ là điểm cuối của một trận. Nó còn là dữ liệu đầu vào cho history, tournament match hoặc thống kê sau giải. Do đó, việc lưu và hiển thị kết quả phải nhất quán với dữ liệu trong database.",
        ]),
        (SECTIONS["CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG"][13], [
            "Tournament được xây dựng như phần mở rộng từ một phòng đấu riêng lẻ. Khi số trận nhiều hơn một, ban tổ chức cần quản lý người tham gia, vòng đấu, match, winner và roomCode liên kết với từng trận.",
            "Dữ liệu TournamentParticipant giúp lưu thông tin người chơi hoặc đội tham gia giải. TournamentMatch giúp mô tả cặp đấu ở từng round và trạng thái winner. Khi match được liên kết với roomCode, trận draft cụ thể trở thành một phần của bracket.",
            "Trong phiên bản hiện tại, tournament tạo nền tảng cho single elimination và các format cơ bản. Những phần nâng cao như double elimination, round robin hoàn chỉnh hoặc seeding phức tạp cần tiếp tục phát triển thêm.",
            "Việc đưa tournament vào báo cáo giúp thể hiện hướng phát triển của sản phẩm, nhưng cần trình bày trung thực. Đây là nền tảng mở rộng, không nên mô tả như một hệ thống vận hành giải chuyên nghiệp hoàn chỉnh nếu chưa kiểm thử với nhiều đội thật.",
        ]),
        (SECTIONS["CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG"][14], [
            "Khi chuyển từ môi trường phát triển sang môi trường triển khai, vấn đề thường gặp không phải chỉ là code có chạy hay không, mà là cấu hình môi trường có đúng hay không. Database URL, Supabase key và đường dẫn dữ liệu cost catalog đều ảnh hưởng trực tiếp đến ứng dụng.",
            "Docker giúp giảm sai khác giữa máy phát triển và máy chạy production. Khi image được build thành công, ứng dụng có môi trường phụ thuộc rõ ràng hơn so với việc cài thủ công từng gói trên server.",
            "Docker compose trong dự án cho phép cấu hình port, environment và volume dữ liệu. Điều này thuận tiện khi cần chạy ứng dụng theo dạng container, đặc biệt trong trường hợp muốn triển khai ở VPS hoặc môi trường không phải Vercel.",
            "Dù vậy, triển khai không kết thúc ở bước chạy được container. Sản phẩm cần được kiểm tra lại trên môi trường thật: tạo phòng, join, realtime, đọc/ghi database và truy cập dữ liệu ngoài. Đây là bước quan trọng trước khi dùng cho một buổi thi đấu thật.",
        ]),
    ]
    for title, paras in chapter4_extras:
        extra_section(doc, title, paras)


def chapter5(doc):
    heading(doc, "CHƯƠNG 5: TỔNG KẾT, ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN", 1)
    section(doc, SECTIONS["CHƯƠNG 5: TỔNG KẾT, ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN"][0], [
        "Sau quá trình thực hiện, đề tài đã xây dựng được một website hỗ trợ Ban/Pick Genshin Impact theo luồng end-to-end. Người dùng có thể tạo phòng, tham gia đội, thực hiện Ban/Pick, khai báo build, tính Cost, xem Time Handicap và lưu lại dữ liệu trận đấu.",
        "Kết quả đạt được không chỉ nằm ở số lượng màn hình, mà ở việc các màn hình được nối thành một quy trình có ý nghĩa. Dữ liệu từ phòng đấu được dùng cho draft, dữ liệu draft được dùng cho build, dữ liệu build được dùng cho result, và kết quả có thể liên kết với history hoặc tournament.",
        "Về mặt kỹ thuật, dự án áp dụng được stack web hiện đại và kiến trúc phân tầng. Đây là nền tảng để tiếp tục phát triển sản phẩm sau phạm vi học phần.",
    ])
    table(doc, TABLES[10], ["Mục tiêu ban đầu", "Kết quả hiện tại", "Mức độ"], [
        ["Chuẩn hóa quy trình Ban/Pick", "Có phòng, slot đội, lượt draft và validation server", "Đạt"],
        ["Tự động tính Cost/Handicap", "Có build, Cost Catalog, tính tổng và chênh lệch", "Đạt"],
        ["Giao diện theo bối cảnh thi đấu", "Có draft board, result, lobby/tournament cơ bản", "Đạt một phần"],
        ["Lưu lịch sử và mở rộng giải đấu", "Có dữ liệu Room, DraftLog, Tournament, Match", "Có nền tảng"],
        ["Triển khai thực tế", "Có Docker, cấu hình Supabase/PostgreSQL", "Đạt mức đồ án"],
    ])
    for title, paras in [
        (SECTIONS["CHƯƠNG 5: TỔNG KẾT, ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN"][1], [
            "Khi đối chiếu với mục tiêu ban đầu, phần cốt lõi của đề tài đã được đáp ứng. Luồng Ban/Pick không còn phụ thuộc hoàn toàn vào trọng tài ghi tay; thao tác sai lượt hoặc sai quyền được kiểm tra phía server; Cost và Handicap được tính tự động từ dữ liệu build.",
            "Một số mục tiêu mở rộng như tournament, lobby, profile và overlay đã có nền tảng trong mã nguồn và giao diện, nhưng vẫn cần tiếp tục hoàn thiện trước khi xem là sản phẩm cộng đồng quy mô lớn. Việc đánh giá theo mức độ giúp báo cáo trung thực hơn, không khẳng định quá mức những phần mới ở giai đoạn nền tảng.",
        ]),
        (SECTIONS["CHƯƠNG 5: TỔNG KẾT, ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN"][2], [
            "Ưu điểm đầu tiên của hệ thống là quy trình thi đấu được chuẩn hóa. Người chơi biết khi nào đến lượt, trọng tài có dữ liệu để đối chiếu, còn khán giả có thể theo dõi trạng thái rõ ràng hơn so với cách làm bằng tin nhắn rời rạc.",
            "Ưu điểm thứ hai là logic quan trọng được kiểm tra ở backend. Điều này giúp hạn chế việc người dùng thao tác sai hoặc cố tình gửi request không hợp lệ. Đây là điểm khác biệt giữa một giao diện chọn nhân vật đơn giản và một công cụ có thể dùng trong thi đấu.",
            "Ưu điểm thứ ba là khả năng mở rộng. Kiến trúc phân tầng, draft template, cost catalog và nhóm dữ liệu tournament tạo điều kiện để sản phẩm phát triển tiếp theo nhu cầu cộng đồng.",
        ]),
        (SECTIONS["CHƯƠNG 5: TỔNG KẾT, ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN"][3], [
            "Bên cạnh kết quả đạt được, hệ thống vẫn còn một số hạn chế. Trước hết, dữ liệu build vẫn phụ thuộc một phần vào khai báo của người chơi hoặc kiểm tra của trọng tài. Việc lấy profile qua UID chỉ hỗ trợ hiển thị và tham khảo, chưa thể thay thế hoàn toàn quy trình xác minh trong game.",
            "Thứ hai, website chưa tự động lấy kết quả clear time từ Genshin Impact. Hệ thống tính Handicap và hỗ trợ tổ chức trận đấu, nhưng kết quả thi đấu thực tế vẫn cần được người chơi hoặc trọng tài nhập/xác nhận bên ngoài.",
            "Thứ ba, realtime và triển khai production cần kiểm thử thêm với nhiều phòng đồng thời. Trong phạm vi đồ án, hệ thống phù hợp với môi trường thử nghiệm và cộng đồng nhỏ; để vận hành lớn cần giám sát, rate limiting, backup và logging chi tiết hơn.",
            "Cuối cùng, các hình trong báo cáo nên được thay bằng ảnh chụp thật từ website trước khi nộp chính thức. Điều này giúp minh chứng sản phẩm thuyết phục hơn so với placeholder.",
        ]),
        (SECTIONS["CHƯƠNG 5: TỔNG KẾT, ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN"][4], [
            "Hướng phát triển đầu tiên là hoàn thiện giao diện quản trị luật giải đấu. Ban tổ chức nên có khả năng cấu hình draft template, constraints, Cost Catalog và series format trực tiếp trên giao diện thay vì phụ thuộc vào chỉnh sửa kỹ thuật.",
            "Hướng thứ hai là tăng khả năng xác minh build. Hệ thống có thể khai thác thêm dữ liệu từ Enka hoặc quy trình upload minh chứng để trọng tài kiểm tra nhanh hơn, dù vẫn cần lưu ý giới hạn từ API ngoài và quyền riêng tư của người chơi.",
            "Hướng thứ ba là phát triển thống kê meta và broadcast overlay. Khi dữ liệu đủ lớn, website có thể thống kê nhân vật bị ban nhiều nhất, đội hình phổ biến hoặc kết quả theo tournament. Overlay giúp sản phẩm phù hợp hơn với livestream.",
            "Về kỹ thuật, cần bổ sung test tự động, audit log, monitoring, backup database và tối ưu mobile UX. Đây là các bước quan trọng nếu sản phẩm được dùng thường xuyên trong cộng đồng.",
        ]),
        (SECTIONS["CHƯƠNG 5: TỔNG KẾT, ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN"][5], [
            "Bài học quan trọng nhất trong quá trình thực hiện là phải bắt đầu từ nghiệp vụ trước khi chọn công nghệ. Khi hiểu rõ trọng tài, người chơi và khán giả cần gì, các quyết định về realtime, dữ liệu, phân quyền và giao diện trở nên có lý do hơn.",
            "Bài học thứ hai là không nên đặt toàn bộ logic vào giao diện. Một website thi đấu cần backend kiểm tra các thao tác quan trọng, vì kết quả trận đấu phải minh bạch và có thể đối chiếu.",
            "Tổng kết lại, đề tài đã xây dựng được nền tảng website Ban/Pick Genshin Impact có tính thực tiễn, có quy trình rõ ràng và có khả năng phát triển tiếp. Dù vẫn còn hạn chế, sản phẩm đã thể hiện được cách ứng dụng công nghệ web hiện đại vào một nhu cầu cụ thể của cộng đồng game.",
        ]),
    ]:
        section(doc, title, paras)
    extra_section(doc, SECTIONS["CHƯƠNG 5: TỔNG KẾT, ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN"][6], [
        "Giá trị của đề tài không nằm ở việc mô tả thật nhiều công nghệ, mà ở việc các công nghệ đó cùng phục vụ một quy trình sử dụng cụ thể. Người dùng cuối không nhìn thấy DraftPolicy hay Prisma Repository, nhưng họ cảm nhận được việc lượt chọn đúng, kết quả rõ và thao tác ít nhầm hơn.",
        "Vì vậy, nếu tiếp tục phát triển, sản phẩm nên giữ trọng tâm là hỗ trợ tổ chức trận đấu công bằng. Các tính năng mới chỉ nên được thêm khi chúng làm quy trình thi đấu rõ ràng hơn, minh bạch hơn hoặc dễ vận hành hơn.",
        "Khi nhìn lại quá trình thực hiện, có thể thấy đề tài này phù hợp với học phần công nghệ mới không chỉ vì áp dụng nhiều công cụ hiện đại, mà còn vì nó kết nối được công nghệ với một nhu cầu thực tế của cộng đồng người chơi Genshin Impact.",
    ])
    extra_section(doc, SECTIONS["CHƯƠNG 5: TỔNG KẾT, ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN"][7], [
        "Một báo cáo tổng kết website cần chứng minh được sản phẩm giải quyết vấn đề nào, đã được thiết kế ra sao và kết quả có phù hợp mục tiêu ban đầu hay không. Vì vậy, bản báo cáo này không tập trung liệt kê toàn bộ route hoặc file mã nguồn, mà tập trung vào mạch sử dụng của website.",
        "Cách trình bày theo luồng end-to-end giúp người đọc dễ theo dõi hơn: từ tạo phòng đến draft, từ draft đến build, từ build đến result và từ result đến history/tournament. Đây là mạch gần với trải nghiệm người dùng thật.",
        "Các phần kỹ thuật vẫn được giữ nhưng đặt đúng vai trò. Next.js, Supabase, Prisma hay Docker được giải thích vì chúng hỗ trợ yêu cầu nào, không phải vì báo cáo cần liệt kê nhiều công nghệ.",
        "Nhờ đó, báo cáo vừa có tính kỹ thuật vừa giữ được tinh thần tổng kết sản phẩm. Đây là hướng phù hợp hơn với hai tài liệu mẫu đã đọc so với việc tạo nhiều bảng hoặc phụ lục kỹ thuật dài.",
    ])
    extra_section(doc, SECTIONS["CHƯƠNG 5: TỔNG KẾT, ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN"][8], [
        "Sau khi hoàn thiện bản đồ án, bước tiếp theo nên là tổ chức một buổi dùng thử với vài người chơi thật. Buổi dùng thử cần mô phỏng đầy đủ: host tạo phòng, hai người chơi vào đội, thực hiện draft, khai báo build và xem kết quả.",
        "Trong quá trình dùng thử, cần ghi nhận những điểm người chơi phải hỏi lại, bấm nhầm hoặc chờ quá lâu. Những phản hồi này có giá trị hơn cảm nhận chủ quan của người phát triển, vì chúng phản ánh trải nghiệm thật.",
        "Nếu người chơi mới hiểu được lượt hiện tại mà không cần hướng dẫn nhiều, giao diện draft có thể xem là đạt mục tiêu. Nếu trọng tài giải thích được cost và handicap ngay trên trang result, phần kết quả có thể xem là minh bạch.",
        "Những dữ liệu phản hồi này nên được dùng cho vòng cải tiến tiếp theo. Một sản phẩm cộng đồng tốt không chỉ được xây dựng từ suy nghĩ của người phát triển, mà cần được điều chỉnh theo cách cộng đồng thật sự sử dụng.",
    ])
    extra_section(doc, SECTIONS["CHƯƠNG 5: TỔNG KẾT, ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN"][9], [
        "Với quy mô hiện tại, website phù hợp để áp dụng cho các trận giao hữu, buổi test luật hoặc giải đấu nhỏ trong cộng đồng. Các chức năng lõi đã đủ để giảm lỗi thủ công và trình bày kết quả rõ hơn.",
        "Nếu cộng đồng muốn dùng thường xuyên, cần chuẩn bị thêm hướng dẫn sử dụng cho host và player. Hướng dẫn này nên ngắn, tập trung vào các bước chính, tránh biến công cụ thành thứ khó tiếp cận với người không rành kỹ thuật.",
        "Ban tổ chức cũng cần thống nhất quy trình xác minh build và kết quả trong game. Website hỗ trợ tính toán và lưu dữ liệu, nhưng vẫn cần quy tắc vận hành bên ngoài để xử lý tranh chấp hoặc trường hợp dữ liệu người chơi không rõ ràng.",
        "Nhìn chung, đề tài có khả năng áp dụng thực tế nếu được hoàn thiện thêm ảnh giao diện, kiểm thử người dùng, cấu hình triển khai ổn định và hướng dẫn vận hành rõ ràng.",
    ])



def references(doc):
    heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "Next.js Documentation", "React Documentation", "TypeScript Documentation", "Tailwind CSS Documentation",
        "Supabase Documentation", "Prisma Documentation", "Docker Documentation", "Vercel Documentation",
        "Genshin.dev API", "Enka.network", "Genshin Impact / HoYoverse public resources",
    ]
    for i, ref in enumerate(refs, 1):
        para(doc, f"[{i}] {ref}.", first=False)


def main():
    doc = Document()
    setup(doc)
    cover(doc)
    front(doc)
    chapter1(doc)
    doc.add_page_break()
    chapter2(doc)
    doc.add_page_break()
    chapter3(doc)
    doc.add_page_break()
    chapter4(doc)
    doc.add_page_break()
    chapter5(doc)
    references(doc)
    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    main()
