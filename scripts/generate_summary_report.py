from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "BaoCao_TomTat_Website_BanPick_Genshin_TheoMau.docx"


def set_font(run, size=13, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def setup(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    for style_name, size in [("Normal", 13), ("Heading 1", 14), ("Heading 2", 13)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        if style_name.startswith("Heading"):
            style.font.bold = True
    doc.styles["Normal"].paragraph_format.line_spacing = 1.5
    doc.styles["Normal"].paragraph_format.space_after = Pt(4)


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, bold=False, italic=False, size=13):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    if first_line and text:
        p.paragraph_format.first_line_indent = Cm(1)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("- " + text)
    set_font(r, 13)
    return p


def heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    set_font(r, size=13, bold=True)
    return p


def set_cell(cell, text, bold=False):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    set_font(r, size=10.3, bold=bold)


def borders(table):
    tbl_pr = table._tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:color"), "444444")
        tbl_borders.append(element)
    tbl_pr.append(tbl_borders)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(t)
    for idx, header in enumerate(headers):
        set_cell(t.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            set_cell(cells[idx], value)
    return t


def cover(doc):
    for line in ["TRƯỜNG ĐẠI HỌC ĐÀ LẠT", "KHOA CÔNG NGHỆ THÔNG TIN"]:
        paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, bold=True, size=14)
    paragraph(doc, "\n\nBÁO CÁO TÓM TẮT", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, bold=True, size=18)
    paragraph(doc, "HỌC PHẦN: CÁC CÔNG NGHỆ MỚI", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, bold=True, size=14)
    paragraph(doc, "TRONG PHÁT TRIỂN PHẦN MỀM", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, bold=True, size=14)
    paragraph(doc, "ĐỀ TÀI: XÂY DỰNG WEBSITE HỖ TRỢ CẤM/CHỌN NHÂN VẬT", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, bold=True, size=14)
    paragraph(doc, "TRONG GENSHIN IMPACT", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, bold=True, size=14)
    paragraph(doc, "\n\nGiảng viên hướng dẫn:  KS. Nguyễn Trọng Hiếu", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    paragraph(doc, "Sinh viên thực hiện:  2212377 - Trần Ngọc Hưng", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    paragraph(doc, "Lớp: [Cần bổ sung]", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    paragraph(doc, "\n\n\n\nLâm Đồng, tháng 05 năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    doc.add_page_break()


def teacher_comment(doc):
    paragraph(doc, "NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, bold=True, size=14)
    for _ in range(24):
        paragraph(doc, "." * 106, first_line=False)
    doc.add_page_break()


def content(doc):
    add_page_number(doc.sections[0])

    heading(doc, "1. Mục tiêu chính của đề tài")
    paragraph(doc, "Đề tài xây dựng website hỗ trợ cấm/chọn nhân vật trong Genshin Impact được thực hiện nhằm giải quyết nhu cầu tổ chức các trận đấu cộng đồng có luật cấm/chọn rõ ràng, có trọng tài điều phối, có dữ liệu lưu vết và có kết quả tính Cost, Time Handicap minh bạch. Trong thực tế, các trận đấu cộng đồng thường được tổ chức bằng Discord, bảng tính hoặc ghi chú thủ công, dễ phát sinh sai lượt, sai nhân vật hoặc tranh luận về kết quả bù trừ thời gian.")
    paragraph(doc, "Đề tài hướng tới ba mục tiêu cốt lõi:")
    bullet(doc, "Về nghiệp vụ và sản phẩm: xây dựng một luồng sử dụng hoàn chỉnh từ tạo phòng, tham gia đội, thực hiện cấm/chọn, khai báo build, tính Cost đến hiển thị kết quả Time Handicap. Website phục vụ các vai trò Host/Referee, Player Blue, Player Red và Spectator/Caster.")
    bullet(doc, "Về kỹ thuật - công nghệ: xây dựng hệ thống web theo hướng full-stack với Next.js App Router, TypeScript, Supabase/PostgreSQL, Prisma ORM, Supabase Realtime, Tailwind CSS và Docker. Các quy tắc quan trọng như đúng lượt, đúng đội, không chọn trùng và tính Cost được xử lý ở tầng server/service thay vì đặt rải rác trong giao diện.")
    bullet(doc, "Về khả năng mở rộng: chuẩn bị nền tảng dữ liệu cho tournament, lịch sử trận, thông báo, hồ sơ người chơi, overlay cho caster và thống kê sau giải đấu. Nhờ đó, sản phẩm không chỉ là màn hình chọn nhân vật mà có thể phát triển thành công cụ tổ chức giải đấu cộng đồng.")

    heading(doc, "2. Bảng mô tả ngắn các công nghệ sử dụng")
    paragraph(doc, "Dự án ứng dụng các công nghệ web hiện đại theo từng phân lớp kiến trúc để đảm bảo khả năng phát triển, bảo trì và triển khai:")
    table(doc, ["Phân lớp kiến trúc", "Công nghệ / Framework", "Vai trò cốt lõi trong dự án"], [
        ["Giao diện (Frontend)", "Next.js App Router, React, Tailwind CSS, TypeScript", "Xây dựng giao diện theo component, bố cục Blue/Pool/Red, hỗ trợ responsive và biểu thị rõ trạng thái lượt, đội, nhân vật."],
        ["Dịch vụ (Backend)", "Next.js API Routes, Service/Policy Layer", "Xử lý các use case như tạo phòng, cấm/chọn, khai báo build, tính Cost; kiểm tra quyền và luật nghiệp vụ ở server."],
        ["Dữ liệu (Database)", "Supabase PostgreSQL, Prisma ORM", "Lưu Room, DraftLog, CharacterBuild, Tournament; Prisma hỗ trợ truy vấn có kiểu và mô hình hóa quan hệ dữ liệu."],
        ["Đồng bộ thời gian thực", "Supabase Realtime", "Cập nhật trạng thái phòng, lượt cấm/chọn và build cho nhiều client mà không cần tải lại trang."],
        ["Vận hành (DevOps)", "Docker, Docker Compose", "Đóng gói ứng dụng bằng multi-stage build, hỗ trợ chạy ứng dụng nhất quán giữa môi trường phát triển và triển khai."],
        ["Tiện ích bổ sung", "Zustand, html2canvas", "Quản lý trạng thái giao diện và hỗ trợ xuất ảnh kết quả khi cần chia sẻ minh chứng trận đấu."],
    ])

    heading(doc, "3. Công cụ AI Agents đã sử dụng")
    paragraph(doc, "Trong quá trình phát triển, em sử dụng công cụ AI như một trợ lý hỗ trợ phân tích, lập kế hoạch, rà soát và cải thiện báo cáo. AI không thay thế quá trình thực hiện của sinh viên; các nội dung do AI gợi ý đều cần được đối chiếu với mã nguồn, yêu cầu học phần và sản phẩm thực tế trước khi đưa vào báo cáo.")
    bullet(doc, "Phân tích yêu cầu học phần: sử dụng AI để tóm tắt checklist bắt buộc, xác định các nhóm cần chứng minh như Next.js App Router, Supabase, Docker, deployment, GitHub và minh chứng AI.")
    bullet(doc, "Thiết kế hệ thống: sử dụng AI để hỗ trợ mô tả actor, use case, trạng thái phòng WAITING/DRAFTING/BUILDING/FINISHED, luồng cấm/chọn và quan hệ dữ liệu Room - DraftLog - CharacterBuild - Tournament.")
    bullet(doc, "Rà soát báo cáo: sử dụng AI để kiểm tra lỗi đánh số chương mục, lỗi danh mục bảng/hình, văn phong chưa học thuật và các phần còn thiếu trong báo cáo tóm tắt.")
    bullet(doc, "Hỗ trợ trình bày kỹ thuật: sử dụng AI để diễn giải vai trò của Supabase Realtime, Prisma ORM, Docker multi-stage build, Cost và công thức Time Handicap theo văn phong dễ hiểu hơn.")
    bullet(doc, "Minh chứng prompt: [Cần bổ sung: danh sách prompt thật, thời điểm sử dụng, công cụ AI đã dùng và ảnh chụp minh chứng nếu giảng viên yêu cầu].")

    heading(doc, "4. Bảng phân tích chức năng/module")
    paragraph(doc, "Hệ thống được thiết kế theo hướng module hóa, tập trung vào luồng cấm/chọn cốt lõi và các chức năng mở rộng phục vụ cộng đồng:")
    table(doc, ["Phân hệ (Module)", "Đặc tả tính năng chi tiết"], [
        ["Xác thực & phân quyền", "Đăng nhập, kiểm tra vai trò; Host/Referee điều phối phòng, Player thao tác theo đội, Spectator/Caster chỉ theo dõi, Admin quản lý dữ liệu."],
        ["Phòng đấu", "Tạo mã phòng, lưu trạng thái, cấu hình costPerPoint, bankTime, draftTemplate; quản lý vòng đời WAITING, DRAFTING, BUILDING, FINISHED."],
        ["Cấm/chọn realtime", "Điều phối lượt BAN/PICK; kiểm tra đúng đội, đúng lượt, phòng đúng trạng thái và nhân vật chưa bị khóa; lưu DraftLog để truy vết."],
        ["Build & Cost", "Người chơi khai báo rarity, consLevel, weaponRarity cho nhân vật đã pick; hệ thống tổng hợp Cost từng đội."],
        ["Kết quả", "Tính Time Handicap theo công thức: chênh lệch Cost x số giây mỗi điểm Cost; trình bày kết quả rõ để trọng tài và người chơi đối chiếu."],
        ["Tournament", "Quản lý giải đấu, participant, match và liên kết từng match với roomCode của phòng cấm/chọn."],
        ["Cộng đồng", "Hồ sơ người chơi, lịch sử hoạt động, thông báo, bạn bè và định hướng overlay cho livestream."],
    ])

    heading(doc, "5. Bảng phân tích CSDL (Supabase)")
    paragraph(doc, "Cơ sở dữ liệu sử dụng Supabase PostgreSQL, được mô hình hóa bằng Prisma Schema. Thiết kế dữ liệu bám theo vòng đời trận đấu, trong đó Room là trung tâm, DraftLog ghi lịch sử lượt, CharacterBuild lưu dữ liệu đội hình và Tournament mở rộng từ trận đơn sang giải đấu.")
    table(doc, ["Bảng (Table)", "Vai trò và các trường quan trọng"], [
        ["User", "Quản lý tài khoản, email, tên hiển thị, vai trò và thông tin liên quan đến quyền truy cập."],
        ["Room", "Thực thể trung tâm của trận đấu; chứa code, status, host, blueClientId, redClientId, costPerPoint, bankTime, draftTemplate."],
        ["DraftLog", "Lưu từng lượt cấm/chọn; gồm roomId, player, action, characterId, turnNumber. Hỗ trợ khôi phục lịch sử và tránh tranh chấp."],
        ["CharacterBuild", "Lưu build của nhân vật đã pick; gồm roomId, player, characterId, rarity, consLevel, weaponRarity, totalCost."],
        ["ChatMessage", "Lưu nội dung trao đổi trong phòng, hỗ trợ giao tiếp giữa người chơi và trọng tài."],
        ["Tournament", "Lưu thông tin giải đấu như slug, name, format, status, rulesText, costCap và cấu hình chung."],
        ["TournamentParticipant", "Lưu đội/người tham gia giải, teamName, seed, logoUrl và danh sách thành viên."],
        ["TournamentMatch", "Lưu trận trong bracket, round, matchNumber, roomCode và winnerParticipantId."],
        ["Notification, ActivityEvent, UserSettings", "Hỗ trợ thông báo, lịch sử hoạt động, cài đặt cá nhân và định hướng phát triển cộng đồng."],
    ])

    heading(doc, "6. Bảng mô tả các chức năng đã hoàn thiện")
    paragraph(doc, "Các chức năng đã được xây dựng theo mức độ phù hợp với phạm vi học phần, trong đó luồng cấm/chọn, build, result và realtime là trọng tâm chính:")
    table(doc, ["Phân hệ (Module)", "Tính năng đã hoàn thiện", "Trạng thái triển khai"], [
        ["UI/UX Client", "Giao diện nền tối, bố cục đội Blue/Red và pool nhân vật; hỗ trợ nhận diện trạng thái lượt, nhân vật bị cấm/chọn.", "Đã có nền tảng"],
        ["Xác thực", "Đăng nhập và kiểm tra vai trò người dùng trong các luồng cần bảo vệ.", "Đã triển khai cơ bản"],
        ["Phòng đấu", "Tạo phòng, tham gia đội, lưu trạng thái phòng và cấu hình luật ban đầu.", "Đã hoàn thiện lõi"],
        ["Draft engine", "Kiểm tra đúng lượt, đúng đội, đúng trạng thái phòng; lưu DraftLog và đồng bộ realtime.", "Đã hoàn thiện lõi"],
        ["Build & Result", "Khai báo build, tính Cost, tính Time Handicap và hiển thị kết quả cuối trận.", "Đã hoàn thiện lõi"],
        ["Tournament", "Có dữ liệu và giao diện nền tảng cho giải đấu, participant và match.", "Đã có nền tảng"],
        ["Docker", "Có Dockerfile và Docker Compose để đóng gói, build và chạy ứng dụng.", "Đã triển khai"],
        ["Minh chứng triển khai", "[Cần bổ sung: URL domain/VPS/SSL, ảnh docker compose ps, ảnh trang web chạy thực tế nếu nộp theo yêu cầu triển khai].", "Cần bổ sung minh chứng"],
    ])

    heading(doc, "7. Kết luận, tự nhận xét, đánh giá")
    paragraph(doc, "Kết luận: Đề tài đã xây dựng được một nền tảng website hỗ trợ cấm/chọn nhân vật trong Genshin Impact theo hướng có quy trình, có dữ liệu, có phân quyền vai trò và có đồng bộ realtime. Sản phẩm giải quyết bài toán thực tế của cộng đồng người chơi khi cần tổ chức trận đấu có luật rõ ràng, hạn chế sai sót thủ công và trình bày kết quả Time Handicap minh bạch.")
    paragraph(doc, "Tự nhận xét và đánh giá hệ thống:")
    bullet(doc, "Ưu điểm: Hệ thống có bài toán thực tế, mô hình dữ liệu rõ ràng và luồng nghiệp vụ nhất quán. Việc sử dụng Next.js, TypeScript, Supabase Realtime, Prisma và Docker giúp sản phẩm đáp ứng được nhiều yêu cầu của học phần. Các quy tắc quan trọng được xử lý ở server/service, giúp tăng tính đúng đắn của dữ liệu so với việc chỉ xử lý ở giao diện.")
    bullet(doc, "Hạn chế: Hệ thống chưa thể tự động xác minh tuyệt đối dữ liệu build trong trò chơi, chưa có kiểm thử tải realtime ở quy mô lớn và phần minh chứng triển khai thực tế cần được bổ sung đầy đủ hơn. Một số chức năng như tournament nâng cao, overlay và thống kê sau giải vẫn ở mức định hướng hoặc nền tảng.")
    bullet(doc, "Hướng phát triển tương lai: Ưu tiên hoàn thiện xác minh build, mở rộng tournament cho BO3/BO5, bổ sung overlay cho caster, thống kê tỉ lệ cấm/chọn nhân vật, kiểm thử tự động cho DraftPolicy/CostPolicy và hoàn thiện quy trình deploy trên VPS với domain, SSL, backup và monitoring.")
    paragraph(doc, "Qua quá trình thực hiện, em rút ra bài học rằng việc phân tích nghiệp vụ trước khi xây dựng giao diện là rất quan trọng. Khi actor, use case, trạng thái phòng và dữ liệu được xác định rõ, quá trình triển khai giao diện, API và cơ sở dữ liệu có định hướng hơn, đồng thời báo cáo cũng có cơ sở kỹ thuật chặt chẽ hơn.")


def main():
    doc = Document()
    setup(doc)
    cover(doc)
    teacher_comment(doc)
    content(doc)
    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    main()
