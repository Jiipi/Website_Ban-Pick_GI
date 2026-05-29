from __future__ import annotations

import math
import re
import zipfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Cm, Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "BaoCao_TongHop_Website_BanPick_Genshin_CaiThien_v2.docx"
OUT_DOCX = ROOT / "BaoCao_TongHop_Website_BanPick_Genshin_HoanChinh_StyleGuard_Refs.docx"
ASSET_DIR = ROOT / ".generated_baocao_assets"
MEDIA_DIR = ASSET_DIR / "source_media"


FONT_BODY = "Times New Roman"
FONT_SANS = "Arial"
INK = "111111"
MUTED = "555555"
TABLE_HEADER = "D9EAF7"
TABLE_SUBHEADER = "EEF4FA"
ACCENT_BLUE = "1F4E79"
LIGHT_GRAY = "F3F4F6"


def ensure_dirs() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    MEDIA_DIR.mkdir(parents=True, exist_ok=True)


def extract_media() -> None:
    if any(MEDIA_DIR.iterdir()):
        return
    with zipfile.ZipFile(SOURCE_DOCX) as zf:
        for name in zf.namelist():
            if name.startswith("word/media/"):
                target = MEDIA_DIR / Path(name).name
                target.write_bytes(zf.read(name))


def font_path(name: str = "arial.ttf") -> str:
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/times.ttf"),
    ]
    for p in candidates:
        if p.exists():
            return str(p)
    return name


def draw_centered(draw: ImageDraw.ImageDraw, box, text, font, fill="black", max_width=None, line_gap=4):
    x1, y1, x2, y2 = box
    width = x2 - x1 if max_width is None else max_width
    words = text.split()
    lines = []
    current = ""
    for word in words:
        test = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), test, font=font)[2] <= width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    line_heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total_h = sum(line_heights) + line_gap * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line, h in zip(lines, line_heights):
        tw = draw.textbbox((0, 0), line, font=font)[2]
        draw.text((x1 + ((x2 - x1) - tw) / 2, y), line, font=font, fill=fill)
        y += h + line_gap


def rounded_label(draw, box, text, font, fill, outline="#1f2937"):
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=2)
    draw_centered(draw, box, text, font, fill="black", max_width=(box[2] - box[0] - 20))


def stick_actor(draw, x, y, label, font):
    draw.ellipse((x - 18, y, x + 18, y + 36), outline="black", width=2)
    draw.line((x, y + 36, x, y + 96), fill="black", width=2)
    draw.line((x - 34, y + 58, x + 34, y + 58), fill="black", width=2)
    draw.line((x, y + 96, x - 30, y + 138), fill="black", width=2)
    draw.line((x, y + 96, x + 30, y + 138), fill="black", width=2)
    draw_centered(draw, (x - 80, y + 148, x + 80, y + 190), label, font, max_width=150)


def line_between(draw, p1, p2):
    draw.line((p1[0], p1[1], p2[0], p2[1]), fill="#111111", width=2)


def arrow(draw, p1, p2, fill="#111111", width=2):
    draw.line((p1[0], p1[1], p2[0], p2[1]), fill=fill, width=width)
    angle = math.atan2(p2[1] - p1[1], p2[0] - p1[0])
    size = 12
    for delta in (math.pi * 0.82, -math.pi * 0.82):
        x = p2[0] + size * math.cos(angle + delta)
        y = p2[1] + size * math.sin(angle + delta)
        draw.line((p2[0], p2[1], x, y), fill=fill, width=width)


def create_usecase_diagram(path: Path):
    img = Image.new("RGB", (1400, 900), "white")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(font_path("arialbd.ttf"), 32)
    label_font = ImageFont.truetype(font_path("arial.ttf"), 22)
    actor_font = ImageFont.truetype(font_path("arial.ttf"), 19)
    draw.rectangle((250, 70, 1130, 820), outline="#111111", width=3)
    draw_centered(draw, (250, 80, 1130, 130), "Hệ thống Ban/Pick GI", title_font)

    actors = {
        "host": (120, 140, "Host/Referee"),
        "blue": (120, 380, "Player Blue"),
        "red": (120, 620, "Player Red"),
        "spectator": (1260, 310, "Spectator/Caster"),
        "admin": (1260, 610, "Admin"),
    }
    for x, y, label in actors.values():
        stick_actor(draw, x, y, label, actor_font)

    cases = {
        "create": (390, 170, 650, 235, "Tạo phòng"),
        "control": (710, 170, 1020, 235, "Điều phối trận đấu"),
        "draft": (410, 380, 700, 455, "Tham gia cấm/chọn"),
        "build": (720, 380, 990, 455, "Khai báo build"),
        "watch": (700, 560, 1000, 635, "Theo dõi realtime"),
        "data": (390, 680, 660, 755, "Quản lý dữ liệu"),
        "config": (720, 680, 1010, 755, "Cấu hình hệ thống"),
    }
    for box in cases.values():
        draw.ellipse(box[:4], fill="#F8FAFC", outline="#111111", width=2)
        draw_centered(draw, box[:4], box[4], label_font, max_width=box[2] - box[0] - 28)

    line_between(draw, (190, 205), (390, 200))
    line_between(draw, (190, 205), (710, 200))
    line_between(draw, (190, 445), (410, 418))
    line_between(draw, (190, 445), (720, 418))
    line_between(draw, (190, 685), (410, 418))
    line_between(draw, (190, 685), (720, 418))
    line_between(draw, (1190, 380), (1000, 598))
    line_between(draw, (1190, 680), (660, 718))
    line_between(draw, (1190, 680), (1010, 718))
    img.save(path)


def create_architecture_diagram(path: Path):
    img = Image.new("RGB", (1400, 780), "white")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(font_path("arialbd.ttf"), 32)
    label_font = ImageFont.truetype(font_path("arialbd.ttf"), 24)
    small = ImageFont.truetype(font_path("arial.ttf"), 20)
    draw_centered(draw, (0, 30, 1400, 80), "Kiến trúc phân tầng của website", title_font)
    layers = [
        ("Presentation", "src/app, src/components - giao diện, route, component tương tác", "#DBEAFE"),
        ("Application", "services xử lý use case, điều phối request và kiểm tra trạng thái", "#E0F2FE"),
        ("Domain", "DraftPolicy, CostPolicy, RoomAccessPolicy - quy tắc nghiệp vụ lõi", "#DCFCE7"),
        ("Infrastructure", "Prisma repositories, Supabase gateway, dữ liệu ngoài", "#FEF3C7"),
        ("Composition", "wiring services, cấu hình môi trường và khởi tạo phụ thuộc", "#F3E8FF"),
    ]
    y = 115
    for idx, (name, desc, fill) in enumerate(layers):
        box = (220, y, 1180, y + 92)
        draw.rounded_rectangle(box, radius=12, fill=fill, outline="#1f2937", width=2)
        draw.text((250, y + 18), name, font=label_font, fill="#111111")
        draw.text((500, y + 24), desc, font=small, fill="#111111")
        if idx < len(layers) - 1:
            arrow(draw, (700, y + 92), (700, y + 126), fill="#374151", width=3)
        y += 126
    img.save(path)


def create_data_model_diagram(path: Path):
    img = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(font_path("arialbd.ttf"), 32)
    head_font = ImageFont.truetype(font_path("arialbd.ttf"), 22)
    body_font = ImageFont.truetype(font_path("arial.ttf"), 18)
    draw_centered(draw, (0, 30, 1500, 85), "Mô hình dữ liệu chính theo vòng đời trận đấu", title_font)

    def entity(x, y, w, h, title, fields, fill):
        draw.rounded_rectangle((x, y, x + w, y + h), radius=10, fill=fill, outline="#111111", width=2)
        draw.rectangle((x, y, x + w, y + 44), fill="#E5E7EB", outline="#111111")
        draw_centered(draw, (x, y, x + w, y + 44), title, head_font)
        fy = y + 60
        for f in fields:
            draw.text((x + 18, fy), f, font=body_font, fill="#111111")
            fy += 28
        return (x, y, x + w, y + h)

    room = entity(80, 180, 290, 260, "Room", ["id, code, status", "costPerPoint", "blue/red client", "draftTemplate", "constraints"], "#F8FAFC")
    log = entity(520, 110, 300, 210, "DraftLog", ["roomId", "player", "action", "characterId", "turnNumber"], "#EFF6FF")
    build = entity(520, 370, 300, 230, "CharacterBuild", ["roomId", "player", "characterId", "rarity", "totalCost"], "#EFF6FF")
    chat = entity(520, 650, 300, 170, "ChatMessage", ["roomId", "sender", "message", "role"], "#EFF6FF")
    tour = entity(980, 150, 330, 230, "Tournament", ["slug, name, format", "status, maxTeams", "costCap, bankTime", "rulesText"], "#F0FDF4")
    participant = entity(1030, 450, 330, 190, "TournamentParticipant", ["tournamentId", "playerUid", "teamName", "seed"], "#F0FDF4")
    match = entity(1030, 690, 330, 170, "TournamentMatch", ["tournamentId", "round", "roomCode", "winner"], "#F0FDF4")

    arrow(draw, (370, 240), (520, 205))
    draw.text((390, 192), "1-n", font=body_font, fill="#111111")
    arrow(draw, (370, 315), (520, 480))
    draw.text((390, 380), "1-n", font=body_font, fill="#111111")
    arrow(draw, (370, 390), (520, 720))
    draw.text((390, 590), "1-n", font=body_font, fill="#111111")
    arrow(draw, (1145, 380), (1160, 450))
    draw.text((1185, 405), "1-n", font=body_font, fill="#111111")
    arrow(draw, (1160, 640), (1160, 690))
    draw.text((1185, 657), "1-n", font=body_font, fill="#111111")
    draw.text((74, 118), "Lobby/Social liên kết người dùng theo UID/clientId", font=body_font, fill="#374151")
    img.save(path)


def create_flow_diagram(path: Path):
    img = Image.new("RGB", (1100, 1100), "white")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(font_path("arialbd.ttf"), 31)
    label_font = ImageFont.truetype(font_path("arial.ttf"), 23)
    small = ImageFont.truetype(font_path("arial.ttf"), 19)
    draw_centered(draw, (0, 35, 1100, 90), "Luồng xử lý một lượt cấm/chọn", title_font)
    boxes = [
        (330, 130, 770, 195, "Client chọn nhân vật", "rect"),
        (330, 250, 770, 315, "API Draft nhận request", "rect"),
        (290, 370, 810, 445, "DraftService kiểm tra phòng và quyền", "rect"),
        (420, 510, 680, 610, "Hợp lệ?", "diamond"),
        (290, 690, 810, 765, "DraftPolicy kiểm tra lượt và nhân vật", "rect"),
        (330, 830, 770, 895, "Repository lưu DraftLog", "rect"),
        (330, 960, 770, 1025, "Realtime cập nhật giao diện", "rect"),
    ]
    for x1, y1, x2, y2, text, kind in boxes:
        if kind == "diamond":
            pts = [(550, y1), (x2, (y1 + y2) // 2), (550, y2), (x1, (y1 + y2) // 2)]
            draw.polygon(pts, fill="#FEF3C7", outline="#111111")
            draw_centered(draw, (x1, y1, x2, y2), text, label_font)
        else:
            rounded_label(draw, (x1, y1, x2, y2), text, label_font, "#F8FAFC")
    arrow(draw, (550, 195), (550, 250))
    arrow(draw, (550, 315), (550, 370))
    arrow(draw, (550, 445), (550, 510))
    arrow(draw, (550, 610), (550, 690))
    arrow(draw, (550, 765), (550, 830))
    arrow(draw, (550, 895), (550, 960))
    arrow(draw, (420, 560), (185, 560))
    rounded_label(draw, (45, 515, 185, 605), "Trả lỗi", small, "#FEE2E2")
    draw.text((330, 628), "Có", font=small, fill="#111111")
    draw.text((245, 530), "Không", font=small, fill="#111111")
    img.save(path)


def create_layout_diagram(path: Path):
    img = Image.new("RGB", (1400, 760), "white")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(font_path("arialbd.ttf"), 32)
    label_font = ImageFont.truetype(font_path("arialbd.ttf"), 25)
    small = ImageFont.truetype(font_path("arial.ttf"), 21)
    draw_centered(draw, (0, 30, 1400, 90), "Bố cục MOBA layout Blue/Pool/Red", title_font)
    draw.text((90, 130), "Desktop", font=label_font, fill="#111111")
    panels = [
        (90, 180, 405, 530, "Đội Xanh", "#DBEAFE"),
        (435, 180, 965, 530, "Pool nhân vật", "#F8FAFC"),
        (995, 180, 1310, 530, "Đội Đỏ", "#FEE2E2"),
    ]
    for box in panels:
        draw.rounded_rectangle(box[:4], radius=10, fill=box[5], outline="#111111", width=2)
        draw_centered(draw, box[:4], box[4], label_font)
    draw.text((90, 585), "Mobile", font=label_font, fill="#111111")
    mobile = [
        (250, 585, 570, 660, "Pool nhân vật", "#F8FAFC"),
        (610, 585, 920, 660, "Đội Xanh", "#DBEAFE"),
        (960, 585, 1270, 660, "Đội Đỏ", "#FEE2E2"),
    ]
    for box in mobile:
        draw.rounded_rectangle(box[:4], radius=10, fill=box[5], outline="#111111", width=2)
        draw_centered(draw, box[:4], box[4], small)
    img.save(path)


def create_diagrams() -> dict[str, Path]:
    diagrams = {
        "usecase": ASSET_DIR / "hinh_3_1_usecase.png",
        "architecture": ASSET_DIR / "hinh_3_2_architecture.png",
        "data": ASSET_DIR / "hinh_3_3_data_model.png",
        "flow": ASSET_DIR / "hinh_3_4_banpick_flow.png",
        "layout": ASSET_DIR / "hinh_3_5_layout.png",
    }
    create_usecase_diagram(diagrams["usecase"])
    create_architecture_diagram(diagrams["architecture"])
    create_data_model_diagram(diagrams["data"])
    create_flow_diagram(diagrams["flow"])
    create_layout_diagram(diagrams["layout"])
    return diagrams


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def length_to_dxa(length) -> int:
    return int(round(length.twips))


def content_width_dxa(doc: Document) -> int:
    section = doc.sections[0]
    return length_to_dxa(section.page_width) - length_to_dxa(section.left_margin) - length_to_dxa(section.right_margin)


def widths_from_weights(weights, total_width):
    total = float(sum(weights))
    out = [int(round(total_width * (w / total))) for w in weights]
    out[-1] += total_width - sum(out)
    return out


def ensure_child(parent, tag):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_width_node(parent, tag, width):
    node = ensure_child(parent, tag)
    node.set(qn("w:type"), "dxa")
    node.set(qn("w:w"), str(int(width)))


def apply_exact_table_geometry(doc: Document, table, weights=None, indent_dxa=120):
    total_width = content_width_dxa(doc)
    if weights is None:
        weights = [1] * len(table.columns)
    col_widths = widths_from_weights(weights, total_width)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    set_width_node(tbl_pr, "w:tblW", total_width)
    ind = ensure_child(tbl_pr, "w:tblInd")
    ind.set(qn("w:type"), "dxa")
    ind.set(qn("w:w"), str(indent_dxa))
    layout = ensure_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in col_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for idx, width in enumerate(col_widths):
        table.columns[idx].width = Twips(width)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Twips(col_widths[idx])
            set_width_node(cell._tc.get_or_add_tcPr(), "w:tcW", col_widths[idx])


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_layout_fixed(table):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def add_field(paragraph, instr: str):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def set_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(13)
    normal.font.color.rgb = RGBColor(17, 17, 17)
    pf = normal.paragraph_format
    pf.first_line_indent = Cm(1.0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)

    for name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = styles[name]
        style.font.name = FONT_BODY
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Heading 1"].paragraph_format.space_before = Pt(12)
    styles["Heading 1"].paragraph_format.space_after = Pt(12)
    styles["Heading 1"].paragraph_format.keep_with_next = True
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].paragraph_format.keep_with_next = True
    styles["Heading 3"].font.size = Pt(13)
    styles["Heading 3"].font.italic = True
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    styles["Heading 3"].paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = FONT_BODY
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    caption.font.size = Pt(12)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    for list_style in ["List Bullet", "List Number"]:
        st = styles[list_style]
        st.font.name = FONT_BODY
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        st.font.size = Pt(13)
        st.paragraph_format.first_line_indent = None
        st.paragraph_format.left_indent = Cm(1.25)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.5


def footer_page_numbers(doc: Document):
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Trang ")
    run.font.name = FONT_BODY
    run.font.size = Pt(11)
    add_field(footer, "PAGE")


def p(doc, text="", style=None, bold=False, italic=False, align=None, first_line=True):
    para = doc.add_paragraph(style=style)
    if not first_line:
        para.paragraph_format.first_line_indent = Cm(0)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = FONT_BODY
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    return para


def add_hyperlink(paragraph, url: str, text: str | None = None):
    rel_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_BODY)
    r_fonts.set(qn("w:hAnsi"), FONT_BODY)
    r_fonts.set(qn("w:eastAsia"), FONT_BODY)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "26")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.append(r_fonts)
    run_pr.append(size)
    run_pr.append(color)
    run_pr.append(underline)
    run.append(run_pr)

    text_node = OxmlElement("w:t")
    text_node.text = text or url
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_heading(doc, text, level=1, page_break=False):
    if page_break and len(doc.paragraphs) > 0:
        doc.add_page_break()
    para = doc.add_paragraph(text, style=f"Heading {level}")
    para.paragraph_format.first_line_indent = Cm(0)
    if level == 1:
        para.runs[0].text = text.upper()
    return para


def add_bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(item, style="List Bullet")
        para.paragraph_format.first_line_indent = None


def add_numbered(doc, items):
    for item in items:
        para = doc.add_paragraph(item, style="List Number")
        para.paragraph_format.first_line_indent = None


def add_table(doc, headers, rows, widths=None, header_fill=TABLE_HEADER, font_size=11):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_layout_fixed(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.name = FONT_BODY
        run.font.size = Pt(font_size)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
    for row_values in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.first_line_indent = Cm(0)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for part_i, part in enumerate(str(value).split("\n")):
                if part_i:
                    para.add_run().add_break()
                run = para.add_run(part)
                run.font.name = FONT_BODY
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
                run.font.size = Pt(font_size)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
    if widths:
        apply_exact_table_geometry(doc, table, widths)
    else:
        apply_exact_table_geometry(doc, table)
    doc.add_paragraph()
    return table


def shade_table_column(table, col_idx, fills_by_text):
    for row in table.rows[1:]:
        cell = row.cells[col_idx]
        text = cell.text.strip().split()[0] if cell.text.strip() else ""
        fill = fills_by_text.get(cell.text.strip()) or fills_by_text.get(text)
        if fill:
            set_cell_shading(cell, fill)


def add_caption(doc, text, kind="Hình"):
    para = doc.add_paragraph(f"{kind} {text}", style="Caption")
    para.paragraph_format.first_line_indent = Cm(0)
    return para


def add_image(doc, image_path: Path, caption: str, width_cm=15.5):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.keep_with_next = True
    run = para.add_run()
    inline_shape = run.add_picture(str(image_path), width=Cm(width_cm))
    inline_shape._inline.docPr.set("descr", caption)
    inline_shape._inline.docPr.set("title", caption)
    add_caption(doc, caption)


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    para = cell.paragraphs[0]
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(text.splitlines()):
        if idx:
            para.add_run().add_break()
        run = para.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
    apply_exact_table_geometry(doc, table, [1], indent_dxa=160)
    doc.add_paragraph()


def add_cover(doc: Document):
    logo = MEDIA_DIR / "image1.jpg"
    p(doc, "TRƯỜNG ĐẠI HỌC ĐÀ LẠT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    p(doc, "KHOA CÔNG NGHỆ THÔNG TIN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    if logo.exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Cm(0)
        inline_shape = para.add_run().add_picture(str(logo), width=Cm(3.4))
        inline_shape._inline.docPr.set("descr", "Logo Trường Đại học Đà Lạt")
        inline_shape._inline.docPr.set("title", "Logo Trường Đại học Đà Lạt")
    for _ in range(2):
        doc.add_paragraph()
    p(doc, "BÁO CÁO TỔNG KẾT HỌC PHẦN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    p(doc, "HỌC PHẦN: CÁC CÔNG NGHỆ MỚI TRONG PHÁT TRIỂN PHẦN MỀM", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    doc.add_paragraph()
    title = p(
        doc,
        "ĐỀ TÀI: XÂY DỰNG WEBSITE HỖ TRỢ CẤM/CHỌN NHÂN VẬT TRONG GENSHIN IMPACT",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
    )
    title.runs[0].font.size = Pt(16)
    doc.add_paragraph()
    info = [
        ("Giảng viên hướng dẫn:", "KS. Nguyễn Trọng Hiếu"),
        ("Sinh viên thực hiện:", "2212377 - Trần Ngọc Hưng"),
        ("Lớp:", "Công nghệ thông tin"),
        ("Thời gian thực hiện:", "Tháng 05 năm 2026"),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    for i, (k, v) in enumerate(info):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        for cell in table.row_cells(i):
            for para in cell.paragraphs:
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.name = FONT_BODY
                    run.font.size = Pt(13)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(i, 0).paragraphs[0].runs[0].bold = True
    apply_exact_table_geometry(doc, table, [4.5, 8.5])
    doc.add_paragraph()
    for _ in range(4):
        doc.add_paragraph()
    p(doc, "Lâm Đồng, tháng 05 năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)


def add_front_matter(doc: Document):
    add_heading(doc, "NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN", 1, page_break=True)
    for _ in range(20):
        para = p(doc, "." * 118, first_line=False)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_heading(doc, "LỜI CẢM ƠN", 1, page_break=True)
    p(doc, "Trong quá trình thực hiện đề tài, em xin chân thành cảm ơn quý thầy cô Khoa Công nghệ Thông tin, Trường Đại học Đà Lạt đã truyền đạt những kiến thức nền tảng về phân tích thiết kế hệ thống, lập trình web, cơ sở dữ liệu và triển khai phần mềm. Đây là cơ sở quan trọng giúp em có thể tiếp cận đề tài theo hướng hệ thống, không chỉ dừng lại ở việc xây dựng giao diện mà còn chú trọng đến quy trình nghiệp vụ, tổ chức dữ liệu và khả năng vận hành thực tế.")
    p(doc, "Em xin gửi lời cảm ơn đến thầy Nguyễn Trọng Hiếu đã định hướng nội dung học phần và góp ý trong quá trình lựa chọn, phân tích và hoàn thiện đề tài. Những định hướng của thầy giúp em nhìn nhận website hỗ trợ cấm/chọn nhân vật Genshin Impact như một sản phẩm phần mềm có nhiều vai trò người dùng, có trạng thái thời gian thực, có quy tắc nghiệp vụ và có yêu cầu triển khai cụ thể.")
    p(doc, "Trong quá trình xây dựng sản phẩm, em có cơ hội vận dụng nhiều công nghệ hiện đại như Next.js, TypeScript, Prisma, Supabase, Tailwind CSS, Docker và các công cụ hỗ trợ phát triển. Bên cạnh kết quả sản phẩm, quá trình thực hiện giúp em rèn luyện khả năng phân tích bài toán, tổ chức mã nguồn, kiểm thử luồng người dùng và trình bày kết quả theo văn phong báo cáo kỹ thuật.")
    p(doc, "Do giới hạn về thời gian và kinh nghiệm, đề tài khó tránh khỏi thiếu sót. Em rất mong nhận được góp ý từ quý thầy cô để tiếp tục hoàn thiện sản phẩm và nâng cao chất lượng báo cáo trong các phiên bản tiếp theo.")
    p(doc, "Em xin chân thành cảm ơn!", align=WD_ALIGN_PARAGRAPH.RIGHT)

    add_heading(doc, "MỤC LỤC", 1, page_break=True)
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0)
    add_field(para, r'TOC \o "1-3" \h \z \u')
    p(doc, "Ghi chú: Khi mở bằng Microsoft Word, chọn Update Field để cập nhật số trang tự động nếu Word chưa tự cập nhật.", italic=True, first_line=False)

    add_heading(doc, "DANH MỤC HÌNH ẢNH", 1, page_break=True)
    figures = [
        "Hình 3.1. Sơ đồ Use Case tổng quan của hệ thống",
        "Hình 3.2. Kiến trúc phân tầng của website",
        "Hình 3.3. Mô hình dữ liệu chính theo vòng đời trận đấu",
        "Hình 3.4. Luồng xử lý một lượt cấm/chọn",
        "Hình 3.5. Bố cục MOBA layout ba cột Blue/Pool/Red",
        "Hình 4.1. Giao diện chọn vai trò và bắt đầu phiên sử dụng",
        "Hình 4.2. Giao diện phòng cấm/chọn trong trạng thái thi đấu",
        "Hình 4.3. Giao diện tổng hợp sau lượt cấm/chọn",
        "Hình 4.4. Giao diện kết quả và Time Handicap",
        "Hình 4.5. Giao diện giải đấu và danh sách người tham gia",
        "Hình 4.6. Giao diện hồ sơ, thống kê và lịch sử hoạt động",
    ]
    for item in figures:
        p(doc, item, first_line=False)

    add_heading(doc, "DANH MỤC BẢNG", 1, page_break=True)
    tables = [
        "Bảng 0.1. Danh mục từ viết tắt",
        "Bảng 1.1. Phạm vi thực hiện của đề tài",
        "Bảng 2.1. Vai trò của các công nghệ trong dự án",
        "Bảng 3.1. Tác nhân và nhu cầu sử dụng hệ thống",
        "Bảng 3.2. Yêu cầu chức năng chính",
        "Bảng 3.3. Yêu cầu phi chức năng",
        "Bảng 3.4. Use Case chính của hệ thống",
        "Bảng 3.5. Thiết kế chi tiết một số bảng dữ liệu chính",
        "Bảng 3.6. Nhóm API tiêu biểu của hệ thống",
        "Bảng 3.7. Bảng màu chủ đạo của giao diện",
        "Bảng 3.8. Quy ước trạng thái và phản hồi tương tác",
        "Bảng 4.1. Tổ chức mã nguồn theo tầng",
        "Bảng 4.2. Kết quả kiểm thử chức năng",
        "Bảng 5.1. Đối chiếu mục tiêu ban đầu và kết quả đạt được",
        "Bảng A.1. Minh chứng sử dụng công cụ AI trong quá trình phát triển",
    ]
    for item in tables:
        p(doc, item, first_line=False)

    add_heading(doc, "DANH MỤC TỪ VIẾT TẮT", 1, page_break=True)
    add_caption(doc, "0.1. Danh mục từ viết tắt", "Bảng")
    add_table(
        doc,
        ["Từ viết tắt", "Tên đầy đủ", "Ý nghĩa trong báo cáo"],
        [
            ["API", "Application Programming Interface", "Giao diện lập trình ứng dụng, dùng để trao đổi dữ liệu giữa client và server."],
            ["CSDL", "Cơ sở dữ liệu", "Nơi lưu trữ các thông tin về phòng đấu, lượt cấm/chọn, build, giải đấu và người dùng."],
            ["CRUD", "Create, Read, Update, Delete", "Nhóm thao tác cơ bản đối với dữ liệu trong hệ thống."],
            ["DTO", "Data Transfer Object", "Cấu trúc dữ liệu truyền giữa tầng giao diện, API và service."],
            ["ORM", "Object Relational Mapping", "Cơ chế ánh xạ giữa mã nguồn TypeScript và bảng dữ liệu quan hệ thông qua Prisma."],
            ["Realtime", "Real-time synchronization", "Cơ chế đồng bộ trạng thái gần thời gian thực cho các vai trò đang cùng theo dõi một phòng."],
            ["UID", "User Identifier", "Mã định danh người chơi Genshin Impact hoặc định danh client trong phạm vi hệ thống."],
        ],
        widths=[3.0, 4.7, 8.0],
    )


def chapter_1(doc):
    add_heading(doc, "CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI", 1, page_break=True)
    add_heading(doc, "1.1. Bối cảnh hình thành đề tài", 2)
    add_heading(doc, "1.1.1. Bối cảnh cộng đồng Genshin Impact", 3)
    p(doc, "Genshin Impact là trò chơi nhập vai hành động có cộng đồng người chơi lớn, trong đó nhiều hoạt động giao lưu, thử thách và thi đấu cộng đồng được tổ chức bên ngoài trò chơi chính thức. Các hoạt động này thường dựa trên luật chơi do cộng đồng xây dựng, chẳng hạn giới hạn nhân vật, giới hạn vũ khí, cấm/chọn trước trận hoặc quy đổi lợi thế tài khoản thành mức bù trừ thời gian.")
    p(doc, "Trong bối cảnh thi đấu cộng đồng, yếu tố công bằng không chỉ phụ thuộc vào kỹ năng điều khiển nhân vật mà còn chịu ảnh hưởng bởi tài nguyên tài khoản của mỗi người chơi. Một người chơi sở hữu nhiều nhân vật 5 sao, cung mệnh cao hoặc vũ khí 5 sao có thể có lợi thế rõ rệt. Vì vậy, các giải đấu cộng đồng thường cần cơ chế kiểm soát lựa chọn đội hình và cơ chế quy đổi lợi thế thành Time Handicap.")
    add_heading(doc, "1.1.2. Đặc thù của quy trình cấm/chọn", 3)
    p(doc, "Quy trình cấm/chọn là giai đoạn tiền trận, trong đó hai đội lần lượt loại bỏ hoặc lựa chọn nhân vật theo một thứ tự đã được quy định. Đây là giai đoạn có tính chiến thuật cao vì mỗi lựa chọn đều ảnh hưởng đến đội hình, tài nguyên còn lại và khả năng thi đấu của đối thủ. Nếu quy trình không được ghi nhận rõ ràng, người chơi và trọng tài rất dễ tranh luận về lượt hiện tại, nhân vật đã bị cấm hoặc quyền thao tác của từng bên.")
    p(doc, "Đối với trận đấu có phát sóng hoặc có khán giả theo dõi, yêu cầu về minh bạch còn cao hơn. Người xem cần biết đội nào đang thao tác, nhân vật nào đã bị cấm, nhân vật nào đã được chọn, trạng thái trận đấu đang ở giai đoạn nào và kết quả tính Cost được hình thành từ dữ liệu nào. Những thông tin này khó được trình bày mạch lạc nếu chỉ sử dụng bảng tính hoặc tin nhắn rời rạc.")
    add_heading(doc, "1.1.3. Hạn chế của cách tổ chức thủ công", 3)
    p(doc, "Trước khi có công cụ chuyên dụng, ban tổ chức thường kết hợp Discord, bảng tính, hình ảnh nhân vật và ghi chú thủ công để điều phối trận đấu. Cách làm này có thể đáp ứng các trận nhỏ, nhưng bộc lộ nhiều hạn chế khi trận đấu có nhiều lượt, nhiều người theo dõi hoặc cần lưu lại kết quả. Trọng tài phải đồng thời kiểm tra lượt, cập nhật danh sách nhân vật, tính Cost và giải thích kết quả cho người chơi.")
    p(doc, "Sai sót trong một lượt cấm/chọn có thể làm gián đoạn toàn bộ trận đấu. Tương tự, sai sót khi tính Cost hoặc Time Handicap có thể ảnh hưởng trực tiếp đến kết quả thi đấu. Do đó, việc xây dựng một website chuyên dụng nhằm chuẩn hóa quy trình, lưu lại dữ liệu và hiển thị trạng thái theo thời gian thực là nhu cầu có cơ sở thực tiễn.")
    add_heading(doc, "1.2. Lý do chọn đề tài", 2)
    add_heading(doc, "1.2.1. Giá trị thực tiễn của sản phẩm", 3)
    p(doc, "Đề tài được lựa chọn vì xuất phát từ một nhu cầu cụ thể của cộng đồng: cần một công cụ hỗ trợ tổ chức phiên cấm/chọn nhân vật Genshin Impact có luật, có vai trò điều phối và có dữ liệu đối chiếu. Website không thay thế trò chơi và không can thiệp vào máy chủ của nhà phát hành; hệ thống chỉ hỗ trợ phần tổ chức bên ngoài trò chơi, tương tự một công cụ điều phối giải đấu.")
    add_heading(doc, "1.2.2. Sự phù hợp với học phần", 3)
    p(doc, "Về mặt học thuật, đề tài phù hợp với học phần Các công nghệ mới trong phát triển phần mềm vì sử dụng nhiều thành phần của ứng dụng web hiện đại: frontend component hóa, API route, cơ sở dữ liệu quan hệ, realtime, ORM, triển khai container và dịch vụ đám mây. Bài toán không dừng ở chức năng CRUD đơn giản mà có trạng thái nghiệp vụ phức tạp, ràng buộc quyền thao tác và yêu cầu đồng bộ giữa nhiều người dùng.")
    add_heading(doc, "1.2.3. Tính mở rộng của bài toán", 3)
    p(doc, "Từ luồng cấm/chọn cơ bản, hệ thống có thể mở rộng thành nền tảng hỗ trợ giải đấu cộng đồng với các chức năng như lobby, tournament, hồ sơ người chơi, lịch sử trận đấu, thông báo, overlay cho phát sóng và thống kê. Tính mở rộng này giúp đề tài có chiều sâu, đồng thời tạo điều kiện đánh giá khả năng tổ chức kiến trúc phần mềm trong quá trình phát triển.")
    add_heading(doc, "1.3. Mục tiêu của đề tài", 2)
    add_heading(doc, "1.3.1. Mục tiêu nghiệp vụ", 3)
    p(doc, "Mục tiêu nghiệp vụ của đề tài là xây dựng được một luồng tổ chức trận đấu hoàn chỉnh. Trọng tài có thể tạo phòng, mời người chơi, cấu hình luật, bắt đầu phiên cấm/chọn, chuyển sang giai đoạn khai báo build, tổng hợp Cost và hiển thị kết quả. Người chơi có thể thao tác đúng lượt, khai báo thông tin nhân vật và xem kết quả một cách minh bạch.")
    add_heading(doc, "1.3.2. Mục tiêu kỹ thuật", 3)
    p(doc, "Mục tiêu kỹ thuật là xây dựng hệ thống có cấu trúc rõ ràng, tách biệt giữa tầng giao diện, tầng xử lý use case, tầng quy tắc miền và tầng truy cập dữ liệu. Các quy tắc quan trọng như kiểm tra lượt, kiểm tra quyền thao tác, không cho chọn trùng nhân vật và tính Cost cần được đặt ở tầng service hoặc policy, thay vì xử lý rải rác trong giao diện.")
    add_heading(doc, "1.3.3. Mục tiêu sản phẩm", 3)
    p(doc, "Mục tiêu sản phẩm là tạo ra website có giao diện trực quan, hỗ trợ thao tác nhanh trong bối cảnh thi đấu và có khả năng theo dõi realtime. Giao diện cần phân tách rõ đội xanh, đội đỏ, pool nhân vật, trạng thái lượt và khu vực kết quả. Bên cạnh đó, hệ thống cần có nền tảng dữ liệu đủ tốt để mở rộng các chức năng cộng đồng trong tương lai.")
    add_heading(doc, "1.4. Phạm vi thực hiện", 2)
    add_heading(doc, "1.4.1. Phạm vi chức năng", 3)
    p(doc, "Trong phạm vi học phần, đề tài tập trung vào các chức năng liên quan trực tiếp đến vòng đời một trận đấu: tạo phòng, tham gia theo vai trò, thực hiện cấm/chọn, khai báo build, tính Cost, tính Time Handicap, hiển thị kết quả và lưu lại dữ liệu chính. Các chức năng lobby, tournament, hồ sơ người chơi và social được xem là phần mở rộng nhằm chứng minh định hướng phát triển của hệ thống.")
    add_heading(doc, "1.4.2. Giới hạn ngoài phạm vi", 3)
    p(doc, "Hệ thống chưa đặt mục tiêu thay thế cơ chế chống gian lận trong trò chơi. Dữ liệu build có thể được người chơi khai báo thủ công hoặc hỗ trợ lấy một phần từ nguồn ngoài nếu người chơi công khai thông tin phù hợp. Việc xác minh tuyệt đối toàn bộ thông tin trong game không thuộc phạm vi phiên bản hiện tại do giới hạn API chính thức và phạm vi triển khai học phần.")
    add_caption(doc, "1.1. Phạm vi thực hiện của đề tài", "Bảng")
    add_table(
        doc,
        ["Nhóm nội dung", "Trong phạm vi", "Ngoài phạm vi hiện tại"],
        [
            ["Phòng đấu", "Tạo phòng, quản lý trạng thái, phân vai Host/Player/Spectator và khôi phục trạng thái khi tải lại trang.", "Tổ chức nhiều máy chủ realtime phân tán hoặc cân bằng tải quy mô lớn."],
            ["Cấm/chọn", "Kiểm tra lượt, kiểm tra quyền, lưu DraftLog và đồng bộ giao diện cho các vai trò đang theo dõi.", "Tự động điều khiển hoặc can thiệp vào client Genshin Impact."],
            ["Build và Cost", "Khai báo rarity, cung mệnh, vũ khí, refinement và tính tổng Cost theo luật cấu hình.", "Xác minh tuyệt đối toàn bộ dữ liệu tài khoản nếu nguồn dữ liệu ngoài không cung cấp."],
            ["Giải đấu", "Thiết kế nền tảng tournament, participant, match và roomCode để mở rộng theo bracket.", "Vận hành giải đấu lớn với thanh toán, livestream tích hợp đầy đủ hoặc kiểm thử tải chuyên sâu."],
        ],
        widths=[3.3, 6.2, 6.2],
    )
    add_heading(doc, "1.5. Phương pháp thực hiện", 2)
    add_heading(doc, "1.5.1. Phân tích nghiệp vụ trước khi hiện thực", 3)
    p(doc, "Quá trình thực hiện bắt đầu từ việc mô tả quy trình tổ chức một trận cấm/chọn trong cộng đồng, sau đó chuyển quy trình này thành actor, use case, yêu cầu chức năng, yêu cầu phi chức năng và mô hình dữ liệu. Cách tiếp cận này giúp hệ thống được xây dựng dựa trên nghiệp vụ thực tế thay vì chỉ mô phỏng giao diện bên ngoài.")
    add_heading(doc, "1.5.2. Thiết kế theo hướng phân tầng", 3)
    p(doc, "Sau khi xác định yêu cầu, đề tài tổ chức mã nguồn theo các tầng có trách nhiệm rõ ràng. Tầng giao diện tiếp nhận thao tác và hiển thị trạng thái; tầng application điều phối use case; tầng domain chứa policy kiểm tra luật; tầng infrastructure chịu trách nhiệm truy cập cơ sở dữ liệu và tích hợp dịch vụ ngoài. Cách tổ chức này giúp hệ thống dễ kiểm thử và mở rộng.")
    add_heading(doc, "1.5.3. Kiểm thử theo luồng người dùng", 3)
    p(doc, "Kiểm thử được thực hiện theo các luồng end-to-end như tạo phòng, tham gia đội, thực hiện cấm/chọn, khai báo build và xem kết quả. Đối với hệ thống có trạng thái realtime, kiểm thử từng API riêng lẻ chưa đủ để chứng minh tính đúng đắn của sản phẩm; cần đánh giá toàn bộ chuỗi thao tác của người dùng trong cùng một phiên trận đấu.")
    p(doc, "Trong quá trình kiểm thử, mỗi kịch bản được đặt trong bối cảnh có nhiều vai trò cùng tham gia. Ví dụ, khi Player Blue thực hiện lượt chọn, hệ thống không chỉ cần kiểm tra API có lưu DraftLog hay không, mà còn cần quan sát giao diện của Player Red, Host và Spectator có nhận đúng trạng thái mới hay không. Cách kiểm thử này phù hợp với đặc thù của ứng dụng realtime, nơi lỗi thường xuất hiện ở sự không đồng bộ giữa các client.")
    p(doc, "Ngoài các kịch bản thành công, đề tài cũng chú trọng kiểm thử các trường hợp bị từ chối như thao tác sai lượt, chọn lại nhân vật đã bị khóa, người xem cố gửi request điều khiển phòng hoặc người chơi tải lại trang giữa trận. Những trường hợp này giúp đánh giá khả năng bảo vệ trạng thái chính thức của hệ thống và giảm rủi ro khi sản phẩm được sử dụng trong một trận đấu có người theo dõi.")
    add_heading(doc, "1.6. Ý nghĩa thực tiễn của đề tài", 2)
    add_heading(doc, "1.6.1. Ý nghĩa đối với trọng tài và ban tổ chức", 3)
    p(doc, "Đối với trọng tài, website giúp giảm khối lượng thao tác thủ công trong quá trình điều phối trận đấu. Trọng tài không phải liên tục đối chiếu bảng tính, ghi chú lượt hiện tại hoặc tự tính lại chênh lệch Cost bằng tay. Khi trạng thái được hệ thống ghi nhận và hiển thị, trọng tài có nhiều thời gian hơn để xử lý tình huống phát sinh, giải thích luật và bảo đảm nhịp độ trận đấu.")
    p(doc, "Đối với ban tổ chức, việc lưu trữ dữ liệu có cấu trúc giúp quá trình tổng hợp kết quả và tra cứu lịch sử thuận tiện hơn. Sau mỗi trận, dữ liệu về nhân vật bị cấm, nhân vật được chọn, đội hình, Cost và kết quả có thể được dùng để thống kê hoặc làm minh chứng khi có khiếu nại. Đây là điểm khác biệt đáng kể so với cách lưu dữ liệu rời rạc bằng ảnh chụp hoặc tin nhắn.")
    add_heading(doc, "1.6.2. Ý nghĩa đối với người chơi và khán giả", 3)
    p(doc, "Đối với người chơi, hệ thống tạo ra môi trường thao tác rõ ràng hơn. Người chơi biết khi nào đến lượt của họ, thao tác nào hợp lệ và vì sao một thao tác bị từ chối. Điều này giúp giảm tranh luận không cần thiết, đồng thời giúp người chơi tập trung vào chiến thuật lựa chọn đội hình thay vì kiểm tra thủ công trạng thái trận đấu.")
    p(doc, "Đối với khán giả và caster, giao diện realtime giúp quá trình theo dõi trận đấu trở nên mạch lạc hơn. Thay vì phải nghe trọng tài giải thích liên tục, người xem có thể quan sát trực tiếp danh sách nhân vật đã cấm, đã chọn, lượt hiện tại và kết quả bù trừ thời gian. Đây là yếu tố quan trọng nếu trận đấu được phát trực tiếp trong cộng đồng.")
    add_heading(doc, "1.7. Cấu trúc báo cáo", 2)
    p(doc, "Báo cáo được tổ chức thành năm chương. Chương 1 trình bày bối cảnh, lý do chọn đề tài, mục tiêu, phạm vi và phương pháp thực hiện. Chương 2 phân tích các công nghệ được sử dụng và lý do lựa chọn. Chương 3 trình bày phần phân tích, thiết kế hệ thống, bao gồm actor, use case, yêu cầu chức năng, kiến trúc, dữ liệu, API và giao diện. Chương 4 mô tả quá trình xây dựng, triển khai và kiểm thử. Chương 5 tổng kết kết quả đạt được, hạn chế và hướng phát triển.")
    p(doc, "Cách tổ chức này nhằm bảo đảm mạch trình bày đi từ nhu cầu thực tế đến giải pháp kỹ thuật, sau đó mới đánh giá kết quả hiện thực. Nhờ đó, người đọc có thể thấy rõ mối liên hệ giữa vấn đề cộng đồng đặt ra, quyết định thiết kế hệ thống và các chức năng đã được xây dựng trong sản phẩm.")
    add_heading(doc, "1.8. Đóng góp của đề tài", 2)
    add_heading(doc, "1.8.1. Đóng góp về mặt nghiệp vụ", 3)
    p(doc, "Đóng góp đầu tiên của đề tài là mô hình hóa được một quy trình tổ chức trận đấu cộng đồng vốn thường được thực hiện bằng nhiều công cụ rời rạc. Thay vì xem cấm/chọn chỉ là thao tác chọn nhân vật trên giao diện, báo cáo phân tích quy trình này như một chuỗi nghiệp vụ có trạng thái, có vai trò, có luật và có dữ liệu cần lưu vết. Cách nhìn này giúp hệ thống có cơ sở thiết kế rõ ràng hơn, đồng thời giúp người đọc hiểu vì sao các thành phần như Room, DraftLog, CharacterBuild và Tournament đều cần thiết.")
    p(doc, "Ở góc độ tổ chức giải đấu, đề tài đề xuất cách chuyển dữ liệu trận đấu từ thông tin tạm thời sang dữ liệu có cấu trúc. Mỗi lượt cấm/chọn, mỗi build và mỗi kết quả Cost đều có thể được ghi lại để đối chiếu. Điều này có giá trị thực tế vì các trận đấu cộng đồng thường phát sinh tranh luận khi kết quả chỉ được ghi nhận bằng ảnh chụp, tin nhắn hoặc bảng tính cá nhân của trọng tài.")
    add_heading(doc, "1.8.2. Đóng góp về mặt kỹ thuật", 3)
    p(doc, "Về kỹ thuật, đề tài thể hiện cách áp dụng kiến trúc web hiện đại vào một bài toán có nhiều trạng thái đồng thời. Việc sử dụng Next.js giúp gom giao diện và API trong cùng một dự án; TypeScript giúp kiểm soát kiểu dữ liệu; Prisma và PostgreSQL hỗ trợ mô hình dữ liệu quan hệ; Supabase Realtime đáp ứng yêu cầu đồng bộ trạng thái cho nhiều người dùng. Những công nghệ này không được liệt kê như danh sách công cụ, mà được gắn với vai trò cụ thể trong từng phần của hệ thống.")
    p(doc, "Một đóng góp khác là việc tách quy tắc nghiệp vụ ra khỏi giao diện. Các quy tắc như đúng lượt, không chọn trùng nhân vật, kiểm tra quyền thao tác và tính Cost đều có thể thay đổi theo từng giải đấu. Nếu các quy tắc này nằm trực tiếp trong component, hệ thống sẽ khó bảo trì. Vì vậy, báo cáo nhấn mạnh vai trò của service và policy như lớp trung gian bảo vệ tính đúng đắn của dữ liệu.")
    add_heading(doc, "1.8.3. Đóng góp về mặt sản phẩm", 3)
    p(doc, "Về mặt sản phẩm, website hướng đến trải nghiệm sử dụng trong bối cảnh thi đấu thật. Người dùng không chỉ cần một giao diện đẹp, mà cần biết chính xác trạng thái hiện tại, lượt tiếp theo, thao tác hợp lệ và kết quả cuối cùng. Do đó, thiết kế giao diện ưu tiên khả năng nhận diện đội, trạng thái nhân vật, thông tin Cost và khu vực kết quả. Đây là hướng tiếp cận phù hợp với công cụ hỗ trợ thi đấu, khác với website giới thiệu thông tin thông thường.")
    p(doc, "Ngoài luồng cấm/chọn cốt lõi, đề tài cũng mở ra định hướng phát triển thành nền tảng cộng đồng thông qua tournament, profile, social, notification và overlay. Các chức năng mở rộng này chưa phải trọng tâm tuyệt đối của phiên bản hiện tại, nhưng chúng cho thấy sản phẩm có khả năng phát triển dài hạn nếu tiếp tục được đầu tư sau học phần.")


def chapter_2(doc):
    add_heading(doc, "CHƯƠNG 2: CÔNG NGHỆ VÀ CÔNG CỤ PHÁT TRIỂN", 1, page_break=True)
    add_heading(doc, "2.1. Next.js App Router và React", 2)
    add_heading(doc, "2.1.1. Vai trò của Next.js trong kiến trúc full-stack", 3)
    p(doc, "Next.js được sử dụng làm framework chính để phát triển ứng dụng web. Với App Router, hệ thống có thể tổ chức route theo thư mục, kết hợp giao diện người dùng và API route trong cùng một dự án. Điều này phù hợp với phạm vi học phần vì giảm độ phức tạp khi triển khai nhưng vẫn cho phép tách các xử lý nghiệp vụ ra thành service riêng.")
    add_heading(doc, "2.1.2. Component hóa giao diện với React", 3)
    p(doc, "React hỗ trợ chia giao diện thành các component nhỏ như DraftBoard, CharacterCard, CostCalculator, khu vực đội xanh, khu vực đội đỏ, bảng kết quả và thanh điều hướng. Component hóa giúp mã nguồn dễ đọc, giảm lặp lại và thuận tiện khi thay đổi giao diện theo từng giai đoạn của trận đấu.")
    add_heading(doc, "2.1.3. Lý do lựa chọn trong đề tài", 3)
    p(doc, "Website có nhiều màn hình chức năng như trang chủ, lobby, phòng cấm/chọn, trang build, trang kết quả, overlay và tournament. Cấu trúc route của Next.js giúp ánh xạ tương đối trực tiếp giữa thư mục mã nguồn và chức năng người dùng, từ đó hỗ trợ cả quá trình phát triển lẫn quá trình trình bày trong báo cáo.")
    add_heading(doc, "2.2. TypeScript và kiểm soát dữ liệu", 2)
    add_heading(doc, "2.2.1. Kiểu dữ liệu trong luồng nghiệp vụ", 3)
    p(doc, "TypeScript bổ sung hệ thống kiểu tĩnh cho JavaScript, giúp giảm lỗi khi dữ liệu luân chuyển giữa giao diện, API, service và cơ sở dữ liệu. Trong đề tài, các giá trị như đội BLUE/RED, hành động BAN/PICK, trạng thái phòng WAITING/DRAFTING/BUILDING/FINISHED và cấu trúc build cần được mô tả rõ để tránh truyền sai dữ liệu.")
    add_heading(doc, "2.2.2. Giới hạn của kiểm tra kiểu", 3)
    p(doc, "TypeScript không thay thế kiểm thử nghiệp vụ. Công cụ này giúp phát hiện lỗi kiểu dữ liệu khi phát triển, nhưng các quy tắc như đúng lượt, không chọn trùng nhân vật, kiểm tra quyền thao tác hoặc tính Time Handicap vẫn cần được triển khai trong service/policy và kiểm chứng bằng kịch bản kiểm thử cụ thể.")
    add_heading(doc, "2.3. Tailwind CSS và thiết kế giao diện thi đấu", 2)
    add_heading(doc, "2.3.1. Thiết kế theo trạng thái", 3)
    p(doc, "Tailwind CSS hỗ trợ xây dựng giao diện nhanh thông qua các lớp tiện ích. Đối với website cấm/chọn, giao diện cần thể hiện rõ trạng thái lượt hiện tại, nhân vật đã bị cấm, nhân vật đã được chọn, đội đang thao tác và kết quả Cost. Tailwind phù hợp với yêu cầu này vì cho phép mô tả màu sắc, khoảng cách, layout và trạng thái tương tác trực tiếp trong component.")
    add_heading(doc, "2.3.2. Định hướng responsive", 3)
    p(doc, "Giao diện phòng đấu được thiết kế theo định hướng MOBA layout. Trên desktop, ba khu vực đội xanh, pool nhân vật và đội đỏ được hiển thị song song để người dùng so sánh nhanh. Trên mobile, nội dung được sắp xếp theo thứ tự ưu tiên, trong đó pool nhân vật và trạng thái lượt được đặt ở vị trí dễ thao tác.")
    add_heading(doc, "2.4. Supabase, PostgreSQL và Realtime", 2)
    add_heading(doc, "2.4.1. PostgreSQL làm nền tảng lưu trữ", 3)
    p(doc, "PostgreSQL được sử dụng làm hệ quản trị cơ sở dữ liệu quan hệ chính. Các thực thể như Room, DraftLog, CharacterBuild, ChatMessage, Tournament và TournamentMatch có quan hệ rõ ràng, do đó cơ sở dữ liệu quan hệ phù hợp hơn so với cách lưu trữ phi cấu trúc. Việc dùng PostgreSQL cũng giúp hệ thống thuận lợi khi cần truy vấn lịch sử, thống kê và mở rộng dữ liệu giải đấu.")
    add_heading(doc, "2.4.2. Realtime trong phòng đấu", 3)
    p(doc, "Supabase Realtime hỗ trợ đồng bộ thay đổi dữ liệu cho các client đang theo dõi cùng một phòng. Khi một người chơi thực hiện lượt cấm/chọn, trạng thái cần được cập nhật cho trọng tài, đối thủ và khán giả. Cơ chế realtime giúp giảm thao tác tải lại trang và tăng tính minh bạch trong quá trình thi đấu.")
    add_heading(doc, "2.4.3. Khả năng khôi phục trạng thái", 3)
    p(doc, "Dữ liệu realtime cần được kết hợp với lưu trữ bền vững. Trạng thái chính thức của trận đấu vẫn được ghi trong PostgreSQL, nhờ vậy khi người dùng mất kết nối tạm thời hoặc tải lại trang, hệ thống có thể khôi phục trạng thái từ cơ sở dữ liệu thay vì phụ thuộc vào trạng thái tạm thời trên client.")
    add_heading(doc, "2.5. Prisma ORM", 2)
    add_heading(doc, "2.5.1. Vai trò ánh xạ dữ liệu", 3)
    p(doc, "Prisma ORM đóng vai trò ánh xạ giữa mô hình dữ liệu trong mã nguồn TypeScript và bảng dữ liệu PostgreSQL. Các model như Room, DraftLog, CharacterBuild, Tournament hoặc Notification được khai báo trong Prisma Schema, từ đó sinh Prisma Client để truy vấn dữ liệu có kiểu rõ ràng.")
    add_heading(doc, "2.5.2. Lợi ích trong bảo trì", 3)
    p(doc, "Khi cấu trúc dữ liệu thay đổi, Prisma giúp người phát triển theo dõi thay đổi schema và hạn chế lỗi truy vấn thủ công. Trong đề tài có nhiều quan hệ một-nhiều và nhiều ràng buộc unique/index, ORM giúp mã nguồn truy cập dữ liệu nhất quán hơn, đặc biệt ở các service xử lý phòng đấu và giải đấu.")
    add_heading(doc, "2.6. Dữ liệu ngoài Enka và Genshin.dev", 2)
    add_heading(doc, "2.6.1. Mục đích tích hợp", 3)
    p(doc, "Website có thể sử dụng dữ liệu nhân vật từ các nguồn ngoài nhằm giảm lượng nhập liệu thủ công và tăng tính nhất quán khi hiển thị nhân vật. Genshin.dev hỗ trợ dữ liệu nhân vật cơ bản, trong khi Enka.network có thể cung cấp một số thông tin hồ sơ nếu người chơi công khai dữ liệu phù hợp.")
    add_heading(doc, "2.6.2. Giới hạn dữ liệu ngoài", 3)
    p(doc, "Nguồn dữ liệu ngoài không phải lúc nào cũng đầy đủ hoặc ổn định. Vì vậy, hệ thống cần cho phép khai báo thủ công và cần sự kiểm tra của trọng tài trong các trường hợp dữ liệu không thể xác minh tự động. Báo cáo không xem cơ chế tích hợp dữ liệu ngoài là giải pháp chống gian lận hoàn chỉnh, mà là công cụ hỗ trợ giảm thao tác và tăng tính tiện lợi.")
    add_heading(doc, "2.7. Docker và triển khai ứng dụng", 2)
    add_heading(doc, "2.7.1. Đóng gói môi trường chạy", 3)
    p(doc, "Docker được sử dụng để đóng gói ứng dụng, giảm sự khác biệt giữa môi trường phát triển và môi trường triển khai. Dockerfile multi-stage giúp tách giai đoạn cài đặt phụ thuộc, build ứng dụng và chạy production, từ đó làm bản triển khai gọn hơn và dễ tái lập hơn.")
    add_heading(doc, "2.7.2. Cấu hình vận hành", 3)
    p(doc, "Docker Compose hỗ trợ mô tả các container, cổng truy cập và biến môi trường cần thiết. Đối với một ứng dụng sử dụng cơ sở dữ liệu ngoài, Supabase và các khóa API, việc quản lý biến môi trường có vai trò quan trọng nhằm tránh đưa thông tin nhạy cảm vào mã nguồn.")
    p(doc, "Trong môi trường phát triển, Docker giúp kiểm tra khả năng build và chạy ứng dụng trong điều kiện gần với production hơn so với việc chỉ chạy bằng lệnh dev. Điều này đặc biệt hữu ích khi dự án sử dụng nhiều phụ thuộc như Prisma Client, Next.js build output và biến môi trường. Nếu ứng dụng chỉ hoạt động trong môi trường máy cá nhân nhưng lỗi khi đóng gói, khả năng vận hành thực tế sẽ bị ảnh hưởng.")
    add_heading(doc, "2.8. Công cụ hỗ trợ phát triển và quản lý chất lượng", 2)
    add_heading(doc, "2.8.1. Zustand trong quản lý trạng thái giao diện", 3)
    p(doc, "Zustand được sử dụng như một thư viện quản lý trạng thái nhẹ cho những phần giao diện cần chia sẻ dữ liệu giữa nhiều component. Trong một phòng cấm/chọn, các trạng thái như người dùng hiện tại, đội đang quan sát, pool nhân vật, bộ lọc hoặc trạng thái thao tác tạm thời có thể được quản lý tập trung để tránh truyền props quá sâu giữa các component.")
    p(doc, "Việc sử dụng state management ở client không có nghĩa là client được quyền quyết định trạng thái chính thức của trận đấu. State phía client chỉ phục vụ trải nghiệm hiển thị và thao tác; dữ liệu chính thức vẫn phải lấy từ server hoặc cơ sở dữ liệu. Cách phân biệt này giúp hệ thống vừa có giao diện phản hồi nhanh, vừa giữ được tính đúng đắn của nghiệp vụ.")
    add_heading(doc, "2.8.2. html2canvas và nhu cầu xuất kết quả", 3)
    p(doc, "Thư viện html2canvas hỗ trợ chuyển một vùng giao diện HTML thành ảnh. Trong ngữ cảnh đề tài, chức năng này có thể được dùng để xuất ảnh kết quả trận đấu, ảnh đội hình hoặc ảnh tóm tắt Cost. Đây là nhu cầu thực tế của cộng đồng vì kết quả thường được chia sẻ lại trong nhóm chat, bài đăng thông báo hoặc kênh tổ chức giải.")
    p(doc, "Khi sử dụng chức năng xuất ảnh, giao diện kết quả cần được thiết kế ổn định về kích thước, màu sắc và khả năng hiển thị font. Nếu bố cục phụ thuộc quá nhiều vào chiều rộng màn hình hoặc có phần tử tràn ra ngoài, ảnh xuất ra sẽ khó dùng làm minh chứng. Vì vậy, thiết kế trang kết quả cần ưu tiên cấu trúc rõ ràng và có khả năng chụp lại nhất quán.")
    add_heading(doc, "2.8.3. ESLint và quy ước mã nguồn", 3)
    p(doc, "ESLint hỗ trợ phát hiện một số lỗi phổ biến trong mã nguồn JavaScript/TypeScript, chẳng hạn import không dùng, biểu thức có khả năng gây lỗi hoặc vi phạm quy ước của Next.js. Trong một dự án có nhiều màn hình và nhiều service, công cụ lint giúp duy trì chất lượng mã nguồn ổn định hơn khi chức năng được mở rộng.")
    p(doc, "Tuy nhiên, lint chỉ là một lớp kiểm tra tĩnh. Các lỗi nghiệp vụ như sai thứ tự lượt, tính Cost sai hoặc quyền thao tác không đúng vẫn phải được kiểm tra bằng test case và rà soát logic service. Vì vậy, công cụ quản lý chất lượng cần được kết hợp với thiết kế kiến trúc và quy trình kiểm thử, thay vì xem như giải pháp duy nhất.")
    add_caption(doc, "2.1. Vai trò của các công nghệ trong dự án", "Bảng")
    add_table(
        doc,
        ["Công nghệ", "Vai trò trong hệ thống", "Lý do phù hợp với đề tài"],
        [
            ["Next.js", "Tổ chức route, giao diện và API route trong cùng dự án.", "Phù hợp với ứng dụng full-stack quy mô học phần nhưng vẫn có khả năng triển khai thực tế."],
            ["React", "Xây dựng component giao diện cho board cấm/chọn, thẻ nhân vật và kết quả.", "Giúp giao diện phức tạp được chia nhỏ, dễ tái sử dụng và bảo trì."],
            ["TypeScript", "Kiểm soát kiểu dữ liệu giữa client, API, service và Prisma.", "Giảm lỗi truyền sai cấu trúc dữ liệu trong luồng có nhiều trạng thái."],
            ["Tailwind CSS", "Thiết kế layout, màu sắc và trạng thái tương tác.", "Phù hợp giao diện thi đấu cần nhận diện nhanh đội, lượt và trạng thái nhân vật."],
            ["Supabase Realtime", "Đồng bộ trạng thái phòng gần thời gian thực.", "Đáp ứng nhu cầu nhiều vai trò cùng theo dõi một phiên trận đấu."],
            ["PostgreSQL/Prisma", "Lưu trữ và truy vấn dữ liệu quan hệ.", "Phù hợp với quan hệ Room - DraftLog - Build - Tournament."],
            ["Docker", "Đóng gói và triển khai ứng dụng.", "Giúp môi trường chạy có khả năng tái lập và dễ cấu hình hơn."],
            ["Zustand", "Quản lý trạng thái giao diện cần chia sẻ giữa nhiều component.", "Phù hợp với phòng đấu có nhiều vùng giao diện phụ thuộc cùng một trạng thái."],
            ["html2canvas", "Xuất một phần giao diện thành ảnh kết quả.", "Phục vụ nhu cầu chia sẻ kết quả trận đấu trong cộng đồng."],
            ["ESLint", "Kiểm tra quy ước và lỗi tĩnh trong mã nguồn.", "Hỗ trợ duy trì chất lượng mã nguồn khi chức năng mở rộng."],
        ],
        widths=[3.3, 6.0, 6.4],
    )
    add_heading(doc, "2.9. Đánh giá tổng hợp về lựa chọn công nghệ", 2)
    add_heading(doc, "2.9.1. Sự phù hợp giữa công nghệ và nghiệp vụ realtime", 3)
    p(doc, "Điểm đáng chú ý trong đề tài là công nghệ được lựa chọn dựa trên đặc điểm nghiệp vụ, không chỉ dựa trên độ phổ biến. Bài toán cấm/chọn có nhiều người dùng theo dõi cùng một trạng thái và có yêu cầu cập nhật nhanh sau mỗi lượt thao tác. Vì vậy, việc kết hợp Next.js với Supabase Realtime và PostgreSQL tạo ra một cấu trúc phù hợp: Next.js xử lý giao diện và API, PostgreSQL lưu trạng thái chính thức, còn Realtime truyền tín hiệu thay đổi cho các client.")
    p(doc, "Nếu chỉ sử dụng một ứng dụng frontend tĩnh và lưu dữ liệu cục bộ, hệ thống khó đáp ứng yêu cầu nhiều người cùng theo dõi một phòng. Ngược lại, nếu xây dựng một backend riêng quá phức tạp ngay từ đầu, phạm vi đề tài sẽ vượt quá thời gian học phần. Cách lựa chọn hiện tại cân bằng giữa tính thực tế và khả năng triển khai: đủ mạnh để chứng minh luồng nghiệp vụ, nhưng không tạo ra quá nhiều gánh nặng hạ tầng.")
    add_heading(doc, "2.9.2. Sự phù hợp giữa TypeScript, Prisma và thiết kế dữ liệu", 3)
    p(doc, "Dữ liệu trong hệ thống có nhiều thực thể liên kết với nhau. Room liên kết đến DraftLog, CharacterBuild và ChatMessage; Tournament liên kết đến Participant và Match; Notification và ActivityEvent liên quan đến trải nghiệm cộng đồng. Với cấu trúc như vậy, TypeScript và Prisma hỗ trợ người phát triển kiểm soát dữ liệu tốt hơn so với việc viết truy vấn SQL rời rạc ở nhiều nơi.")
    p(doc, "Prisma Schema đồng thời đóng vai trò như một tài liệu kỹ thuật về cơ sở dữ liệu. Khi đọc schema, người phát triển có thể biết bảng nào có quan hệ một-nhiều, bảng nào có ràng buộc unique và trường nào được đánh index. Điều này đặc biệt hữu ích trong báo cáo vì mô hình dữ liệu không chỉ được trình bày bằng sơ đồ mà còn có cơ sở đối chiếu với mã nguồn thực tế.")
    add_heading(doc, "2.9.3. Sự phù hợp giữa Tailwind CSS và giao diện thi đấu", 3)
    p(doc, "Trong một website thi đấu, giao diện cần thể hiện trạng thái nhanh và rõ hơn là trang trí phức tạp. Tailwind CSS giúp xây dựng các trạng thái giao diện như highlight lượt, màu đội xanh/đỏ, thẻ nhân vật bị khóa, nút đang chờ xử lý hoặc thông báo lỗi. Việc mô tả trạng thái bằng utility class giúp người phát triển dễ điều chỉnh giao diện theo từng component mà không phải quản lý quá nhiều file CSS riêng.")
    p(doc, "Tuy nhiên, Tailwind cũng có rủi ro nếu sử dụng thiếu quy ước. Khi class bị lặp lại quá nhiều hoặc mỗi component tự định nghĩa màu sắc riêng, giao diện dễ mất nhất quán. Vì vậy, báo cáo bổ sung Design System và bảng màu chủ đạo để xác định vai trò từng màu trong hệ thống. Đây là bước cần thiết để giao diện không chỉ hoạt động được mà còn có tính nhận diện và tính chuyên nghiệp.")
    add_heading(doc, "2.9.4. Sự phù hợp giữa Docker và yêu cầu triển khai", 3)
    p(doc, "Docker giúp quá trình triển khai có tính tái lập cao hơn, đặc biệt khi ứng dụng cần build Next.js, generate Prisma Client và đọc biến môi trường. Trong bối cảnh học phần, Docker không chỉ là công cụ vận hành mà còn là minh chứng rằng sản phẩm có thể được đóng gói thành một ứng dụng chạy độc lập, thay vì chỉ hoạt động trong môi trường phát triển cá nhân.")
    p(doc, "Dù vậy, Docker không giải quyết toàn bộ vấn đề vận hành. Khi hệ thống sử dụng dịch vụ ngoài như Supabase hoặc nguồn dữ liệu nhân vật, quá trình triển khai vẫn cần kiểm tra kết nối, quyền truy cập, realtime channel và biến môi trường. Do đó, báo cáo xem Docker là một thành phần trong quy trình triển khai, không phải là giải pháp duy nhất cho toàn bộ bài toán vận hành.")


def chapter_3(doc, diagrams):
    add_heading(doc, "CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1, page_break=True)
    add_heading(doc, "3.1. Quy trình nghiệp vụ tổng quát", 2)
    add_heading(doc, "3.1.1. Giai đoạn chuẩn bị phòng", 3)
    p(doc, "Quy trình bắt đầu khi Host hoặc trọng tài tạo phòng đấu. Phòng có mã truy cập, trạng thái ban đầu, cấu hình luật, thời gian bank time và hệ số quy đổi Cost sang thời gian. Sau khi phòng được tạo, người chơi hai bên tham gia theo vai trò Blue hoặc Red, trong khi khán giả hoặc caster có thể theo dõi nếu phòng được mở công khai.")
    add_heading(doc, "3.1.2. Giai đoạn cấm/chọn", 3)
    p(doc, "Khi trọng tài bắt đầu trận, hệ thống chuyển phòng sang trạng thái cấm/chọn. Mỗi lượt gồm thông tin đội thực hiện, hành động BAN hoặc PICK, số lượt hiện tại và danh sách nhân vật còn hợp lệ. Client chỉ gửi yêu cầu khi đến lượt của người dùng tương ứng, còn server là nơi quyết định thao tác có hợp lệ hay không.")
    add_heading(doc, "3.1.3. Giai đoạn build và tổng kết", 3)
    p(doc, "Sau khi hoàn tất cấm/chọn, người chơi khai báo build cho các nhân vật đã chọn. Hệ thống tính tổng Cost từng bên, so sánh chênh lệch và quy đổi thành Time Handicap theo cấu hình phòng. Kết quả được hiển thị trên trang tổng kết để trọng tài, người chơi và khán giả có thể đối chiếu.")
    add_heading(doc, "3.1.4. Luồng dữ liệu xuyên suốt vòng đời trận đấu", 3)
    p(doc, "Một điểm quan trọng trong thiết kế là dữ liệu của trận đấu không xuất hiện rời rạc ở từng màn hình. Khi phòng được tạo, Room đóng vai trò bản ghi trung tâm. Khi người chơi thao tác cấm/chọn, DraftLog bổ sung lịch sử theo từng lượt. Khi chuyển sang giai đoạn build, CharacterBuild bổ sung dữ liệu về đội hình và Cost. Khi trận đấu thuộc một giải, TournamentMatch liên kết phòng đấu với vòng đấu và kết quả. Nhờ vậy, vòng đời trận đấu có thể được truy vết từ lúc chuẩn bị đến lúc tổng kết.")
    p(doc, "Cách tổ chức dữ liệu này giúp hệ thống tránh tình trạng mỗi màn hình tự lưu một kiểu dữ liệu riêng. Nếu dữ liệu phòng, dữ liệu lượt và dữ liệu kết quả không cùng nằm trong một mô hình quan hệ rõ ràng, việc khôi phục trạng thái sau khi tải lại trang hoặc thống kê sau giải đấu sẽ rất khó thực hiện. Vì vậy, thiết kế dữ liệu là nền tảng để bảo đảm tính liên tục của toàn bộ sản phẩm.")
    add_heading(doc, "3.2. Tác nhân hệ thống", 2)
    p(doc, "Hệ thống được thiết kế cho nhiều vai trò sử dụng cùng lúc. Việc phân tách actor giúp xác định quyền thao tác, nhu cầu hiển thị và phạm vi dữ liệu của từng nhóm người dùng.")
    add_caption(doc, "3.1. Tác nhân và nhu cầu sử dụng hệ thống", "Bảng")
    add_table(
        doc,
        ["Tác nhân", "Vai trò", "Nhu cầu chính"],
        [
            ["Host/Referee", "Người tạo phòng và điều phối trận đấu.", "Tạo phòng, cấu hình luật, chuyển phase, xử lý tranh chấp và xác nhận kết quả."],
            ["Player Blue", "Người chơi thuộc đội xanh.", "Tham gia phòng, thực hiện lượt cấm/chọn, khai báo build và xem kết quả."],
            ["Player Red", "Người chơi thuộc đội đỏ.", "Thao tác tương tự Player Blue nhưng theo lượt của đội đỏ."],
            ["Spectator/Caster", "Người xem hoặc bình luận viên.", "Theo dõi trạng thái trận đấu realtime, xem pool nhân vật, lượt hiện tại và kết quả."],
            ["Admin", "Người quản trị dữ liệu hệ thống.", "Quản lý dữ liệu nhân vật, bảng Cost, cấu hình hệ thống và các thông tin mở rộng."],
        ],
        widths=[3.3, 5.0, 7.4],
    )
    add_image(doc, diagrams["usecase"], "3.1. Sơ đồ Use Case tổng quan của hệ thống", 15.7)
    add_heading(doc, "3.3. Yêu cầu chức năng", 2)
    add_heading(doc, "3.3.1. Nhóm chức năng phòng đấu", 3)
    p(doc, "Nhóm chức năng phòng đấu chịu trách nhiệm tạo mã phòng, lưu trạng thái, quản lý người tham gia và kiểm soát phase. Đây là nền tảng để các chức năng cấm/chọn, build và kết quả hoạt động trên cùng một trạng thái dữ liệu.")
    add_heading(doc, "3.3.2. Nhóm chức năng cấm/chọn", 3)
    p(doc, "Nhóm chức năng cấm/chọn cần bảo đảm thao tác chỉ được chấp nhận khi đúng lượt, đúng đội, đúng trạng thái phòng và nhân vật chưa bị cấm hoặc chọn trước đó. Mọi thao tác hợp lệ được lưu thành DraftLog để có thể khôi phục lịch sử và hiển thị lại trạng thái.")
    add_heading(doc, "3.3.3. Nhóm chức năng build và kết quả", 3)
    p(doc, "Nhóm chức năng build cho phép người chơi khai báo rarity, cung mệnh, rarity vũ khí, refinement và nguồn dữ liệu. Từ dữ liệu này, hệ thống tính tổng Cost từng bên và kết quả Time Handicap. Kết quả cần được trình bày rõ để người chơi có thể đối chiếu.")
    add_caption(doc, "3.2. Yêu cầu chức năng chính", "Bảng")
    add_table(
        doc,
        ["Mã", "Yêu cầu", "Mô tả", "Màn hình/API liên quan"],
        [
            ["FR-01", "Tạo phòng", "Host tạo phòng có mã truy cập, trạng thái ban đầu, cấu hình Cost và luật draft.", "Trang lobby, API room"],
            ["FR-02", "Tham gia phòng", "Người chơi tham gia theo đội Blue/Red; caster hoặc spectator tham gia chế độ theo dõi.", "Trang room, lobby"],
            ["FR-03", "Cấm/chọn đúng lượt", "Server kiểm tra lượt hiện tại, quyền thao tác, trạng thái nhân vật và ghi DraftLog.", "API draft, DraftBoard"],
            ["FR-04", "Khai báo build", "Player khai báo thông tin nhân vật và vũ khí cho các nhân vật đã pick.", "Build page, API build"],
            ["FR-05", "Tính Cost", "Hệ thống tính tổng Cost, chênh lệch Cost và Time Handicap theo cấu hình phòng.", "CostService, Result page"],
            ["FR-06", "Theo dõi realtime", "Các client trong cùng phòng nhận cập nhật khi DraftLog, build hoặc trạng thái phòng thay đổi.", "Supabase channel"],
            ["FR-07", "Quản lý giải đấu", "Tạo tournament, participant, match và liên kết match với roomCode khi bắt đầu trận.", "Tournaments page"],
        ],
        widths=[1.6, 3.0, 7.2, 3.9],
        font_size=10.5,
    )
    add_heading(doc, "3.4. Yêu cầu phi chức năng", 2)
    add_caption(doc, "3.3. Yêu cầu phi chức năng", "Bảng")
    add_table(
        doc,
        ["Nhóm", "Yêu cầu", "Cách đáp ứng trong thiết kế"],
        [
            ["Tính đúng đắn", "Không chấp nhận thao tác sai lượt, sai quyền hoặc chọn trùng nhân vật.", "Đặt kiểm tra ở tầng DraftPolicy và RoomAccessPolicy phía server."],
            ["Tính minh bạch", "Người chơi và trọng tài có thể đối chiếu từng lượt và kết quả Cost.", "Lưu DraftLog, CharacterBuild và hiển thị kết quả tổng hợp."],
            ["Realtime", "Thay đổi trạng thái cần được cập nhật cho các client liên quan.", "Sử dụng Supabase Realtime kết hợp truy vấn lại trạng thái chính thức."],
            ["Khả năng mở rộng", "Có thể mở rộng từ phòng đơn sang tournament, social và overlay.", "Thiết kế model Tournament, TournamentMatch, Notification và ActivityEvent."],
            ["Bảo trì", "Mã nguồn cần có phân tầng, hạn chế logic nghiệp vụ nằm trong component.", "Tách service, policy, repository và component giao diện."],
            ["Bảo mật", "Không để người dùng tự ý chuyển phase hoặc sửa dữ liệu không thuộc quyền.", "Kiểm tra hostUserId/clientId/role tại API và service."],
        ],
        widths=[3.0, 5.5, 7.2],
    )
    add_heading(doc, "3.5. Use Case chính", 2)
    add_heading(doc, "3.5.1. Cách xác định Use Case", 3)
    p(doc, "Use Case của hệ thống được xác định từ các vai trò thực tế trong một trận đấu cộng đồng. Mỗi use case phải gắn với một mục tiêu cụ thể của actor, không chỉ mô tả tên màn hình. Ví dụ, 'Thực hiện cấm/chọn' không đơn thuần là thao tác trên một thẻ nhân vật, mà bao gồm việc xác định lượt hiện tại, kiểm tra quyền, ghi nhận lịch sử và đồng bộ trạng thái cho các bên liên quan.")
    p(doc, "Việc mô tả use case theo mục tiêu giúp tránh tình trạng thiết kế chức năng rời rạc. Một màn hình có thể phục vụ nhiều use case, và một use case có thể liên quan đến nhiều thành phần như API, service, database và realtime. Do đó, phần đặc tả use case trong báo cáo được trình bày như cầu nối giữa yêu cầu nghiệp vụ và thiết kế kỹ thuật.")
    add_caption(doc, "3.4. Use Case chính của hệ thống", "Bảng")
    add_table(
        doc,
        ["Use Case", "Tác nhân", "Tiền điều kiện", "Kết quả mong muốn"],
        [
            ["Tạo phòng đấu", "Host/Referee", "Người dùng có quyền tạo phòng.", "Phòng được tạo với mã truy cập, cấu hình luật và trạng thái WAITING."],
            ["Tham gia đội", "Player Blue/Red", "Phòng tồn tại và còn vị trí phù hợp.", "Người chơi được gán vào đội, trạng thái phòng cập nhật cho các client."],
            ["Thực hiện cấm/chọn", "Player Blue/Red", "Phòng đang ở phase DRAFTING và đến lượt đội tương ứng.", "DraftLog được lưu; pool nhân vật và lượt tiếp theo được cập nhật."],
            ["Khai báo build", "Player Blue/Red", "Draft đã hoàn tất và nhân vật thuộc danh sách pick của đội.", "CharacterBuild được lưu; tổng Cost được tính lại."],
            ["Theo dõi trận đấu", "Spectator/Caster", "Phòng cho phép theo dõi hoặc có quyền truy cập.", "Giao diện chỉ xem hiển thị trạng thái realtime và kết quả."],
            ["Quản lý cấu hình", "Admin", "Người dùng có quyền quản trị.", "Dữ liệu nhân vật, bảng Cost hoặc thông tin hệ thống được cập nhật có kiểm soát."],
        ],
        widths=[3.0, 3.0, 5.5, 4.2],
        font_size=10.5,
    )
    add_heading(doc, "3.5.2. Đặc tả Use Case tạo phòng", 3)
    p(doc, "Use case tạo phòng bắt đầu khi Host truy cập lobby và yêu cầu tạo một phòng mới. Hệ thống sinh mã phòng, thiết lập trạng thái WAITING, lưu cấu hình mặc định như costPerPoint, bankTime, draftTemplate và gắn thông tin host nếu có. Kết quả trả về là mã phòng để Host chia sẻ cho người chơi hoặc caster.")
    p(doc, "Điều kiện thành công của use case này không chỉ là tạo được một bản ghi Room. Phòng được tạo phải đủ dữ liệu để chuyển sang các giai đoạn sau, bao gồm thông tin trạng thái, khả năng gán người chơi và cấu hình luật. Nếu thiếu các giá trị mặc định cần thiết, lỗi có thể xuất hiện muộn hơn ở giai đoạn cấm/chọn hoặc tính kết quả.")
    add_heading(doc, "3.5.3. Đặc tả Use Case thực hiện cấm/chọn", 3)
    p(doc, "Use case thực hiện cấm/chọn là use case có rủi ro nghiệp vụ cao nhất. Người chơi gửi lựa chọn, nhưng hệ thống chỉ chấp nhận khi phòng đang ở trạng thái phù hợp, actor đúng đội, lượt hiện tại đúng hành động và nhân vật chưa bị khóa. Nếu bất kỳ điều kiện nào không thỏa, API phải trả về lỗi có thông điệp đủ rõ để người dùng biết nguyên nhân.")
    p(doc, "Sau khi thao tác hợp lệ được ghi nhận, hệ thống cần cập nhật lượt kế tiếp và phát tín hiệu realtime. Kết quả của use case không chỉ là một dòng DraftLog mới, mà là sự thay đổi trạng thái của toàn bộ phòng. Vì vậy, service xử lý use case cần trả về trạng thái đã chuẩn hóa để giao diện không phải tự suy diễn quá nhiều.")
    add_heading(doc, "3.5.4. Đặc tả Use Case khai báo build", 3)
    p(doc, "Use case khai báo build diễn ra sau khi phase draft hoàn tất. Người chơi chỉ được khai báo build cho các nhân vật mà đội của họ đã pick. Điều kiện này giúp tránh trường hợp người dùng gửi build cho nhân vật không thuộc đội, từ đó làm sai kết quả Cost. Mỗi lần cập nhật build cần ghi lại thời điểm cập nhật để trọng tài có thể đối chiếu nếu có tranh chấp.")
    p(doc, "Trong tương lai, build có thể được xác minh một phần bằng dữ liệu ngoài, ảnh minh chứng hoặc xác nhận của trọng tài. Tuy nhiên, ở phiên bản hiện tại, thiết kế vẫn phải chấp nhận khả năng nhập thủ công để không phụ thuộc hoàn toàn vào dịch vụ bên thứ ba. Đây là quyết định phù hợp với phạm vi học phần và điều kiện API của trò chơi.")
    add_heading(doc, "3.6. Thiết kế luật cấm/chọn và kiểm tra phía server", 2)
    add_heading(doc, "3.6.1. Nguyên tắc kiểm tra lượt", 3)
    p(doc, "Client có thể hiển thị lượt hiện tại để hỗ trợ người dùng, nhưng server phải là nơi quyết định cuối cùng. Khi nhận request cấm/chọn, API chuyển dữ liệu cho DraftService. Service truy vấn trạng thái phòng, danh sách DraftLog đã có và thông tin người gửi, sau đó gọi DraftPolicy để xác định thao tác có hợp lệ hay không.")
    add_heading(doc, "3.6.2. Kiểm tra nhân vật và trạng thái phòng", 3)
    p(doc, "Một nhân vật đã bị cấm hoặc đã được chọn không được phép xuất hiện lại trong cùng phòng. Ngoài ra, thao tác chỉ hợp lệ khi phòng đang ở trạng thái DRAFTING. Nếu phòng đã chuyển sang BUILDING hoặc FINISHED, API phải trả lỗi rõ ràng thay vì im lặng bỏ qua thao tác.")
    add_heading(doc, "3.6.3. Ghi nhận lịch sử thao tác", 3)
    p(doc, "Mỗi thao tác hợp lệ được lưu thành DraftLog gồm roomId, player, action, characterId và turnNumber. Thiết kế này giúp hệ thống có thể dựng lại toàn bộ tiến trình cấm/chọn, hỗ trợ kiểm tra tranh chấp và khôi phục trạng thái khi người dùng tải lại trang.")
    add_heading(doc, "3.6.4. Xử lý thao tác đồng thời", 3)
    p(doc, "Trong môi trường realtime, hai client có thể gửi request gần như cùng thời điểm, đặc biệt khi người dùng mở nhiều tab hoặc kết nối mạng có độ trễ. Vì vậy, hệ thống cần dựa vào trạng thái server và ràng buộc dữ liệu để quyết định kết quả cuối cùng. Ràng buộc unique trên DraftLog theo roomId và characterId giúp giảm nguy cơ một nhân vật được ghi nhận nhiều lần trong cùng phòng.")
    p(doc, "Ngoài ràng buộc dữ liệu, service cần đọc lại danh sách DraftLog mới nhất trước khi quyết định lượt hợp lệ. Nếu chỉ tin vào trạng thái client, người dùng có thể gửi request dựa trên dữ liệu cũ. Đây là lý do các kiểm tra nghiệp vụ quan trọng được đặt ở tầng server, còn client chỉ đóng vai trò hỗ trợ hiển thị và giảm thao tác sai.")
    add_heading(doc, "3.6.5. Nguyên tắc thông báo lỗi nghiệp vụ", 3)
    p(doc, "Thông báo lỗi trong luồng cấm/chọn cần được thiết kế như một phần của nghiệp vụ, không chỉ là phản hồi kỹ thuật. Khi thao tác bị từ chối, người dùng cần biết nguyên nhân là sai lượt, sai quyền, nhân vật đã bị khóa hay phòng không còn ở trạng thái draft. Nếu chỉ trả về thông báo chung, người chơi khó hiểu tình huống và trọng tài phải giải thích lại bằng lời.")
    p(doc, "Vì vậy, API nên phân biệt lỗi theo nhóm. Lỗi dữ liệu đầu vào cho biết request thiếu thông tin hoặc sai định dạng. Lỗi phân quyền cho biết người gửi không có quyền thao tác. Lỗi nghiệp vụ cho biết thao tác không phù hợp với trạng thái hiện tại. Cách phân loại này giúp giao diện có thể hiển thị thông báo rõ hơn, đồng thời hỗ trợ người phát triển kiểm tra lỗi trong quá trình vận hành.")
    add_image(doc, diagrams["flow"], "3.4. Luồng xử lý một lượt cấm/chọn", 12.5)
    add_heading(doc, "3.7. Thiết kế Cost và Time Handicap", 2)
    add_heading(doc, "3.7.1. Nguyên tắc tính Cost", 3)
    p(doc, "Cost được thiết kế để phản ánh lợi thế tài nguyên của đội hình. Nhân vật hoặc vũ khí có độ hiếm cao, cung mệnh cao hoặc refinement cao có thể được quy đổi thành điểm Cost theo luật của ban tổ chức. Trong phiên bản hiện tại, bảng Cost được thiết kế theo hướng có thể cấu hình, tránh khóa chặt luật vào giao diện.")
    add_heading(doc, "3.7.2. Quy đổi Time Handicap", 3)
    p(doc, "Sau khi có tổng Cost của hai đội, hệ thống tính chênh lệch Cost và nhân với hệ số costPerPoint của phòng. Đội có tổng Cost cao hơn phải bù trừ bằng thời gian theo công thức: thời gian bù trừ = chênh lệch Cost x số giây mỗi Cost. Kết quả này giúp trận đấu cân bằng hơn khi tài nguyên tài khoản không tương đồng.")
    add_heading(doc, "3.8. Kiến trúc phân tầng", 2)
    p(doc, "Hệ thống được tổ chức theo kiến trúc phân tầng để tránh việc logic nghiệp vụ nằm rải rác trong component giao diện. Cách tiếp cận này giúp các chức năng quan trọng như kiểm tra lượt, phân quyền phòng, tính Cost và lưu dữ liệu có trách nhiệm rõ ràng hơn.")
    add_image(doc, diagrams["architecture"], "3.2. Kiến trúc phân tầng của website", 15.5)
    add_heading(doc, "3.8.1. Tầng Presentation", 3)
    p(doc, "Tầng Presentation bao gồm các trang trong src/app và component trong src/components. Tầng này chịu trách nhiệm hiển thị trạng thái, nhận thao tác người dùng và gọi API. Các component không nên tự quyết định tính hợp lệ của nghiệp vụ mà chỉ hỗ trợ hiển thị phản hồi từ server.")
    add_heading(doc, "3.8.2. Tầng Application và Domain", 3)
    p(doc, "Tầng Application chứa service xử lý use case như tạo phòng, thực hiện draft, lưu build và tính kết quả. Tầng Domain chứa các policy mô tả quy tắc nghiệp vụ như DraftPolicy, CostPolicy và RoomAccessPolicy. Đây là nơi đặt các quy tắc có giá trị lâu dài, ít phụ thuộc vào giao diện.")
    add_heading(doc, "3.8.3. Tầng Infrastructure và Composition", 3)
    p(doc, "Tầng Infrastructure chịu trách nhiệm tương tác Prisma repository, Supabase gateway và dữ liệu ngoài. Tầng Composition phụ trách wiring service, cấu hình môi trường và khởi tạo phụ thuộc. Việc tách riêng các tầng này giúp hệ thống có thể thay đổi cách lưu trữ hoặc tích hợp dịch vụ ngoài mà không làm xáo trộn toàn bộ giao diện.")
    add_heading(doc, "3.9. Thiết kế cơ sở dữ liệu", 2)
    add_image(doc, diagrams["data"], "3.3. Mô hình dữ liệu chính theo vòng đời trận đấu", 15.7)
    add_heading(doc, "3.9.1. Nguyên tắc thiết kế dữ liệu", 3)
    p(doc, "Thiết kế cơ sở dữ liệu của hệ thống tuân theo nguyên tắc tách dữ liệu trạng thái, dữ liệu lịch sử và dữ liệu mở rộng. Room lưu trạng thái hiện tại của phòng; DraftLog lưu lịch sử lượt; CharacterBuild lưu dữ liệu đội hình; Tournament lưu cấu hình giải đấu; TournamentMatch liên kết trận trong bracket với phòng cấm/chọn. Việc tách rõ các nhóm dữ liệu giúp hệ thống tránh lưu quá nhiều thông tin không đồng nhất trong một bảng duy nhất.")
    p(doc, "Một nguyên tắc khác là sử dụng ràng buộc dữ liệu để hỗ trợ kiểm soát nghiệp vụ. Ví dụ, unique theo roomId và characterId trong DraftLog giúp hạn chế trùng nhân vật trong cùng phòng. Unique theo roomId, player và characterId trong CharacterBuild giúp mỗi đội chỉ có một bản ghi build cho một nhân vật đã chọn. Những ràng buộc này không thay thế hoàn toàn kiểm tra service, nhưng là lớp bảo vệ quan trọng ở tầng dữ liệu.")
    add_heading(doc, "3.9.2. Thiết kế dữ liệu cho khả năng mở rộng", 3)
    p(doc, "Các bảng tournament, participant và match được thiết kế để chuẩn bị cho nhu cầu mở rộng từ trận đơn sang giải đấu. Một Tournament có nhiều TournamentParticipant và TournamentMatch; mỗi match có thể liên kết đến một roomCode khi trận cấm/chọn bắt đầu. Cách thiết kế này giúp hệ thống không cần tạo một loại phòng riêng cho giải đấu, mà tái sử dụng Room như đơn vị thi đấu cơ bản.")
    p(doc, "Ngoài dữ liệu trận đấu, hệ thống còn có các bảng phục vụ trải nghiệm cộng đồng như Friendship, Notification, ActivityEvent và UserSettings. Những bảng này không bắt buộc để chạy một phiên cấm/chọn cơ bản, nhưng tạo nền tảng cho sản phẩm nếu tiếp tục phát triển thành nền tảng tổ chức giải đấu cộng đồng. Việc đưa chúng vào thiết kế thể hiện định hướng sản phẩm dài hạn, đồng thời cho thấy dữ liệu lõi và dữ liệu mở rộng được phân tách tương đối rõ.")
    add_caption(doc, "3.5. Thiết kế chi tiết một số bảng dữ liệu chính", "Bảng")
    add_table(
        doc,
        ["Bảng", "Trường tiêu biểu", "Khóa/ràng buộc", "Vai trò"],
        [
            ["Room", "code, status, costPerPoint, blueClientId, redClientId, draftTemplate", "code unique; index hostUserId, seriesId", "Lưu trạng thái chính của một phòng cấm/chọn."],
            ["DraftLog", "roomId, player, action, characterId, turnNumber", "unique roomId-characterId; index roomId-turnNumber", "Ghi nhận từng lượt cấm/chọn để khôi phục lịch sử."],
            ["CharacterBuild", "roomId, player, characterId, rarity, consLevel, weaponRarity, totalCost", "unique roomId-player-characterId", "Lưu dữ liệu build và Cost của nhân vật đã chọn."],
            ["ChatMessage", "roomId, sender, message, role, createdAt", "index roomId-createdAt", "Hỗ trợ trao đổi trong phòng đấu."],
            ["Tournament", "slug, name, format, status, costCap, bankTime", "slug unique; index status", "Lưu thông tin giải đấu và cấu hình thi đấu."],
            ["TournamentMatch", "tournamentId, round, matchNumber, roomCode, winnerParticipantId", "unique tournamentId-round-matchNumber", "Liên kết trận đấu trong bracket với phòng cấm/chọn."],
            ["Notification", "recipientUid, type, title, link, read", "index recipientUid-read-createdAt", "Gửi thông báo lời mời, kết quả hoặc sự kiện cộng đồng."],
        ],
        widths=[3.0, 5.3, 4.0, 3.4],
        font_size=9.8,
    )
    add_heading(doc, "3.10. Thiết kế API theo nhóm nghiệp vụ", 2)
    add_heading(doc, "3.10.1. Nguyên tắc thiết kế API", 3)
    p(doc, "API của hệ thống được thiết kế theo hướng phục vụ use case, không chỉ phản ánh trực tiếp tên bảng dữ liệu. Ví dụ, API draft không đơn thuần tạo một DraftLog, mà xử lý toàn bộ nghiệp vụ của một lượt cấm/chọn. Cách tiếp cận này giúp API phù hợp hơn với hành vi người dùng và giảm nguy cơ client phải tự ghép nhiều thao tác nhỏ để hoàn thành một nghiệp vụ.")
    p(doc, "Mỗi API cần có ba lớp kiểm tra: kiểm tra dữ liệu đầu vào, kiểm tra quyền thao tác và kiểm tra quy tắc nghiệp vụ. Dữ liệu đầu vào sai cần bị từ chối sớm; quyền thao tác sai cần trả lỗi bảo mật; quy tắc nghiệp vụ sai cần trả lỗi có ngữ cảnh, chẳng hạn sai lượt hoặc nhân vật đã bị khóa. Việc phân biệt các nhóm lỗi giúp giao diện phản hồi chính xác hơn cho người dùng.")
    add_heading(doc, "3.10.2. Thiết kế phản hồi API", 3)
    p(doc, "Phản hồi API cần đủ thông tin để giao diện cập nhật trạng thái mà không phải tự suy diễn quá nhiều. Với thao tác draft thành công, phản hồi nên bao gồm DraftLog mới, lượt tiếp theo và trạng thái phòng sau thao tác. Với thao tác build thành công, phản hồi nên bao gồm build đã lưu và kết quả Cost đã tính lại. Cách trả dữ liệu này giúp client cập nhật giao diện theo nguồn dữ liệu chính thức từ server.")
    p(doc, "Đối với lỗi, phản hồi nên có mã lỗi nghiệp vụ và thông điệp thân thiện. Ví dụ, mã lỗi TURN_MISMATCH có thể tương ứng với thông báo 'Chưa đến lượt thao tác của đội này', còn CHARACTER_LOCKED tương ứng với 'Nhân vật đã bị cấm hoặc đã được chọn'. Cấu trúc này giúp giao diện dễ ánh xạ lỗi thành thông báo, đồng thời giúp việc ghi log và kiểm thử trở nên rõ ràng hơn.")
    add_caption(doc, "3.6. Nhóm API tiêu biểu của hệ thống", "Bảng")
    add_table(
        doc,
        ["Nhóm API", "Chức năng", "Điều kiện/ý nghĩa"],
        [
            ["Room", "Tạo phòng, đọc trạng thái, chuyển phase, cập nhật cấu hình.", "Chỉ Host hoặc quyền phù hợp được phép điều khiển trạng thái phòng."],
            ["Draft", "Nhận request cấm/chọn và ghi DraftLog.", "Bắt buộc kiểm tra lượt, quyền thao tác và trạng thái nhân vật phía server."],
            ["Build", "Lưu build, tính Cost và cập nhật kết quả.", "Chỉ nhận build cho nhân vật đã được đội đó pick."],
            ["Tournament", "Tạo giải, thêm participant, sinh match và liên kết room.", "Phục vụ mở rộng từ trận đơn sang giải đấu."],
            ["Profile/Social", "Quản lý hồ sơ, bạn bè, thông báo và hoạt động.", "Tăng khả năng vận hành như một nền tảng cộng đồng."],
        ],
        widths=[3.4, 6.0, 6.3],
    )
    add_heading(doc, "3.11. Thiết kế giao diện và trải nghiệm người dùng", 2)
    add_heading(doc, "3.11.1. Nguyên tắc ưu tiên thông tin", 3)
    p(doc, "Giao diện phòng đấu cần ưu tiên những thông tin ảnh hưởng trực tiếp đến thao tác: đội đang đến lượt, hành động cần thực hiện, nhân vật đã bị cấm, nhân vật đã được chọn và danh sách nhân vật còn khả dụng. Các thông tin phụ như chat, lịch sử hoặc hướng dẫn cần được đặt ở vị trí không làm che khuất thao tác chính.")
    add_heading(doc, "3.11.2. Bố cục Blue/Pool/Red", 3)
    p(doc, "Bố cục desktop sử dụng ba cột Blue/Pool/Red nhằm mô phỏng cách quan sát quen thuộc trong các công cụ thi đấu. Pool nhân vật ở giữa giúp người chơi nhìn ngay nguồn lựa chọn còn lại; hai đội ở hai bên giúp so sánh trạng thái đội hình. Trên màn hình nhỏ, bố cục được xếp lại theo thứ tự ưu tiên để tránh làm các thẻ nhân vật quá nhỏ.")
    add_image(doc, diagrams["layout"], "3.5. Bố cục MOBA layout ba cột Blue/Pool/Red", 15.2)
    add_heading(doc, "3.12. Design System và Style Guide", 2)
    add_caption(doc, "3.7. Bảng màu chủ đạo của giao diện", "Bảng")
    color_table = add_table(
        doc,
        ["Vai trò màu", "Mã màu gợi ý", "Mẫu màu", "Cách sử dụng trong website"],
        [
            ["Nền chính", "#0b0c10 / #101524", "", "Tạo nền tối, giảm chói khi người dùng theo dõi trận đấu trong thời gian dài."],
            ["Primary accent", "#66FCF1", "", "Nút hành động chính, viền highlight lượt hiện tại và trạng thái cần chú ý."],
            ["Vàng kim", "#F7C948", "", "Nhấn mạnh kết quả, Cost, Time Handicap và thông tin quan trọng."],
            ["Đội Xanh", "#3B82F6", "", "Khu vực Blue, trạng thái đến lượt Blue và màu nhận diện đội."],
            ["Đội Đỏ", "#EF4444", "", "Khu vực Red, trạng thái đến lượt Red và màu nhận diện đội."],
            ["Tím 4 sao", "#A855F7", "", "Nền hoặc viền nhân vật/vũ khí 4 sao."],
            ["Vàng 5 sao", "#F59E0B", "", "Nền hoặc viền nhân vật/vũ khí 5 sao."],
        ],
        widths=[3.0, 3.5, 2.3, 6.8],
        font_size=10.5,
    )
    colors = ["0B0C10", "66FCF1", "F7C948", "3B82F6", "EF4444", "A855F7", "F59E0B"]
    for row, fill in zip(color_table.rows[1:], colors):
        set_cell_shading(row.cells[2], fill)
        row.cells[2].text = ""
    add_caption(doc, "3.8. Quy ước trạng thái và phản hồi tương tác", "Bảng")
    add_table(
        doc,
        ["Trạng thái", "Biểu hiện giao diện", "Ý nghĩa UX"],
        [
            ["Đến lượt", "Viền sáng, thông báo ngắn, nút thao tác được bật.", "Người chơi biết rõ khi nào có thể thao tác."],
            ["Đã cấm", "Thẻ nhân vật bị làm mờ và có nhãn BAN.", "Tránh nhầm lẫn với nhân vật còn khả dụng."],
            ["Đã chọn", "Thẻ được gắn vào khu vực đội và khóa khỏi pool.", "Thể hiện quyền sở hữu nhân vật sau lượt pick."],
            ["Lỗi thao tác", "Thông báo rõ nguyên nhân như sai lượt, hết thời gian hoặc nhân vật không hợp lệ.", "Giúp người dùng hiểu vì sao thao tác bị từ chối."],
            ["Hoàn tất", "Trang kết quả hiển thị tổng Cost, chênh lệch và Time Handicap.", "Tạo cơ sở đối chiếu minh bạch trước khi vào trận."],
        ],
        widths=[3.0, 6.0, 6.7],
    )
    add_heading(doc, "3.13. Phân quyền và an toàn dữ liệu", 2)
    add_heading(doc, "3.13.1. Kiểm soát quyền theo vai trò", 3)
    p(doc, "Host/Referee được quyền điều khiển trạng thái phòng; Player chỉ được thao tác với lượt và build của đội của họ; Spectator/Caster chỉ được theo dõi; Admin có quyền quản lý dữ liệu hệ thống. Phân quyền này cần được kiểm tra ở server vì giao diện client có thể bị thay đổi bởi người dùng.")
    add_heading(doc, "3.13.2. Bảo vệ trạng thái chính thức", 3)
    p(doc, "Trạng thái chính thức của trận đấu được xác định bởi dữ liệu trong cơ sở dữ liệu, không phải bởi state tạm thời ở trình duyệt. Khi có xung đột giữa dữ liệu client và server, hệ thống ưu tiên trạng thái server và trả về thông báo lỗi rõ ràng cho thao tác không hợp lệ.")
    add_heading(doc, "3.13.3. Bảo vệ dữ liệu cấu hình", 3)
    p(doc, "Các cấu hình như costPerPoint, draftTemplate, spectatorDelay hoặc tournament rules có ảnh hưởng trực tiếp đến kết quả trận đấu. Vì vậy, hệ thống không cho phép mọi người dùng tự do sửa các trường này. Việc cập nhật cấu hình cần gắn với quyền Host hoặc Admin và nên được giới hạn ở những thời điểm phù hợp, chẳng hạn trước khi bắt đầu draft hoặc khi trọng tài tạm dừng trận.")
    p(doc, "Đối với dữ liệu nhạy cảm như webhook, khóa dịch vụ hoặc thông tin kết nối cơ sở dữ liệu, hệ thống phải sử dụng biến môi trường và không đưa trực tiếp vào client. Đây là nguyên tắc cơ bản nhưng quan trọng vì ứng dụng web có nhiều mã chạy ở trình duyệt, nơi người dùng có thể kiểm tra gói JavaScript đã tải xuống.")
    add_heading(doc, "3.14. Khả năng mở rộng", 2)
    p(doc, "Thiết kế dữ liệu và mã nguồn cho phép mở rộng theo hướng tournament, social và overlay mà không phá vỡ luồng cấm/chọn cốt lõi. Các bảng Tournament, TournamentParticipant và TournamentMatch giúp liên kết nhiều phòng đấu thành một giải. Các bảng Notification, ActivityEvent và Friendship tạo nền tảng cho trải nghiệm cộng đồng.")
    add_heading(doc, "3.14.1. Mở rộng theo chiều chức năng", 3)
    p(doc, "Theo chiều chức năng, hệ thống có thể bổ sung các module như overlay cho livestream, thống kê tỉ lệ chọn nhân vật, lịch sử đối đầu, hệ thống nhiệm vụ cộng đồng hoặc bảng xếp hạng. Các module này không bắt buộc trong luồng cấm/chọn cơ bản nhưng có thể khai thác dữ liệu đã được lưu trong Room, DraftLog, CharacterBuild và TournamentMatch.")
    add_heading(doc, "3.14.2. Mở rộng theo chiều vận hành", 3)
    p(doc, "Theo chiều vận hành, hệ thống cần được bổ sung giám sát log, kiểm thử tải realtime, sao lưu cơ sở dữ liệu và cơ chế phục hồi khi dịch vụ ngoài gặp lỗi. Khi số lượng phòng đồng thời tăng, vấn đề không chỉ là giao diện có hoạt động hay không, mà còn là độ trễ cập nhật, số kết nối realtime và khả năng duy trì trạng thái nhất quán.")
    add_heading(doc, "3.15. Rủi ro thiết kế và biện pháp kiểm soát", 2)
    add_heading(doc, "3.15.1. Rủi ro phụ thuộc vào client", 3)
    p(doc, "Nếu quá nhiều logic nghiệp vụ được đặt ở client, người dùng có thể thao tác vượt quyền bằng cách sửa request hoặc can thiệp vào mã chạy trên trình duyệt. Biện pháp kiểm soát là đặt kiểm tra bắt buộc ở server, trong đó client chỉ đóng vai trò hướng dẫn người dùng và hiển thị trạng thái hiện tại.")
    add_heading(doc, "3.15.2. Rủi ro dữ liệu realtime không đồng nhất", 3)
    p(doc, "Realtime giúp giao diện cập nhật nhanh nhưng không nên được xem là nguồn dữ liệu chính thức duy nhất. Khi client mất kết nối hoặc nhận sự kiện trễ, giao diện có thể hiển thị trạng thái cũ. Biện pháp kiểm soát là cho phép truy vấn lại trạng thái phòng từ database và thiết kế API trả về trạng thái mới sau mỗi thao tác thành công.")
    add_heading(doc, "3.15.3. Rủi ro mở rộng quá sớm", 3)
    p(doc, "Hệ thống có nhiều hướng mở rộng như social, tournament, shop hoặc leaderboard. Tuy nhiên, nếu mở rộng quá sớm khi luồng cấm/chọn cốt lõi chưa ổn định, mã nguồn dễ bị phân tán. Vì vậy, thiết kế ưu tiên giữ vững vòng đời trận đấu trước, sau đó mới mở rộng các chức năng cộng đồng dựa trên dữ liệu đã ổn định.")
    add_heading(doc, "3.16. Truy vết yêu cầu đến thiết kế", 2)
    add_heading(doc, "3.16.1. Liên hệ giữa yêu cầu chức năng và dữ liệu", 3)
    p(doc, "Mỗi yêu cầu chức năng chính đều được ánh xạ đến một hoặc nhiều bảng dữ liệu. Yêu cầu tạo phòng tương ứng với Room; yêu cầu cấm/chọn tương ứng với DraftLog; yêu cầu khai báo build tương ứng với CharacterBuild; yêu cầu theo dõi realtime sử dụng Room kết hợp DraftLog và các kênh đồng bộ; yêu cầu giải đấu mở rộng tương ứng với Tournament, TournamentParticipant và TournamentMatch. Việc truy vết này giúp chứng minh thiết kế dữ liệu không được tạo ra một cách tùy ý, mà xuất phát từ yêu cầu nghiệp vụ.")
    p(doc, "Cách truy vết yêu cầu cũng hỗ trợ kiểm thử. Khi một yêu cầu bị lỗi, người phát triển có thể xác định cần kiểm tra bảng nào, API nào và service nào. Ví dụ, nếu trang kết quả hiển thị sai Time Handicap, cần kiểm tra CharacterBuild, CostPolicy và Result page; nếu lượt cấm/chọn sai, cần kiểm tra DraftLog, DraftPolicy và API Draft. Điều này làm giảm thời gian khoanh vùng lỗi.")
    add_heading(doc, "3.16.2. Liên hệ giữa yêu cầu phi chức năng và kiến trúc", 3)
    p(doc, "Yêu cầu phi chức năng cũng được phản ánh trong kiến trúc hệ thống. Tính đúng đắn được hỗ trợ bằng policy phía server và ràng buộc dữ liệu. Tính minh bạch được hỗ trợ bằng DraftLog và trang kết quả. Tính realtime được hỗ trợ bằng Supabase Realtime kết hợp truy vấn lại trạng thái chính thức. Tính bảo trì được hỗ trợ bằng phân tầng và component hóa. Như vậy, kiến trúc không chỉ là sơ đồ minh họa, mà là cách trả lời các yêu cầu chất lượng của hệ thống.")
    p(doc, "Trong báo cáo kỹ thuật, việc trình bày mối liên hệ giữa yêu cầu và thiết kế có ý nghĩa quan trọng. Nếu chỉ liệt kê công nghệ, người đọc khó đánh giá vì sao hệ thống cần các thành phần đó. Khi yêu cầu được liên hệ với kiến trúc, báo cáo cho thấy quyết định kỹ thuật có căn cứ và phù hợp với bài toán.")
    add_heading(doc, "3.17. Đánh giá tính nhất quán của thiết kế", 2)
    add_heading(doc, "3.17.1. Nhất quán giữa vòng đời phòng và giao diện", 3)
    p(doc, "Giao diện của hệ thống cần phản ánh đúng vòng đời phòng. Ở trạng thái WAITING, giao diện tập trung vào việc tham gia đội và chuẩn bị. Ở trạng thái DRAFTING, giao diện ưu tiên pool nhân vật, lượt hiện tại và lịch sử cấm/chọn. Ở trạng thái BUILDING, giao diện ưu tiên danh sách nhân vật đã pick và biểu mẫu khai báo build. Ở trạng thái FINISHED, giao diện ưu tiên kết quả Cost và Time Handicap. Nếu giao diện không thay đổi theo trạng thái, người dùng dễ thao tác sai hoặc hiểu nhầm quy trình.")
    p(doc, "Sự nhất quán này cần được duy trì ở cả desktop và mobile. Trên desktop, bố cục ba cột giúp quan sát tổng thể; trên mobile, cần sắp xếp theo thứ tự ưu tiên để thao tác không bị rối. Dù bố cục thay đổi theo kích thước màn hình, thông tin cốt lõi như lượt hiện tại, trạng thái nhân vật và kết quả vẫn phải được giữ rõ ràng.")
    add_heading(doc, "3.17.2. Nhất quán giữa dữ liệu chính thức và dữ liệu hiển thị", 3)
    p(doc, "Dữ liệu hiển thị trên giao diện phải được đồng bộ từ dữ liệu chính thức. Khi người chơi thực hiện một lượt, giao diện có thể phản hồi nhanh bằng trạng thái tạm thời, nhưng kết quả cuối cùng vẫn phải dựa trên phản hồi từ server. Nguyên tắc này giúp hạn chế trường hợp client hiển thị đã chọn thành công trong khi server từ chối vì sai lượt hoặc nhân vật đã bị khóa.")
    p(doc, "Để đạt được sự nhất quán, hệ thống cần kết hợp ba yếu tố: service kiểm tra nghiệp vụ, database lưu trạng thái chính thức và realtime phát tín hiệu cập nhật. Nếu thiếu một trong ba yếu tố này, trải nghiệm có thể bị lệch. Ví dụ, có realtime nhưng không có lưu trữ bền vững sẽ khó khôi phục sau khi tải lại trang; có database nhưng không có realtime sẽ khiến các vai trò phải tải lại thủ công; có giao diện đẹp nhưng thiếu service kiểm tra sẽ không bảo vệ được tính công bằng.")


def chapter_4(doc):
    add_heading(doc, "CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG", 1, page_break=True)
    add_heading(doc, "4.1. Tổ chức mã nguồn", 2)
    add_heading(doc, "4.1.1. Cấu trúc theo thư mục chức năng", 3)
    p(doc, "Mã nguồn được tổ chức theo cách kết hợp route của Next.js với các thư mục component, service và dữ liệu. Các route trong src/app phản ánh các màn hình hoặc nhóm API, trong khi src/components chứa các thành phần giao diện dùng lại. Cách tổ chức này giúp người phát triển xác định nhanh vị trí liên quan khi cần sửa một chức năng.")
    add_heading(doc, "4.1.2. Tách logic nghiệp vụ khỏi giao diện", 3)
    p(doc, "Các chức năng quan trọng không nên được xử lý hoàn toàn trong component. Component chỉ nên hiển thị dữ liệu, thu nhận thao tác và gọi API. Logic như kiểm tra lượt, tính Cost hoặc kiểm tra quyền được đặt trong service/policy để giảm rủi ro sai lệch giữa các màn hình khác nhau.")
    p(doc, "Trong thực tế phát triển, một số xử lý hiển thị như lọc nhân vật, tô màu trạng thái hoặc mở modal có thể đặt ở component. Tuy nhiên, các quyết định ảnh hưởng đến dữ liệu chính thức như chấp nhận lượt pick, chuyển phase hoặc tính kết quả phải nằm ở tầng server. Ranh giới này giúp mã nguồn dễ kiểm tra hơn và hạn chế lỗi khi giao diện được thay đổi.")
    add_caption(doc, "4.1. Tổ chức mã nguồn theo tầng", "Bảng")
    add_table(
        doc,
        ["Tầng", "Thư mục/file tiêu biểu", "Vai trò"],
        [
            ["Presentation", "src/app, src/components", "Hiển thị giao diện, nhận thao tác và gọi API."],
            ["Application", "services xử lý use case", "Điều phối các bước nghiệp vụ như tạo phòng, draft, build và result."],
            ["Domain", "DraftPolicy, CostPolicy, RoomAccessPolicy", "Định nghĩa quy tắc nghiệp vụ cốt lõi, ít phụ thuộc giao diện."],
            ["Infrastructure", "Prisma repositories, Supabase gateway", "Truy cập dữ liệu, đồng bộ realtime và tích hợp dịch vụ ngoài."],
            ["Composition", "Cấu hình app, wiring service", "Khởi tạo phụ thuộc, cấu hình môi trường và kết nối các tầng."],
        ],
        widths=[3.2, 5.3, 7.2],
    )
    add_heading(doc, "4.2. Xây dựng nền tảng phòng đấu", 2)
    add_heading(doc, "4.2.1. Tạo phòng và mã truy cập", 3)
    p(doc, "Khi Host tạo phòng, hệ thống sinh mã phòng duy nhất, lưu cấu hình mặc định và đặt trạng thái WAITING. Mã phòng là thông tin ngắn gọn để người chơi tham gia, đồng thời là khóa nghiệp vụ giúp các màn hình liên quan truy vấn trạng thái phòng.")
    add_heading(doc, "4.2.2. Quản lý vai trò trong phòng", 3)
    p(doc, "Phòng lưu thông tin host, clientId của đội xanh, clientId của đội đỏ và danh sách caster. Dữ liệu này được dùng để kiểm tra quyền thao tác. Ví dụ, người chơi đội xanh không được gửi lượt cho đội đỏ, spectator không được chuyển phase và người không thuộc phòng không được sửa build.")
    add_heading(doc, "4.2.3. Trạng thái vòng đời phòng", 3)
    p(doc, "Vòng đời phòng được mô tả bằng các trạng thái như WAITING, DRAFTING, BUILDING và FINISHED. Trạng thái WAITING dùng khi phòng vừa được tạo và người chơi chưa sẵn sàng. Trạng thái DRAFTING dùng cho giai đoạn cấm/chọn. Trạng thái BUILDING dùng khi người chơi khai báo đội hình và trang bị. Trạng thái FINISHED dùng khi kết quả đã được tổng hợp.")
    p(doc, "Việc mô hình hóa trạng thái giúp API kiểm soát thao tác tốt hơn. Ví dụ, request cấm/chọn chỉ hợp lệ ở trạng thái DRAFTING, còn request cập nhật build chỉ hợp lệ ở trạng thái BUILDING. Nếu không có trạng thái rõ ràng, hệ thống dễ chấp nhận thao tác ngoài ngữ cảnh, dẫn đến dữ liệu khó phục hồi hoặc kết quả không nhất quán.")
    add_heading(doc, "4.3. Xây dựng draft engine realtime", 2)
    add_heading(doc, "4.3.1. Xử lý request cấm/chọn", 3)
    p(doc, "Khi người chơi chọn nhân vật, client gửi request gồm roomCode, team, action và characterId. API Draft không ghi dữ liệu ngay mà chuyển qua service để kiểm tra trạng thái phòng, quyền thao tác, lượt hiện tại và tính hợp lệ của nhân vật. Chỉ khi tất cả điều kiện đều hợp lệ, DraftLog mới được tạo.")
    add_heading(doc, "4.3.2. Đồng bộ giao diện sau thao tác", 3)
    p(doc, "Sau khi DraftLog được lưu, các client đang theo dõi phòng nhận cập nhật realtime và truy vấn lại trạng thái cần thiết. Cách làm này giúp giao diện của trọng tài, người chơi và khán giả đồng nhất hơn so với việc mỗi client tự suy diễn trạng thái từ thao tác cục bộ.")
    add_heading(doc, "4.3.3. Phản hồi lỗi trong draft engine", 3)
    p(doc, "Draft engine cần trả lỗi có ý nghĩa đối với người dùng. Thay vì chỉ trả lỗi chung, hệ thống nên phân biệt các trường hợp như phòng không tồn tại, sai trạng thái, sai lượt, nhân vật đã bị khóa hoặc người dùng không có quyền. Thông điệp lỗi rõ ràng giúp người chơi hiểu vấn đề, đồng thời giúp trọng tài xử lý nhanh khi có sự cố.")
    p(doc, "Phản hồi lỗi cũng hỗ trợ kiểm thử. Khi test case thất bại, thông điệp cụ thể giúp xác định lỗi thuộc về dữ liệu đầu vào, phân quyền hay quy tắc nghiệp vụ. Điều này làm quá trình sửa lỗi nhanh hơn so với việc toàn bộ trường hợp đều trả về một thông báo chung.")
    add_heading(doc, "4.4. Xây dựng khai báo build và tính Cost", 2)
    add_heading(doc, "4.4.1. Dữ liệu build của nhân vật", 3)
    p(doc, "Mỗi bản ghi CharacterBuild gắn với roomId, player và characterId. Các trường rarity, consLevel, weaponRarity, source và enkaSnapshot cho phép hệ thống vừa hỗ trợ nhập liệu thủ công, vừa có không gian lưu dữ liệu lấy từ nguồn ngoài nếu cần.")
    add_heading(doc, "4.4.2. Tính toán kết quả", 3)
    p(doc, "CostService tổng hợp Cost của từng đội, tính chênh lệch và quy đổi Time Handicap. Kết quả cần được tính từ dữ liệu đã lưu để đảm bảo người dùng tải lại trang hoặc mở trên thiết bị khác vẫn nhận được kết quả thống nhất.")
    add_heading(doc, "4.4.3. Khả năng cấu hình luật Cost", 3)
    p(doc, "Luật Cost trong cộng đồng có thể thay đổi theo từng mùa giải hoặc từng nhóm tổ chức. Vì vậy, thiết kế không nên cố định hoàn toàn mọi giá trị trong giao diện. Các thông tin như costPerPoint, costCap, bankTime hoặc quy tắc fearless draft cần được lưu ở cấu hình phòng hoặc giải đấu. Cách tiếp cận này giúp hệ thống linh hoạt hơn khi ban tổ chức thay đổi thể thức.")
    p(doc, "Dù vậy, cấu hình luật cần được kiểm soát quyền và thời điểm chỉnh sửa. Nếu luật thay đổi sau khi trận đã bắt đầu mà không có cơ chế ghi nhận, kết quả có thể gây tranh luận. Do đó, khi phát triển tiếp, hệ thống nên bổ sung lịch sử thay đổi cấu hình để trọng tài và người chơi có thể đối chiếu.")
    add_heading(doc, "4.5. Xây dựng trải nghiệm trận đấu", 2)
    add_heading(doc, "4.5.1. Giao diện chọn vai trò", 3)
    p(doc, "Giao diện chọn vai trò giúp người dùng vào đúng luồng sử dụng. Host có thể tạo hoặc điều phối phòng; player tham gia theo đội; spectator/caster theo dõi trạng thái. Việc phân biệt vai trò ngay từ đầu giúp giảm nhầm lẫn trong quá trình thao tác.")
    add_image(doc, MEDIA_DIR / "image8.png", "4.1. Giao diện chọn vai trò và bắt đầu phiên sử dụng", 15.5)
    add_heading(doc, "4.5.2. Giao diện phòng cấm/chọn", 3)
    p(doc, "Giao diện phòng cấm/chọn đặt pool nhân vật ở trung tâm và hai đội ở hai bên. Trạng thái lượt được thể hiện bằng màu sắc, timer và vùng thao tác. Các nhân vật đã bị cấm hoặc đã chọn được khóa trạng thái để người dùng không thao tác nhầm.")
    add_image(doc, MEDIA_DIR / "image7.png", "4.2. Giao diện phòng cấm/chọn trong trạng thái thi đấu", 15.5)
    add_image(doc, MEDIA_DIR / "image10.png", "4.3. Giao diện tổng hợp sau lượt cấm/chọn", 15.5)
    add_heading(doc, "4.5.3. Giao diện kết quả", 3)
    p(doc, "Trang kết quả hiển thị tổng Cost của hai đội, chênh lệch và Time Handicap. Thông tin được bố trí đối xứng để trọng tài và người chơi dễ đối chiếu. Đây là màn hình quan trọng vì kết quả trên trang này là cơ sở trước khi hai bên bước vào phần thi đấu trong game.")
    add_image(doc, MEDIA_DIR / "image11.png", "4.4. Giao diện kết quả và Time Handicap", 15.5)
    add_heading(doc, "4.6. Mở rộng cộng đồng và giải đấu", 2)
    add_heading(doc, "4.6.1. Tournament và participant", 3)
    p(doc, "Phân hệ tournament giúp mở rộng hệ thống từ một trận đơn sang giải đấu. Tournament lưu thông tin format, status, maxTeams, costCap, bankTime và rulesText. Participant lưu người tham gia, teamName, logoUrl, seed và danh sách thành viên. Đây là nền tảng để sinh bracket và liên kết từng match với một room cụ thể.")
    add_image(doc, MEDIA_DIR / "image13.png", "4.5. Giao diện giải đấu và danh sách người tham gia", 15.5)
    add_heading(doc, "4.6.2. Hồ sơ và lịch sử hoạt động", 3)
    p(doc, "Các bảng ActivityEvent, Notification, UserSettings và Friendship giúp hệ thống có khả năng vận hành như một nền tảng cộng đồng. Người chơi có thể xem lịch sử hoạt động, thông báo, bạn bè hoặc cấu hình cá nhân. Dù đây không phải chức năng lõi của cấm/chọn, các thành phần này giúp sản phẩm có hướng phát triển dài hạn hơn.")
    add_image(doc, MEDIA_DIR / "image14.png", "4.6. Giao diện hồ sơ, thống kê và lịch sử hoạt động", 15.5)
    add_heading(doc, "4.6.3. Bảng xếp hạng và dữ liệu công khai", 3)
    p(doc, "Bên cạnh phòng đấu và giải đấu, hệ thống có thể hiển thị bảng xếp hạng hoặc danh sách người chơi theo dữ liệu công khai. Chức năng này giúp tăng tính gắn kết cộng đồng nhưng cần được thiết kế cẩn trọng về quyền riêng tư. Người chơi cần có tùy chọn hiển thị hoặc ẩn một số thông tin cá nhân, đặc biệt khi dữ liệu liên quan đến UID, hồ sơ hoặc lịch sử hoạt động.")
    p(doc, "Khi triển khai bảng xếp hạng, hệ thống nên ưu tiên dữ liệu có thể kiểm chứng như số trận hoàn thành, kết quả giải đấu hoặc thành tích đã được xác nhận, thay vì chỉ dựa vào dữ liệu tự khai báo. Cách làm này giúp bảng xếp hạng có giá trị tham khảo cao hơn và giảm khả năng phát sinh tranh chấp.")
    add_heading(doc, "4.7. Triển khai cơ sở dữ liệu và ứng dụng", 2)
    add_heading(doc, "4.7.1. Cấu hình biến môi trường", 3)
    p(doc, "Ứng dụng sử dụng các biến môi trường để cấu hình DATABASE_URL, DIRECT_URL, Supabase URL, Supabase key và các khóa tích hợp ngoài. Việc tách cấu hình khỏi mã nguồn giúp bảo vệ thông tin nhạy cảm và cho phép triển khai linh hoạt trên nhiều môi trường.")
    add_heading(doc, "4.7.2. Prisma migration và generate client", 3)
    p(doc, "Sau khi cập nhật Prisma Schema, người phát triển cần đồng bộ cấu trúc dữ liệu với PostgreSQL và sinh Prisma Client. Lệnh generate giúp mã nguồn TypeScript nhận biết model và field mới, giảm lỗi khi truy vấn dữ liệu.")
    add_code_block(doc, "npm run build\nnpx prisma generate\nnpx prisma db push")
    add_heading(doc, "4.7.3. Docker và vận hành", 3)
    p(doc, "Dockerfile multi-stage giúp tối ưu bản build production. Docker Compose hỗ trợ chạy ứng dụng cùng cấu hình cần thiết trong môi trường triển khai. Khi triển khai thực tế, cần kiểm tra HTTPS, domain, biến môi trường, kết nối database và log ứng dụng.")
    add_heading(doc, "4.7.4. Kiểm soát lỗi triển khai", 3)
    p(doc, "Một số lỗi triển khai thường gặp gồm thiếu biến môi trường, Prisma Client chưa được generate, kết nối database sai chuỗi, service realtime chưa bật hoặc domain chưa cấu hình HTTPS. Để giảm rủi ro, quá trình triển khai cần có checklist rõ ràng và ghi nhận log sau khi build.")
    p(doc, "Đối với ứng dụng có realtime, sau khi triển khai không chỉ kiểm tra trang có mở được hay không, mà còn cần mở nhiều client để xác nhận sự kiện cập nhật được truyền đúng. Đây là bước kiểm tra quan trọng vì lỗi realtime có thể không xuất hiện khi chỉ dùng một trình duyệt.")
    add_heading(doc, "4.8. Kiểm thử chức năng", 2)
    p(doc, "Kiểm thử được thiết kế theo các kịch bản đại diện cho luồng người dùng, thay vì chỉ kiểm tra từng màn hình riêng lẻ. Mục tiêu là chứng minh hệ thống có thể duy trì trạng thái đúng qua nhiều bước thao tác liên tiếp.")
    add_caption(doc, "4.2. Kết quả kiểm thử chức năng", "Bảng")
    add_table(
        doc,
        ["Mã", "Kịch bản", "Dữ liệu thử", "Mong đợi", "Đánh giá"],
        [
            ["TC-01", "Host tạo phòng mới", "Tên host, cấu hình mặc định", "Phòng có mã duy nhất, trạng thái WAITING.", "Đạt"],
            ["TC-02", "Hai player tham gia phòng", "Client Blue và Red", "Hai đội được gán đúng, spectator không chiếm vị trí player.", "Đạt"],
            ["TC-03", "Player thao tác sai lượt", "Red gửi request khi đến lượt Blue", "API từ chối và trả thông báo lỗi rõ ràng.", "Đạt"],
            ["TC-04", "Chọn lại nhân vật đã cấm/chọn", "characterId đã tồn tại trong DraftLog", "API từ chối do vi phạm ràng buộc.", "Đạt"],
            ["TC-05", "Khai báo build sau draft", "Danh sách nhân vật đã pick", "Build được lưu và tổng Cost cập nhật.", "Đạt"],
            ["TC-06", "Tải lại trang trong phòng", "Room đang DRAFTING/BUILDING", "Giao diện khôi phục trạng thái từ database.", "Đạt"],
            ["TC-07", "Xem kết quả", "Cost Blue và Red khác nhau", "Hiển thị chênh lệch và Time Handicap đúng công thức.", "Đạt"],
        ],
        widths=[1.4, 4.0, 3.8, 5.0, 1.5],
        font_size=9.8,
    )
    add_heading(doc, "4.9. Đánh giá quá trình hiện thực", 2)
    add_heading(doc, "4.9.1. Điểm mạnh khi hiện thực", 3)
    p(doc, "Điểm mạnh của quá trình hiện thực là hệ thống đã chuyển được bài toán cộng đồng thành mô hình dữ liệu và luồng thao tác rõ ràng. Các thực thể chính như Room, DraftLog, CharacterBuild và Tournament phản ánh đúng vòng đời trận đấu. Giao diện cũng bám sát ngữ cảnh thi đấu với bố cục phân đội, trạng thái lượt và kết quả trực quan.")
    add_heading(doc, "4.9.2. Những điểm cần tiếp tục hoàn thiện", 3)
    p(doc, "Một số điểm cần tiếp tục hoàn thiện gồm kiểm thử tải realtime, hoàn thiện xác minh dữ liệu build, tối ưu quản lý lỗi khi nhiều client thao tác gần đồng thời và bổ sung tài liệu vận hành. Các điểm này không làm thay đổi kiến trúc cốt lõi nhưng cần được xử lý trước khi hệ thống vận hành ở quy mô giải đấu lớn.")
    add_heading(doc, "4.10. Đánh giá chất lượng mã nguồn", 2)
    add_heading(doc, "4.10.1. Tính dễ đọc và khả năng bảo trì", 3)
    p(doc, "Chất lượng mã nguồn không chỉ được đánh giá bằng việc chức năng chạy được, mà còn bằng khả năng đọc hiểu, sửa lỗi và mở rộng. Với đề tài này, việc đặt tên thư mục theo nhóm chức năng và tách component theo vai trò giao diện giúp người phát triển mới có thể định vị nhanh phần cần chỉnh sửa. Đây là yếu tố quan trọng vì website có nhiều màn hình, từ phòng đấu đến tournament và social.")
    p(doc, "Để tiếp tục nâng cao khả năng bảo trì, hệ thống nên chuẩn hóa thêm quy ước đặt tên service, DTO và repository. Khi các use case tăng lên, việc thiếu quy ước có thể làm mã nguồn khó theo dõi. Một lớp tài liệu kỹ thuật ngắn mô tả luồng request, nơi đặt logic và nơi truy cập dữ liệu sẽ giúp quá trình phát triển nhóm hiệu quả hơn.")
    add_heading(doc, "4.10.2. Tính nhất quán giữa giao diện và nghiệp vụ", 3)
    p(doc, "Một rủi ro thường gặp ở ứng dụng web nhiều trạng thái là giao diện hiển thị một trạng thái trong khi server lại lưu trạng thái khác. Để giảm rủi ro này, giao diện cần được thiết kế theo hướng nhận dữ liệu chuẩn hóa từ server, thay vì tự xây dựng quá nhiều trạng thái suy diễn. Sau mỗi thao tác quan trọng, client nên cập nhật theo phản hồi server hoặc truy vấn lại trạng thái phòng.")
    p(doc, "Đối với phòng cấm/chọn, tính nhất quán này ảnh hưởng trực tiếp đến trải nghiệm người dùng. Nếu một người chơi thấy còn lượt nhưng server đã chuyển lượt, thao tác sẽ bị từ chối và gây khó hiểu. Vì vậy, cập nhật realtime, thông báo lỗi và truy vấn lại trạng thái cần được phối hợp trong cùng một thiết kế.")
    add_heading(doc, "4.10.3. Định hướng kiểm thử tự động", 3)
    p(doc, "Phiên bản hiện tại tập trung vào kiểm thử chức năng theo luồng người dùng. Trong các bước phát triển tiếp theo, hệ thống nên bổ sung kiểm thử tự động cho các policy quan trọng như DraftPolicy, CostPolicy và RoomAccessPolicy. Đây là các phần có tác động lớn đến tính công bằng của trận đấu và ít phụ thuộc vào giao diện, do đó phù hợp để viết unit test.")
    p(doc, "Bên cạnh unit test, một số kịch bản end-to-end cũng nên được tự động hóa, chẳng hạn tạo phòng, tham gia hai đội, thực hiện một lượt pick hợp lệ và kiểm tra realtime. Kiểm thử tự động không thay thế hoàn toàn kiểm thử thủ công, nhưng giúp phát hiện sớm lỗi hồi quy khi mã nguồn thay đổi.")
    add_heading(doc, "4.11. Quy trình kiểm thử và nghiệm thu", 2)
    add_heading(doc, "4.11.1. Kiểm thử theo vai trò người dùng", 3)
    p(doc, "Quy trình kiểm thử cần mô phỏng đầy đủ các vai trò trong một trận đấu. Host được kiểm tra với các thao tác tạo phòng, mời người chơi, bắt đầu draft, chuyển phase, tạm dừng và tổng kết. Player được kiểm tra với thao tác tham gia đội, thực hiện lượt cấm/chọn, khai báo build và xem kết quả. Spectator/Caster được kiểm tra ở chế độ chỉ xem, bảo đảm vai trò này không thể gửi request điều khiển phòng. Admin được kiểm tra với các thao tác quản lý dữ liệu và cấu hình.")
    p(doc, "Kiểm thử theo vai trò giúp phát hiện lỗi phân quyền tốt hơn so với kiểm thử từng màn hình riêng lẻ. Một màn hình phòng đấu có thể hiển thị giống nhau cho nhiều vai trò, nhưng quyền thao tác của mỗi vai trò khác nhau. Nếu không kiểm thử theo vai trò, hệ thống có thể vô tình cho spectator chuyển phase hoặc cho player sửa cấu hình phòng, làm ảnh hưởng trực tiếp đến tính công bằng của trận đấu.")
    add_heading(doc, "4.11.2. Kiểm thử theo trạng thái phòng", 3)
    p(doc, "Bên cạnh vai trò, trạng thái phòng là trục kiểm thử quan trọng. Ở trạng thái WAITING, hệ thống phải cho phép tham gia đội nhưng chưa cho phép cấm/chọn. Ở trạng thái DRAFTING, hệ thống cho phép thao tác draft đúng lượt nhưng chưa cho phép tổng kết kết quả. Ở trạng thái BUILDING, hệ thống cho phép khai báo build cho nhân vật đã pick. Ở trạng thái FINISHED, hệ thống khóa các thao tác có thể làm thay đổi kết quả.")
    p(doc, "Kiểm thử theo trạng thái giúp bảo đảm API không chấp nhận thao tác ngoài ngữ cảnh. Đây là lỗi thường gặp ở các hệ thống có nhiều phase. Nếu API chỉ kiểm tra dữ liệu đầu vào mà không kiểm tra trạng thái phòng, người dùng có thể gửi request build trước khi draft kết thúc hoặc gửi request draft sau khi phòng đã tổng kết.")
    add_heading(doc, "4.11.3. Kiểm thử realtime với nhiều client", 3)
    p(doc, "Đối với chức năng realtime, kiểm thử cần mở ít nhất ba client đại diện cho Host, Player và Spectator. Khi một client thực hiện thao tác, các client còn lại phải nhận đúng cập nhật. Trường hợp kiểm thử không nên chỉ dừng ở việc giao diện thay đổi, mà cần kiểm tra nội dung cập nhật có đúng với dữ liệu server hay không, ví dụ nhân vật được đưa vào đúng đội, lượt tiếp theo đúng thứ tự và pool nhân vật đã khóa đúng thẻ.")
    p(doc, "Ngoài kịch bản kết nối ổn định, cần kiểm tra trường hợp một client mất kết nối hoặc tải lại trang giữa trận. Khi client quay lại, trạng thái phải được khôi phục từ database. Đây là tiêu chí quan trọng vì trong thực tế thi đấu, người dùng có thể mất mạng tạm thời, đổi thiết bị hoặc mở lại trang khi trình duyệt gặp lỗi.")
    add_heading(doc, "4.11.4. Tiêu chí nghiệm thu sản phẩm", 3)
    p(doc, "Một phiên bản có thể được xem là đạt yêu cầu nghiệm thu khi luồng chính hoạt động liên tục từ tạo phòng đến tổng kết kết quả, không cần chỉnh sửa dữ liệu thủ công trong database. Người dùng phải thực hiện được thao tác đúng vai trò, thao tác sai phải bị từ chối có thông báo, kết quả Cost phải có dữ liệu đối chiếu và giao diện theo dõi phải cập nhật cho các vai trò liên quan.")
    p(doc, "Ngoài luồng chính, sản phẩm cần đạt tiêu chí trình bày: giao diện không che khuất thông tin quan trọng, màu đội và trạng thái dễ nhận biết, bảng kết quả đủ rõ để chụp lại làm minh chứng. Các tiêu chí này cho thấy sản phẩm không chỉ đúng về mặt kỹ thuật mà còn có khả năng sử dụng trong bối cảnh tổ chức trận đấu thật.")
    add_heading(doc, "4.12. Ghi nhận kết quả triển khai", 2)
    add_heading(doc, "4.12.1. Kết quả về chức năng lõi", 3)
    p(doc, "Kết quả triển khai cho thấy hệ thống đã hình thành được bộ chức năng lõi phục vụ vòng đời trận đấu. Các màn hình chính như chọn vai trò, phòng cấm/chọn, khai báo build, kết quả và tournament đã có cơ sở giao diện. Mã nguồn cũng thể hiện các nhóm chức năng tương ứng trong src/app và src/components, phù hợp với cấu trúc trình bày ở phần thiết kế.")
    p(doc, "Điểm đáng chú ý là hệ thống không chỉ có trang tĩnh mà đã có định hướng dữ liệu cụ thể. Các bảng trong Prisma Schema phản ánh trực tiếp các thành phần nghiệp vụ, từ phòng đấu đến draft log, build, chat, tournament, notification và activity. Điều này tạo nền tảng để tiếp tục hoàn thiện chức năng thay vì phải thiết kế lại toàn bộ cơ sở dữ liệu.")
    add_heading(doc, "4.12.2. Kết quả về khả năng trình bày sản phẩm", 3)
    p(doc, "Các ảnh giao diện cho thấy sản phẩm có định hướng thị giác phù hợp với công cụ thi đấu: nền tối, hai màu đội rõ ràng, pool nhân vật trực quan và khu vực kết quả nổi bật. Giao diện không chỉ nhằm mục đích trang trí mà phục vụ việc nhận biết trạng thái nhanh trong quá trình cấm/chọn. Đây là yêu cầu quan trọng vì người dùng trong trận đấu thường cần thao tác nhanh và ít có thời gian đọc hướng dẫn dài.")
    p(doc, "Tuy nhiên, để đạt mức sản phẩm hoàn chỉnh hơn, giao diện cần tiếp tục được kiểm tra trên nhiều kích thước màn hình và nhiều trình duyệt. Các thành phần như bảng nhân vật, timer, modal xác nhận và vùng kết quả cần có kích thước ổn định để tránh lỗi tràn nội dung khi tên nhân vật, tên đội hoặc thông báo dài hơn dự kiến.")


def chapter_5(doc):
    add_heading(doc, "CHƯƠNG 5: TỔNG KẾT, ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN", 1, page_break=True)
    add_heading(doc, "5.1. Kết quả đạt được", 2)
    add_heading(doc, "5.1.1. Kết quả về nghiệp vụ", 3)
    p(doc, "Đề tài đã xây dựng được nền tảng website hỗ trợ tổ chức một phiên cấm/chọn nhân vật Genshin Impact theo luồng có trạng thái. Hệ thống hỗ trợ tạo phòng, phân vai người dùng, thực hiện lượt cấm/chọn, khai báo build, tính Cost và hiển thị kết quả Time Handicap. Đây là các chức năng cốt lõi để thay thế quy trình tổ chức thủ công bằng một công cụ có dữ liệu tập trung.")
    add_heading(doc, "5.1.2. Kết quả về kỹ thuật", 3)
    p(doc, "Về kỹ thuật, hệ thống áp dụng Next.js, TypeScript, Prisma, PostgreSQL, Supabase Realtime, Tailwind CSS và Docker. Mã nguồn được định hướng phân tầng, trong đó các quy tắc quan trọng được tách khỏi giao diện. Cơ sở dữ liệu có các bảng phục vụ trận đấu, giải đấu và mở rộng cộng đồng.")
    add_heading(doc, "5.1.3. Kết quả về giao diện", 3)
    p(doc, "Giao diện được thiết kế theo hướng thi đấu, ưu tiên khả năng nhận diện nhanh trạng thái. Khu vực đội xanh, đội đỏ, pool nhân vật, lượt hiện tại và trang kết quả được trình bày rõ ràng. Các ảnh giao diện trong chương 4 cho thấy hệ thống không chỉ có dữ liệu phía sau mà còn có trải nghiệm người dùng tương đối hoàn chỉnh.")
    add_heading(doc, "5.2. Đối chiếu mục tiêu ban đầu", 2)
    add_caption(doc, "5.1. Đối chiếu mục tiêu ban đầu và kết quả đạt được", "Bảng")
    add_table(
        doc,
        ["Mục tiêu ban đầu", "Kết quả hiện tại", "Mức độ"],
        [
            ["Chuẩn hóa quy trình cấm/chọn", "Đã có room, draft turn, DraftLog và giao diện phân đội.", "Hoàn thành"],
            ["Tính Cost và Time Handicap", "Đã thiết kế CharacterBuild, CostService và trang kết quả.", "Hoàn thành"],
            ["Theo dõi realtime", "Đã định hướng sử dụng Supabase Realtime để cập nhật trạng thái nhiều client.", "Hoàn thành ở mức đề tài"],
            ["Mở rộng tournament", "Đã có model Tournament, Participant, Match và giao diện liên quan.", "Hoàn thành một phần"],
            ["Vận hành quy mô lớn", "Chưa có kiểm thử tải và giám sát production đầy đủ.", "Cần phát triển thêm"],
        ],
        widths=[5.2, 7.5, 3.0],
    )
    add_heading(doc, "5.3. Ưu điểm của hệ thống", 2)
    add_heading(doc, "5.3.1. Minh bạch hóa quy trình", 3)
    p(doc, "Ưu điểm quan trọng nhất là hệ thống minh bạch hóa quy trình cấm/chọn. Mỗi lượt thao tác được ghi nhận, trạng thái được hiển thị cho các vai trò liên quan và kết quả Cost có dữ liệu đối chiếu. Điều này giúp giảm tranh luận trong quá trình tổ chức trận đấu.")
    add_heading(doc, "5.3.2. Có nền tảng mở rộng", 3)
    p(doc, "Thiết kế dữ liệu không chỉ phục vụ một phòng đơn mà còn chuẩn bị cho giải đấu, thông báo, hồ sơ người chơi và hoạt động cộng đồng. Đây là điểm có ý nghĩa nếu sản phẩm tiếp tục được phát triển sau học phần.")
    add_heading(doc, "5.3.3. Phù hợp công nghệ hiện đại", 3)
    p(doc, "Việc sử dụng Next.js, TypeScript, Prisma và Supabase giúp đề tài bám sát xu hướng phát triển ứng dụng web hiện đại. Các công nghệ này cũng phù hợp với mục tiêu học phần vì thể hiện được nhiều khía cạnh từ giao diện, backend, database đến triển khai.")
    add_heading(doc, "5.4. Hạn chế còn tồn tại", 2)
    add_heading(doc, "5.4.1. Hạn chế về xác minh dữ liệu", 3)
    p(doc, "Hệ thống chưa thể tự động xác minh tuyệt đối toàn bộ thông tin nhân vật, cung mệnh và vũ khí trong game. Một phần dữ liệu vẫn cần người chơi khai báo hoặc trọng tài kiểm tra, đặc biệt khi nguồn dữ liệu ngoài không đầy đủ hoặc người chơi không công khai hồ sơ.")
    add_heading(doc, "5.4.2. Hạn chế về kiểm thử tải", 3)
    p(doc, "Phiên bản hiện tại chưa có kiểm thử tải với số lượng lớn phòng realtime đồng thời. Do đó, khi triển khai cho giải đấu lớn, hệ thống cần được bổ sung đo đạc hiệu năng, giám sát kết nối và chiến lược xử lý khi nhiều client cập nhật trạng thái cùng lúc.")
    add_heading(doc, "5.4.3. Hạn chế về tài liệu vận hành", 3)
    p(doc, "Tài liệu cấu hình triển khai, sao lưu dữ liệu, phục hồi sự cố và quy trình quản trị cần được bổ sung nếu sản phẩm được vận hành lâu dài. Đây là các nội dung cần thiết để chuyển từ đồ án học phần sang sản phẩm cộng đồng thực tế.")
    add_heading(doc, "5.5. Hướng phát triển", 2)
    add_heading(doc, "5.5.1. Hoàn thiện xác minh build", 3)
    p(doc, "Hướng phát triển đầu tiên là cải thiện cơ chế xác minh build bằng cách kết hợp dữ liệu người chơi, ảnh minh chứng hoặc cơ chế xác nhận của trọng tài. Mục tiêu không nhất thiết là tự động hóa tuyệt đối, mà là giảm tranh luận và tăng tính tin cậy của dữ liệu đầu vào.")
    add_heading(doc, "5.5.2. Phát triển tournament hoàn chỉnh", 3)
    p(doc, "Hệ thống có thể tiếp tục phát triển bracket, lịch thi đấu, seed, kết quả nhiều ván và fearless draft cho các thể thức BO3 hoặc BO5. Khi đó, Room sẽ trở thành đơn vị trận đấu con trong một TournamentMatch, giúp dữ liệu giải đấu được tổ chức nhất quán.")
    add_heading(doc, "5.5.3. Tối ưu realtime và vận hành", 3)
    p(doc, "Về vận hành, cần bổ sung kiểm thử tải, log tập trung, cảnh báo lỗi và chiến lược rollback khi triển khai. Đối với realtime, cần kiểm tra độ trễ, tình huống mất kết nối, mở nhiều tab và thao tác đồng thời để bảo đảm trải nghiệm ổn định.")
    add_heading(doc, "5.5.4. Hoàn thiện trải nghiệm quan sát", 3)
    p(doc, "Một hướng phát triển có giá trị đối với cộng đồng là hoàn thiện overlay cho caster và khán giả. Overlay cần hiển thị gọn trạng thái lượt, danh sách nhân vật đã cấm/chọn, tên đội, tổng Cost và kết quả bù trừ thời gian. Khi được thiết kế tốt, overlay có thể dùng trực tiếp trong phần mềm livestream mà không cần chỉnh sửa thủ công.")
    p(doc, "Trải nghiệm quan sát cũng cần tính đến độ trễ phát sóng. Một số trận đấu có thể cần spectatorDelay để tránh lộ chiến thuật hoặc giảm khả năng người xem truyền thông tin ngược cho người chơi. Vì vậy, hệ thống nên cho phép Host cấu hình độ trễ và hiển thị rõ trạng thái cho caster.")
    add_heading(doc, "5.5.5. Bổ sung thống kê sau giải đấu", 3)
    p(doc, "Sau khi hệ thống lưu đủ dữ liệu DraftLog và CharacterBuild, có thể phát triển thống kê sau giải đấu như tỉ lệ nhân vật bị cấm, tỉ lệ nhân vật được chọn, đội hình phổ biến, Cost trung bình và mức Time Handicap thường gặp. Các thống kê này giúp ban tổ chức đánh giá meta của giải, đồng thời cung cấp nội dung tổng kết hấp dẫn cho cộng đồng.")
    p(doc, "Phần thống kê cần được thiết kế theo hướng truy vấn từ dữ liệu đã xác nhận, tránh dựa vào dữ liệu tạm thời trên client. Khi số lượng trận tăng, dữ liệu thống kê có thể trở thành giá trị riêng của hệ thống, giúp sản phẩm không chỉ phục vụ thao tác trong trận mà còn hỗ trợ phân tích sau giải.")
    add_heading(doc, "5.6. Kết luận", 2)
    p(doc, "Đề tài xây dựng website hỗ trợ cấm/chọn nhân vật trong Genshin Impact đã giải quyết một bài toán cụ thể của cộng đồng người chơi: tổ chức trận đấu có luật, có lượt, có dữ liệu và có kết quả minh bạch. Sản phẩm không can thiệp vào trò chơi chính thức mà đóng vai trò công cụ hỗ trợ điều phối và ghi nhận dữ liệu bên ngoài trò chơi.")
    p(doc, "Thông qua quá trình thực hiện, đề tài cho thấy khả năng vận dụng các công nghệ phát triển web hiện đại vào một bài toán có trạng thái nghiệp vụ rõ ràng. Kết quả đạt được không chỉ là giao diện thao tác, mà còn là mô hình dữ liệu, kiến trúc phân tầng, cơ chế realtime và định hướng mở rộng. Những hạn chế hiện tại là cơ sở để tiếp tục hoàn thiện sản phẩm trong tương lai.")
    add_heading(doc, "5.7. Bài học kinh nghiệm", 2)
    add_heading(doc, "5.7.1. Bài học về phân tích nghiệp vụ", 3)
    p(doc, "Một bài học quan trọng trong quá trình thực hiện là cần phân tích nghiệp vụ trước khi xây dựng giao diện. Nếu bắt đầu từ giao diện, người phát triển dễ tập trung vào bố cục thẻ nhân vật hoặc màu sắc nhưng bỏ sót các câu hỏi quan trọng như ai có quyền thao tác, thao tác hợp lệ trong trạng thái nào và dữ liệu nào cần lưu để khôi phục trận đấu.")
    p(doc, "Khi nghiệp vụ được mô tả rõ bằng actor, use case, trạng thái và bảng dữ liệu, quá trình hiện thực trở nên có định hướng hơn. Mỗi component hoặc API đều có lý do tồn tại và gắn với một bước trong vòng đời trận đấu. Điều này giúp báo cáo cũng có mạch trình bày chặt chẽ hơn, tránh liệt kê chức năng một cách rời rạc.")
    add_heading(doc, "5.7.2. Bài học về tổ chức mã nguồn", 3)
    p(doc, "Đối với một dự án có nhiều trạng thái và nhiều vai trò, tổ chức mã nguồn có ảnh hưởng lớn đến khả năng sửa lỗi. Nếu logic kiểm tra lượt nằm trong nhiều component khác nhau, việc thay đổi luật draft sẽ rất khó kiểm soát. Ngược lại, khi quy tắc được đưa vào policy/service, thay đổi luật có thể được thực hiện tập trung hơn và kiểm thử dễ hơn.")
    p(doc, "Bài học này cho thấy kiến trúc phân tầng không phải là phần trình bày hình thức trong báo cáo, mà có tác động trực tiếp đến chất lượng sản phẩm. Một hệ thống nhỏ vẫn cần ranh giới trách nhiệm rõ ràng nếu bài toán có nhiều trạng thái nghiệp vụ.")
    add_heading(doc, "5.7.3. Bài học về trình bày báo cáo", 3)
    p(doc, "Trong báo cáo kỹ thuật, nội dung cần được trình bày bằng văn phong học thuật, tránh dùng câu văn quá ngắn hoặc thiên về văn nói. Bảng biểu chỉ nên dùng khi dữ liệu có tính so sánh hoặc tra cứu; nếu nội dung là giải thích, nên viết thành đoạn văn hoặc tiểu mục. Việc lạm dụng bảng với nội dung quá ngắn có thể làm báo cáo dài về hình thức nhưng thiếu chiều sâu.")
    p(doc, "Bản báo cáo hoàn chỉnh vì vậy được điều chỉnh theo hướng tăng chất lượng diễn giải, bổ sung heading cấp ba, viết rõ mối liên hệ giữa nghiệp vụ và kỹ thuật, đồng thời giữ bảng biểu cho các phần thật sự cần so sánh. Đây cũng là hướng cải thiện so với bản nháp trước đó.")
    add_heading(doc, "5.8. Đánh giá giá trị học tập của đề tài", 2)
    add_heading(doc, "5.8.1. Giá trị trong phân tích và thiết kế hệ thống", 3)
    p(doc, "Đề tài giúp sinh viên thực hành đầy đủ hơn quy trình phân tích và thiết kế hệ thống. Thay vì chỉ xây dựng một website có nhiều trang, sinh viên phải xác định bối cảnh sử dụng, vai trò người dùng, yêu cầu chức năng, yêu cầu phi chức năng, mô hình dữ liệu, luồng xử lý và tiêu chí kiểm thử. Đây là các nội dung cốt lõi của kỹ nghệ phần mềm và có giá trị vượt ra ngoài phạm vi riêng của trò chơi Genshin Impact.")
    p(doc, "Một điểm học tập quan trọng là cách chuyển yêu cầu không chính thức của cộng đồng thành mô hình kỹ thuật. Ban đầu, yêu cầu có thể chỉ được mô tả bằng ngôn ngữ đời thường như 'cần một trang để ban pick cho tiện'. Tuy nhiên, để xây dựng thành hệ thống, yêu cầu đó phải được chuyển thành use case, trạng thái phòng, actor, ràng buộc dữ liệu và API. Quá trình chuyển đổi này là năng lực cần thiết khi phát triển phần mềm trong thực tế.")
    add_heading(doc, "5.8.2. Giá trị trong phát triển ứng dụng web hiện đại", 3)
    p(doc, "Đề tài cũng giúp sinh viên tiếp cận mô hình phát triển web hiện đại, nơi frontend, backend, cơ sở dữ liệu và triển khai không còn tách rời hoàn toàn. Next.js cho phép xây dựng giao diện và API trong cùng một dự án; Prisma giúp liên kết TypeScript với PostgreSQL; Supabase bổ sung realtime; Docker hỗ trợ đóng gói môi trường chạy. Việc phối hợp các công nghệ này giúp sinh viên hiểu rõ hơn cách một ứng dụng web vận hành từ giao diện đến dữ liệu.")
    p(doc, "Bên cạnh việc sử dụng công nghệ, đề tài cũng đặt ra yêu cầu lựa chọn công nghệ phù hợp với bài toán. Không phải công nghệ nào hiện đại cũng cần đưa vào sản phẩm. Mỗi công nghệ phải có vai trò rõ ràng và giải quyết một vấn đề cụ thể. Đây là bài học quan trọng để tránh tình trạng sử dụng nhiều công cụ nhưng hệ thống thiếu mạch thiết kế.")
    add_heading(doc, "5.8.3. Giá trị trong tư duy sản phẩm", 3)
    p(doc, "Từ góc độ sản phẩm, đề tài cho thấy một phần mềm tốt không chỉ cần có chức năng đúng mà còn cần phù hợp với ngữ cảnh sử dụng. Trong phòng cấm/chọn, người dùng cần thao tác nhanh, nhìn thấy trạng thái rõ ràng và tin tưởng kết quả. Vì vậy, các quyết định về bố cục, màu sắc, thông báo lỗi, realtime và kết quả đều phải phục vụ mục tiêu này.")
    p(doc, "Tư duy sản phẩm cũng thể hiện ở việc phân biệt chức năng lõi và chức năng mở rộng. Chức năng lõi là phòng đấu, lượt cấm/chọn, build, Cost và kết quả. Chức năng mở rộng là tournament, social, overlay và thống kê. Khi biết phân tầng mức độ ưu tiên, quá trình phát triển sẽ tập trung hơn và tránh rơi vào tình trạng mở rộng quá nhiều nhưng phần lõi chưa ổn định.")
    add_heading(doc, "5.9. Nhận xét tổng hợp", 2)
    p(doc, "Nhìn tổng thể, đề tài có giá trị ở việc kết hợp một nhu cầu cụ thể của cộng đồng với các công nghệ phát triển web hiện đại. Sản phẩm có tính ứng dụng rõ ràng, dữ liệu có mô hình quan hệ, giao diện có định hướng thi đấu và kiến trúc có phân tầng. Những yếu tố này giúp báo cáo không chỉ mô tả một website đã làm, mà còn trình bày được tư duy phân tích, thiết kế và đánh giá hệ thống.")
    p(doc, "Phiên bản hiện tại vẫn còn không gian để phát triển, đặc biệt ở kiểm thử tải, xác minh dữ liệu build và vận hành giải đấu quy mô lớn. Tuy nhiên, trong phạm vi học phần, đề tài đã thể hiện được quá trình chuyển từ bài toán thực tế sang sản phẩm phần mềm có cấu trúc. Đây là nền tảng quan trọng để tiếp tục hoàn thiện nếu sản phẩm được triển khai cho cộng đồng người chơi sau này.")


def references_and_appendix(doc):
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1, page_break=True)
    references = [
        ("Vercel", "Next.js Documentation", "Next.js Documentation", "https://nextjs.org/docs"),
        ("Meta", "React Reference Overview", "React Documentation", "https://react.dev/reference/react"),
        ("Microsoft", "TypeScript Documentation", "TypeScript Documentation", "https://www.typescriptlang.org/docs/"),
        ("Tailwind Labs", "Tailwind CSS Documentation", "Tailwind CSS Documentation", "https://tailwindcss.com/docs"),
        ("Supabase", "Supabase Documentation", "Supabase Documentation", "https://supabase.com/docs"),
        ("Supabase", "Realtime", "Supabase Documentation", "https://supabase.com/docs/guides/realtime"),
        ("Prisma Data", "Prisma ORM Documentation", "Prisma Documentation", "https://www.prisma.io/docs/orm"),
        ("PostgreSQL Global Development Group", "Row Security Policies", "PostgreSQL Documentation", "https://www.postgresql.org/docs/current/ddl-rowsecurity.html"),
        ("Docker Inc.", "Docker Documentation", "Docker Docs", "https://docs.docker.com/"),
        ("pmndrs", "Zustand Documentation", "Zustand Documentation", "https://zustand.docs.pmnd.rs/"),
        ("Niklas von Hertzen", "html2canvas Documentation", "html2canvas Documentation", "https://html2canvas.hertzen.com/documentation"),
        ("Genshin.dev", "Fan-Made Genshin Impact API", "Genshin.dev Documentation", "https://genshin.dev/"),
        ("Enka.Network", "API Documentation", "Enka.Network API Documentation", "https://api.enka.network/"),
    ]
    for idx, (author, title, source, url) in enumerate(references, 1):
        entry = doc.add_paragraph()
        entry.paragraph_format.first_line_indent = Cm(0)
        entry.paragraph_format.space_after = Pt(0)
        entry.add_run(f'[{idx}] {author}, "{title}," ')
        source_run = entry.add_run(source)
        source_run.italic = True
        entry.add_run(", 2026.")

        link_line = doc.add_paragraph()
        link_line.paragraph_format.first_line_indent = Cm(0)
        link_line.paragraph_format.space_after = Pt(8)
        link_line.add_run("[Online]. Available: ")
        add_hyperlink(link_line, url)

    add_heading(doc, "PHỤ LỤC A: MINH CHỨNG SỬ DỤNG CÔNG CỤ AI", 1, page_break=True)
    p(doc, "Trong quá trình thực hiện đề tài, công cụ AI được sử dụng như một phương tiện hỗ trợ phân tích, soạn thảo, kiểm tra lỗi và đề xuất cách trình bày. Nội dung do AI gợi ý được sinh viên rà soát, chỉnh sửa và đối chiếu với mã nguồn thực tế trước khi đưa vào báo cáo.")
    add_caption(doc, "A.1. Minh chứng sử dụng công cụ AI trong quá trình phát triển", "Bảng")
    add_table(
        doc,
        ["STT", "Công cụ", "Mục đích sử dụng", "Kết quả/ghi chú"],
        [
            ["1", "ChatGPT/Codex", "Hỗ trợ rà soát văn phong báo cáo và đề xuất cấu trúc heading.", "Nội dung được chỉnh sửa lại theo văn phong học thuật trước khi sử dụng."],
            ["2", "ChatGPT/Codex", "Hỗ trợ mô tả use case, kiến trúc phân tầng và luồng xử lý.", "Các sơ đồ được dựng lại để phù hợp với báo cáo."],
            ["3", "AI hỗ trợ lập trình", "Gợi ý cách kiểm tra lỗi TypeScript, Prisma và tổ chức service.", "Kết quả được kiểm tra với mã nguồn thực tế."],
            ["4", "Công cụ kiểm tra nội dung", "Rà soát lỗi chính tả, câu quá ngắn và bảng bị trình bày rời rạc.", "Bản hoàn chỉnh đã chuẩn hóa lại câu văn, bảng và caption."],
        ],
        widths=[1.2, 3.8, 6.0, 4.7],
    )
    add_heading(doc, "PHỤ LỤC B: GHI CHÚ TRIỂN KHAI", 1, page_break=True)
    p(doc, "Để triển khai hệ thống trong môi trường thực tế, cần chuẩn bị cơ sở dữ liệu PostgreSQL, cấu hình Supabase, khai báo biến môi trường, build ứng dụng và kiểm tra các route quan trọng. Các thông tin nhạy cảm như DATABASE_URL, DIRECT_URL và khóa API không được đưa trực tiếp vào mã nguồn hoặc báo cáo công khai.")
    add_code_block(
        doc,
        "DATABASE_URL=postgresql://...\nDIRECT_URL=postgresql://...\nNEXT_PUBLIC_SUPABASE_URL=https://...\nNEXT_PUBLIC_SUPABASE_ANON_KEY=...\n",
    )


STYLE_GUARD_REPLACEMENTS = {
    "Trong quá trình thực hiện đề tài, em xin chân thành cảm ơn quý thầy cô Khoa Công nghệ Thông tin, Trường Đại học Đà Lạt đã truyền đạt những kiến thức nền tảng về phân tích thiết kế hệ thống, lập trình web, cơ sở dữ liệu và triển khai phần mềm. Đây là cơ sở quan trọng giúp em có thể tiếp cận đề tài theo hướng hệ thống, không chỉ dừng lại ở việc xây dựng giao diện mà còn chú trọng đến quy trình nghiệp vụ, tổ chức dữ liệu và khả năng vận hành thực tế.":
    "Trong quá trình thực hiện đề tài, em xin chân thành cảm ơn quý thầy cô Khoa Công nghệ Thông tin, Trường Đại học Đà Lạt đã truyền đạt những kiến thức nền tảng về phân tích thiết kế hệ thống, lập trình web, cơ sở dữ liệu và triển khai phần mềm. Những kiến thức này là cơ sở để em tiếp cận đề tài theo hướng hệ thống, bao gồm quy trình nghiệp vụ, tổ chức dữ liệu, giao diện và khả năng vận hành thực tế.",
    "Trong bối cảnh thi đấu cộng đồng, yếu tố công bằng không chỉ phụ thuộc vào kỹ năng điều khiển nhân vật mà còn chịu ảnh hưởng bởi tài nguyên tài khoản của mỗi người chơi. Một người chơi sở hữu nhiều nhân vật 5 sao, cung mệnh cao hoặc vũ khí 5 sao có thể có lợi thế rõ rệt. Vì vậy, các giải đấu cộng đồng thường cần cơ chế kiểm soát lựa chọn đội hình và cơ chế quy đổi lợi thế thành Time Handicap.":
    "Trong bối cảnh thi đấu cộng đồng, yếu tố công bằng phụ thuộc vào kỹ năng điều khiển nhân vật và tài nguyên tài khoản của mỗi người chơi. Tài khoản sở hữu nhiều nhân vật 5 sao, cung mệnh cao hoặc vũ khí 5 sao có thể tạo lợi thế rõ rệt. Vì vậy, các giải đấu cộng đồng thường cần cơ chế kiểm soát lựa chọn đội hình và cơ chế quy đổi lợi thế thành Time Handicap.",
    "Hệ thống chưa đặt mục tiêu thay thế cơ chế chống gian lận trong trò chơi. Dữ liệu build có thể được người chơi khai báo thủ công hoặc hỗ trợ lấy một phần từ nguồn ngoài nếu người chơi công khai thông tin phù hợp. Việc xác minh tuyệt đối toàn bộ thông tin trong game không thuộc phạm vi phiên bản hiện tại do giới hạn API chính thức và phạm vi triển khai học phần.":
    "Hệ thống chưa đặt mục tiêu thay thế cơ chế chống gian lận trong trò chơi. Dữ liệu cấu hình nhân vật có thể được người chơi khai báo thủ công hoặc lấy một phần từ nguồn ngoài nếu người chơi công khai thông tin phù hợp. Việc xác minh tuyệt đối toàn bộ thông tin trong trò chơi không thuộc phạm vi phiên bản hiện tại do giới hạn API chính thức và phạm vi triển khai học phần.",
    "Quá trình thực hiện bắt đầu từ việc mô tả quy trình tổ chức một trận cấm/chọn trong cộng đồng, sau đó chuyển quy trình này thành actor, use case, yêu cầu chức năng, yêu cầu phi chức năng và mô hình dữ liệu. Cách tiếp cận này giúp hệ thống được xây dựng dựa trên nghiệp vụ thực tế thay vì chỉ mô phỏng giao diện bên ngoài.":
    "Quá trình thực hiện bắt đầu từ việc mô tả quy trình tổ chức một trận cấm/chọn trong cộng đồng, sau đó chuyển quy trình này thành actor, use case, yêu cầu chức năng, yêu cầu phi chức năng và mô hình dữ liệu. Việc phân tích nghiệp vụ trước khi hiện thực bảo đảm hệ thống được xây dựng dựa trên quy trình sử dụng thực tế, thay vì chỉ mô phỏng giao diện bên ngoài.",
    "Sau khi xác định yêu cầu, đề tài tổ chức mã nguồn theo các tầng có trách nhiệm rõ ràng. Tầng giao diện tiếp nhận thao tác và hiển thị trạng thái; tầng application điều phối use case; tầng domain chứa policy kiểm tra luật; tầng infrastructure chịu trách nhiệm truy cập cơ sở dữ liệu và tích hợp dịch vụ ngoài. Cách tổ chức này giúp hệ thống dễ kiểm thử và mở rộng.":
    "Sau khi xác định yêu cầu, mã nguồn được tổ chức theo các tầng có trách nhiệm rõ ràng. Tầng giao diện tiếp nhận thao tác và hiển thị trạng thái; tầng application điều phối use case; tầng domain chứa policy kiểm tra luật; tầng infrastructure chịu trách nhiệm truy cập cơ sở dữ liệu và tích hợp dịch vụ ngoài. Cấu trúc phân tầng tạo điều kiện kiểm thử từng nhóm trách nhiệm và mở rộng chức năng theo phạm vi kiểm soát được.",
    "Trong quá trình kiểm thử, mỗi kịch bản được đặt trong bối cảnh có nhiều vai trò cùng tham gia. Ví dụ, khi Player Blue thực hiện lượt chọn, hệ thống không chỉ cần kiểm tra API có lưu DraftLog hay không, mà còn cần quan sát giao diện của Player Red, Host và Spectator có nhận đúng trạng thái mới hay không. Cách kiểm thử này phù hợp với đặc thù của ứng dụng realtime, nơi lỗi thường xuất hiện ở sự không đồng bộ giữa các client.":
    "Trong quá trình kiểm thử, mỗi kịch bản được đặt trong bối cảnh có nhiều vai trò cùng tham gia. Khi Player Blue thực hiện lượt chọn, hệ thống cần kiểm tra API có lưu DraftLog chính xác, đồng thời đối chiếu giao diện của Player Red, Host và Spectator sau khi trạng thái mới được đồng bộ. Kiểm thử theo nhiều client phù hợp với đặc thù của ứng dụng realtime, nơi lỗi thường xuất hiện ở sự không đồng bộ giữa các client.",
    "Đối với ban tổ chức, việc lưu trữ dữ liệu có cấu trúc giúp quá trình tổng hợp kết quả và tra cứu lịch sử thuận tiện hơn. Sau mỗi trận, dữ liệu về nhân vật bị cấm, nhân vật được chọn, đội hình, Cost và kết quả có thể được dùng để thống kê hoặc làm minh chứng khi có khiếu nại. Đây là điểm khác biệt đáng kể so với cách lưu dữ liệu rời rạc bằng ảnh chụp hoặc tin nhắn.":
    "Đối với ban tổ chức, dữ liệu có cấu trúc hỗ trợ tổng hợp kết quả và tra cứu lịch sử. Sau mỗi trận, dữ liệu về nhân vật bị cấm, nhân vật được chọn, đội hình, Cost và kết quả có thể được dùng để thống kê hoặc làm minh chứng khi có khiếu nại. So với cách lưu dữ liệu bằng ảnh chụp hoặc tin nhắn, hệ thống lưu DraftLog và CharacterBuild theo cấu trúc có thể truy vấn lại.",
    "Đối với người chơi, hệ thống tạo ra môi trường thao tác rõ ràng hơn. Người chơi biết khi nào đến lượt của họ, thao tác nào hợp lệ và vì sao một thao tác bị từ chối. Điều này giúp giảm tranh luận không cần thiết, đồng thời giúp người chơi tập trung vào chiến thuật lựa chọn đội hình thay vì kiểm tra thủ công trạng thái trận đấu.":
    "Đối với người chơi, hệ thống cung cấp môi trường thao tác rõ ràng hơn. Người chơi biết khi nào đến lượt của họ, thao tác nào hợp lệ và vì sao một thao tác bị từ chối. Cơ chế phản hồi này làm giảm tranh luận không cần thiết và cho phép người chơi tập trung vào chiến thuật lựa chọn đội hình thay vì kiểm tra thủ công trạng thái trận đấu.",
    "Cách tổ chức này nhằm bảo đảm mạch trình bày đi từ nhu cầu thực tế đến giải pháp kỹ thuật, sau đó mới đánh giá kết quả hiện thực. Nhờ đó, người đọc có thể thấy rõ mối liên hệ giữa vấn đề cộng đồng đặt ra, quyết định thiết kế hệ thống và các chức năng đã được xây dựng trong sản phẩm.":
    "Báo cáo trình bày mối liên hệ giữa yêu cầu nghiệp vụ, quyết định thiết kế và kết quả triển khai. Trình tự nội dung đi từ bối cảnh tổ chức cấm/chọn trong cộng đồng đến kiến trúc hệ thống, mô hình dữ liệu, chức năng đã hiện thực và định hướng phát triển.",
    "Về mặt sản phẩm, website hướng đến trải nghiệm sử dụng trong bối cảnh thi đấu thật. Người dùng không chỉ cần một giao diện đẹp, mà cần biết chính xác trạng thái hiện tại, lượt tiếp theo, thao tác hợp lệ và kết quả cuối cùng. Do đó, thiết kế giao diện ưu tiên khả năng nhận diện đội, trạng thái nhân vật, thông tin Cost và khu vực kết quả. Đây là hướng tiếp cận phù hợp với công cụ hỗ trợ thi đấu, khác với website giới thiệu thông tin thông thường.":
    "Về mặt sản phẩm, website hướng đến trải nghiệm sử dụng trong bối cảnh thi đấu thực tế. Giao diện cần cung cấp chính xác trạng thái hiện tại, lượt tiếp theo, thao tác hợp lệ và kết quả cuối cùng. Do đó, thiết kế giao diện ưu tiên khả năng nhận diện đội, trạng thái nhân vật, thông tin Cost và khu vực kết quả. Định hướng này phù hợp với công cụ hỗ trợ thi đấu, khác với website giới thiệu thông tin thông thường.",
    "Docker được sử dụng để đóng gói ứng dụng, giảm sự khác biệt giữa môi trường phát triển và môi trường triển khai. Dockerfile multi-stage giúp tách giai đoạn cài đặt phụ thuộc, build ứng dụng và chạy production, từ đó làm bản triển khai gọn hơn và dễ tái lập hơn.":
    "Docker được sử dụng để đóng gói ứng dụng, giảm sự khác biệt giữa môi trường phát triển và môi trường triển khai. Dockerfile multi-stage tách giai đoạn cài đặt phụ thuộc, build ứng dụng và chạy production, qua đó tối ưu kích thước bản triển khai và tăng khả năng tái lập môi trường.",
    "Điểm đáng chú ý trong đề tài là công nghệ được lựa chọn dựa trên đặc điểm nghiệp vụ, không chỉ dựa trên độ phổ biến. Bài toán cấm/chọn có nhiều người dùng theo dõi cùng một trạng thái và có yêu cầu cập nhật nhanh sau mỗi lượt thao tác. Vì vậy, việc kết hợp Next.js với Supabase Realtime và PostgreSQL tạo ra một cấu trúc phù hợp: Next.js xử lý giao diện và API, PostgreSQL lưu trạng thái chính thức, còn Realtime truyền tín hiệu thay đổi cho các client.":
    "Công nghệ được lựa chọn dựa trên đặc điểm nghiệp vụ, thay vì chỉ dựa trên độ phổ biến. Bài toán cấm/chọn có nhiều người dùng theo dõi cùng một trạng thái và có yêu cầu cập nhật nhanh sau mỗi lượt thao tác. Vì vậy, việc kết hợp Next.js với Supabase Realtime và PostgreSQL tạo ra một cấu trúc phù hợp: Next.js xử lý giao diện và API, PostgreSQL lưu trạng thái chính thức, Realtime truyền tín hiệu thay đổi cho các client.",
    "Prisma Schema đồng thời đóng vai trò như một tài liệu kỹ thuật về cơ sở dữ liệu. Khi đọc schema, người phát triển có thể biết bảng nào có quan hệ một-nhiều, bảng nào có ràng buộc unique và trường nào được đánh index. Điều này đặc biệt hữu ích trong báo cáo vì mô hình dữ liệu không chỉ được trình bày bằng sơ đồ mà còn có cơ sở đối chiếu với mã nguồn thực tế.":
    "Prisma Schema đồng thời là tài liệu kỹ thuật về cơ sở dữ liệu. Schema mô tả bảng có quan hệ một-nhiều, ràng buộc unique và trường được đánh index. Mô hình dữ liệu trong báo cáo vì vậy có cơ sở đối chiếu với mã nguồn thực tế, thay vì chỉ dừng ở sơ đồ khái quát.",
    "Tuy nhiên, Tailwind cũng có rủi ro nếu sử dụng thiếu quy ước. Khi class bị lặp lại quá nhiều hoặc mỗi component tự định nghĩa màu sắc riêng, giao diện dễ mất nhất quán. Vì vậy, báo cáo bổ sung Design System và bảng màu chủ đạo để xác định vai trò từng màu trong hệ thống. Đây là bước cần thiết để giao diện không chỉ hoạt động được mà còn có tính nhận diện và tính chuyên nghiệp.":
    "Tailwind có rủi ro mất nhất quán nếu sử dụng thiếu quy ước. Khi class bị lặp lại quá nhiều hoặc mỗi component tự định nghĩa màu sắc riêng, giao diện dễ thiếu thống nhất. Vì vậy, Design System và bảng màu chủ đạo được sử dụng để xác định vai trò từng màu trong hệ thống. Quy ước này bảo đảm giao diện có tính nhận diện và phù hợp với bối cảnh thi đấu.",
    "Docker giúp quá trình triển khai có tính tái lập cao hơn, đặc biệt khi ứng dụng cần build Next.js, generate Prisma Client và đọc biến môi trường. Trong bối cảnh học phần, Docker không chỉ là công cụ vận hành mà còn là minh chứng rằng sản phẩm có thể được đóng gói thành một ứng dụng chạy độc lập, thay vì chỉ hoạt động trong môi trường phát triển cá nhân.":
    "Docker làm cho quá trình triển khai có tính tái lập cao hơn, đặc biệt khi ứng dụng cần build Next.js, generate Prisma Client và đọc biến môi trường. Trong phạm vi học phần, Docker là cơ chế đóng gói chứng minh ứng dụng có thể vận hành trong môi trường độc lập, thay vì chỉ phụ thuộc vào máy phát triển cá nhân.",
    "Use Case của hệ thống được xác định từ các vai trò thực tế trong một trận đấu cộng đồng. Mỗi use case phải gắn với một mục tiêu cụ thể của actor, không chỉ mô tả tên màn hình. Ví dụ, 'Thực hiện cấm/chọn' không đơn thuần là thao tác trên một thẻ nhân vật, mà bao gồm việc xác định lượt hiện tại, kiểm tra quyền, ghi nhận lịch sử và đồng bộ trạng thái cho các bên liên quan.":
    "Use Case của hệ thống được xác định từ các vai trò thực tế trong một trận đấu cộng đồng. Mỗi use case gắn với một mục tiêu cụ thể của actor, thay vì chỉ mô tả tên màn hình. Ví dụ, 'Thực hiện cấm/chọn' bao gồm thao tác trên thẻ nhân vật, xác định lượt hiện tại, kiểm tra quyền, ghi nhận lịch sử và đồng bộ trạng thái cho các bên liên quan.",
    "Điều kiện thành công của use case này không chỉ là tạo được một bản ghi Room. Phòng được tạo phải đủ dữ liệu để chuyển sang các giai đoạn sau, bao gồm thông tin trạng thái, khả năng gán người chơi và cấu hình luật. Nếu thiếu các giá trị mặc định cần thiết, lỗi có thể xuất hiện muộn hơn ở giai đoạn cấm/chọn hoặc tính kết quả.":
    "Điều kiện thành công của use case này bao gồm việc tạo bản ghi Room và thiết lập đủ dữ liệu để chuyển sang các giai đoạn sau. Phòng cần có thông tin trạng thái, khả năng gán người chơi và cấu hình luật. Nếu thiếu các giá trị mặc định cần thiết, lỗi có thể xuất hiện muộn hơn ở giai đoạn cấm/chọn hoặc tính kết quả.",
    "Sau khi thao tác hợp lệ được ghi nhận, hệ thống cần cập nhật lượt kế tiếp và phát tín hiệu realtime. Kết quả của use case không chỉ là một dòng DraftLog mới, mà là sự thay đổi trạng thái của toàn bộ phòng. Vì vậy, service xử lý use case cần trả về trạng thái đã chuẩn hóa để giao diện không phải tự suy diễn quá nhiều.":
    "Sau khi thao tác hợp lệ được ghi nhận, hệ thống cần cập nhật lượt kế tiếp và phát tín hiệu realtime. Kết quả của use case gồm dòng DraftLog mới và trạng thái mới của toàn bộ phòng. Service xử lý use case cần trả về trạng thái đã chuẩn hóa để giao diện không phải tự suy diễn quá nhiều.",
    "Thông báo lỗi trong luồng cấm/chọn cần được thiết kế như một phần của nghiệp vụ, không chỉ là phản hồi kỹ thuật. Khi thao tác bị từ chối, người dùng cần biết nguyên nhân là sai lượt, sai quyền, nhân vật đã bị khóa hay phòng không còn ở trạng thái draft. Nếu chỉ trả về thông báo chung, người chơi khó hiểu tình huống và trọng tài phải giải thích lại bằng lời.":
    "Thông báo lỗi trong luồng cấm/chọn cần được thiết kế như một phần của nghiệp vụ. Khi thao tác bị từ chối, người dùng cần biết nguyên nhân là sai lượt, sai quyền, nhân vật đã bị khóa hay phòng không còn ở trạng thái draft. Nếu chỉ trả về thông báo chung, người chơi khó xác định tình huống và trọng tài phải giải thích lại bằng lời.",
    "Hệ thống được tổ chức theo kiến trúc phân tầng để tránh việc logic nghiệp vụ nằm rải rác trong component giao diện. Cách tiếp cận này giúp các chức năng quan trọng như kiểm tra lượt, phân quyền phòng, tính Cost và lưu dữ liệu có trách nhiệm rõ ràng hơn.":
    "Hệ thống được tổ chức theo kiến trúc phân tầng để tránh việc logic nghiệp vụ nằm rải rác trong component giao diện. Kiến trúc phân tầng phân định trách nhiệm cho các chức năng quan trọng như kiểm tra lượt, phân quyền phòng, tính Cost và lưu dữ liệu.",
    "API của hệ thống được thiết kế theo hướng phục vụ use case, không chỉ phản ánh trực tiếp tên bảng dữ liệu. Ví dụ, API draft không đơn thuần tạo một DraftLog, mà xử lý toàn bộ nghiệp vụ của một lượt cấm/chọn. Cách tiếp cận này giúp API phù hợp hơn với hành vi người dùng và giảm nguy cơ client phải tự ghép nhiều thao tác nhỏ để hoàn thành một nghiệp vụ.":
    "API của hệ thống được thiết kế theo hướng phục vụ use case, thay vì chỉ phản ánh trực tiếp tên bảng dữ liệu. Ví dụ, API draft xử lý toàn bộ nghiệp vụ của một lượt cấm/chọn, bao gồm kiểm tra quyền, kiểm tra lượt và ghi DraftLog. Thiết kế này làm cho API phù hợp hơn với hành vi người dùng và giảm nguy cơ client phải tự ghép nhiều thao tác nhỏ để hoàn thành một nghiệp vụ.",
    "Theo chiều vận hành, hệ thống cần được bổ sung giám sát log, kiểm thử tải realtime, sao lưu cơ sở dữ liệu và cơ chế phục hồi khi dịch vụ ngoài gặp lỗi. Khi số lượng phòng đồng thời tăng, vấn đề không chỉ là giao diện có hoạt động hay không, mà còn là độ trễ cập nhật, số kết nối realtime và khả năng duy trì trạng thái nhất quán.":
    "Theo chiều vận hành, hệ thống cần được bổ sung giám sát log, kiểm thử tải realtime, sao lưu cơ sở dữ liệu và cơ chế phục hồi khi dịch vụ ngoài gặp lỗi. Khi số lượng phòng đồng thời tăng, phạm vi đánh giá bao gồm độ ổn định giao diện, độ trễ cập nhật, số kết nối realtime và khả năng duy trì trạng thái nhất quán.",
    "Yêu cầu phi chức năng cũng được phản ánh trong kiến trúc hệ thống. Tính đúng đắn được hỗ trợ bằng policy phía server và ràng buộc dữ liệu. Tính minh bạch được hỗ trợ bằng DraftLog và trang kết quả. Tính realtime được hỗ trợ bằng Supabase Realtime kết hợp truy vấn lại trạng thái chính thức. Tính bảo trì được hỗ trợ bằng phân tầng và component hóa. Như vậy, kiến trúc không chỉ là sơ đồ minh họa, mà là cách trả lời các yêu cầu chất lượng của hệ thống.":
    "Yêu cầu phi chức năng được phản ánh trong kiến trúc hệ thống. Tính đúng đắn được hỗ trợ bằng policy phía server và ràng buộc dữ liệu. Tính minh bạch được hỗ trợ bằng DraftLog và trang kết quả. Tính realtime được hỗ trợ bằng Supabase Realtime kết hợp truy vấn lại trạng thái chính thức. Tính bảo trì được hỗ trợ bằng phân tầng và component hóa. Kiến trúc vì vậy là cơ sở kỹ thuật để đáp ứng các yêu cầu chất lượng của hệ thống.",
    "Trong báo cáo kỹ thuật, việc trình bày mối liên hệ giữa yêu cầu và thiết kế có ý nghĩa quan trọng. Nếu chỉ liệt kê công nghệ, người đọc khó đánh giá vì sao hệ thống cần các thành phần đó. Khi yêu cầu được liên hệ với kiến trúc, báo cáo cho thấy quyết định kỹ thuật có căn cứ và phù hợp với bài toán.":
    "Phần thiết kế trình bày mối liên hệ giữa yêu cầu và thành phần kỹ thuật tương ứng. Nếu chỉ liệt kê công nghệ, báo cáo thiếu cơ sở giải thích vì sao hệ thống cần các thành phần đó. Khi yêu cầu được liên hệ với kiến trúc, quyết định kỹ thuật có căn cứ rõ ràng hơn và phù hợp với bài toán.",
    "Sự nhất quán này cần được duy trì ở cả desktop và mobile. Trên desktop, bố cục ba cột giúp quan sát tổng thể; trên mobile, cần sắp xếp theo thứ tự ưu tiên để thao tác không bị rối. Dù bố cục thay đổi theo kích thước màn hình, thông tin cốt lõi như lượt hiện tại, trạng thái nhân vật và kết quả vẫn phải được giữ rõ ràng.":
    "Sự nhất quán này cần được duy trì ở cả desktop và mobile. Trên desktop, bố cục ba cột hỗ trợ quan sát tổng thể; trên mobile, nội dung cần được sắp xếp theo thứ tự ưu tiên để giảm nhầm lẫn khi thao tác. Dù bố cục thay đổi theo kích thước màn hình, thông tin cốt lõi như lượt hiện tại, trạng thái nhân vật và kết quả vẫn phải được giữ rõ ràng.",
    "Mã nguồn được tổ chức theo cách kết hợp route của Next.js với các thư mục component, service và dữ liệu. Các route trong src/app phản ánh các màn hình hoặc nhóm API, trong khi src/components chứa các thành phần giao diện dùng lại. Cách tổ chức này giúp người phát triển xác định nhanh vị trí liên quan khi cần sửa một chức năng.":
    "Mã nguồn kết hợp route của Next.js với các thư mục component, service và dữ liệu. Các route trong src/app phản ánh các màn hình hoặc nhóm API, trong khi src/components chứa các thành phần giao diện dùng lại. Cấu trúc này hỗ trợ người phát triển xác định vị trí liên quan khi cần sửa một chức năng.",
    "Luật Cost trong cộng đồng có thể thay đổi theo từng mùa giải hoặc từng nhóm tổ chức. Vì vậy, thiết kế không nên cố định hoàn toàn mọi giá trị trong giao diện. Các thông tin như costPerPoint, costCap, bankTime hoặc quy tắc fearless draft cần được lưu ở cấu hình phòng hoặc giải đấu. Cách tiếp cận này giúp hệ thống linh hoạt hơn khi ban tổ chức thay đổi thể thức.":
    "Luật Cost trong cộng đồng có thể thay đổi theo từng mùa giải hoặc từng nhóm tổ chức. Vì vậy, thiết kế không nên cố định hoàn toàn mọi giá trị trong giao diện. Các thông tin như costPerPoint, costCap, bankTime hoặc quy tắc fearless draft cần được lưu ở cấu hình phòng hoặc giải đấu. Cấu hình tập trung cho phép hệ thống thay đổi thể thức mà không phải chỉnh sửa nhiều component giao diện.",
    "Trang kết quả hiển thị tổng Cost của hai đội, chênh lệch và Time Handicap. Thông tin được bố trí đối xứng để trọng tài và người chơi dễ đối chiếu. Đây là màn hình quan trọng vì kết quả trên trang này là cơ sở trước khi hai bên bước vào phần thi đấu trong game.":
    "Trang kết quả hiển thị tổng Cost của hai đội, chênh lệch và Time Handicap. Thông tin được bố trí đối xứng để trọng tài và người chơi thuận tiện đối chiếu. Đây là màn hình quan trọng vì kết quả trên trang này là cơ sở trước khi hai bên bước vào phần thi đấu trong trò chơi.",
    "Đối với ứng dụng có realtime, sau khi triển khai không chỉ kiểm tra trang có mở được hay không, mà còn cần mở nhiều client để xác nhận sự kiện cập nhật được truyền đúng. Đây là bước kiểm tra quan trọng vì lỗi realtime có thể không xuất hiện khi chỉ dùng một trình duyệt.":
    "Đối với ứng dụng có realtime, sau khi triển khai cần kiểm tra khả năng truy cập trang và mở nhiều client để xác nhận sự kiện cập nhật được truyền đúng. Đây là bước kiểm tra bắt buộc vì lỗi realtime có thể không xuất hiện khi chỉ dùng một trình duyệt.",
    "Chất lượng mã nguồn không chỉ được đánh giá bằng việc chức năng chạy được, mà còn bằng khả năng đọc hiểu, sửa lỗi và mở rộng. Với đề tài này, việc đặt tên thư mục theo nhóm chức năng và tách component theo vai trò giao diện giúp người phát triển mới có thể định vị nhanh phần cần chỉnh sửa. Đây là yếu tố quan trọng vì website có nhiều màn hình, từ phòng đấu đến tournament và social.":
    "Chất lượng mã nguồn được đánh giá bằng khả năng vận hành chức năng, đọc hiểu, sửa lỗi và mở rộng. Với hệ thống này, việc đặt tên thư mục theo nhóm chức năng và tách component theo vai trò giao diện hỗ trợ người phát triển mới định vị phần cần chỉnh sửa. Yêu cầu này cần được chú trọng vì website có nhiều màn hình, từ phòng đấu đến tournament và social.",
    "Ngoài luồng chính, sản phẩm cần đạt tiêu chí trình bày: giao diện không che khuất thông tin quan trọng, màu đội và trạng thái dễ nhận biết, bảng kết quả đủ rõ để chụp lại làm minh chứng. Các tiêu chí này cho thấy sản phẩm không chỉ đúng về mặt kỹ thuật mà còn có khả năng sử dụng trong bối cảnh tổ chức trận đấu thật.":
    "Ngoài luồng chính, sản phẩm cần đạt tiêu chí trình bày: giao diện không che khuất thông tin quan trọng, màu đội và trạng thái dễ nhận biết, bảng kết quả đủ rõ để chụp lại làm minh chứng. Các tiêu chí này đánh giá khả năng sử dụng của hệ thống trong bối cảnh tổ chức trận đấu thực tế, bên cạnh tính đúng đắn về mặt kỹ thuật.",
    "Kết quả triển khai cho thấy hệ thống đã hình thành được bộ chức năng lõi phục vụ vòng đời trận đấu. Các màn hình chính như chọn vai trò, phòng cấm/chọn, khai báo build, kết quả và tournament đã có cơ sở giao diện. Mã nguồn cũng thể hiện các nhóm chức năng tương ứng trong src/app và src/components, phù hợp với cấu trúc trình bày ở phần thiết kế.":
    "Kết quả triển khai gồm bộ chức năng lõi phục vụ vòng đời trận đấu. Các màn hình chính như chọn vai trò, phòng cấm/chọn, khai báo build, kết quả và tournament đã có cơ sở giao diện. Mã nguồn được tổ chức theo các nhóm chức năng tương ứng trong src/app và src/components, phù hợp với cấu trúc trình bày ở phần thiết kế.",
    "Điểm đáng chú ý là hệ thống không chỉ có trang tĩnh mà đã có định hướng dữ liệu cụ thể. Các bảng trong Prisma Schema phản ánh trực tiếp các thành phần nghiệp vụ, từ phòng đấu đến draft log, build, chat, tournament, notification và activity. Điều này tạo nền tảng để tiếp tục hoàn thiện chức năng thay vì phải thiết kế lại toàn bộ cơ sở dữ liệu.":
    "Hệ thống có định hướng dữ liệu cụ thể thay vì chỉ dừng ở các trang tĩnh. Các bảng trong Prisma Schema phản ánh trực tiếp các thành phần nghiệp vụ, từ phòng đấu đến draft log, build, chat, tournament, notification và activity. Cấu trúc dữ liệu hiện tại cho phép tiếp tục hoàn thiện chức năng mà không phải thiết kế lại toàn bộ cơ sở dữ liệu.",
    "Các ảnh giao diện cho thấy sản phẩm có định hướng thị giác phù hợp với công cụ thi đấu: nền tối, hai màu đội rõ ràng, pool nhân vật trực quan và khu vực kết quả nổi bật. Giao diện không chỉ nhằm mục đích trang trí mà phục vụ việc nhận biết trạng thái nhanh trong quá trình cấm/chọn. Đây là yêu cầu quan trọng vì người dùng trong trận đấu thường cần thao tác nhanh và ít có thời gian đọc hướng dẫn dài.":
    "Các ảnh giao diện ghi nhận định hướng thị giác phù hợp với công cụ thi đấu: nền tối, hai màu đội rõ ràng, pool nhân vật trực quan và khu vực kết quả nổi bật. Giao diện phục vụ việc nhận biết trạng thái nhanh trong quá trình cấm/chọn, thay vì chỉ đóng vai trò trang trí. Yêu cầu này phù hợp với bối cảnh người dùng cần thao tác nhanh trong trận đấu.",
    "Giao diện được thiết kế theo hướng thi đấu, ưu tiên khả năng nhận diện nhanh trạng thái. Khu vực đội xanh, đội đỏ, pool nhân vật, lượt hiện tại và trang kết quả được trình bày rõ ràng. Các ảnh giao diện trong chương 4 cho thấy hệ thống không chỉ có dữ liệu phía sau mà còn có trải nghiệm người dùng tương đối hoàn chỉnh.":
    "Giao diện được thiết kế theo hướng thi đấu, ưu tiên khả năng nhận diện nhanh trạng thái. Khu vực đội xanh, đội đỏ, pool nhân vật, lượt hiện tại và trang kết quả được trình bày rõ ràng. Các ảnh giao diện trong chương 4 ghi nhận mức độ hoàn thiện của trải nghiệm sử dụng bên cạnh phần dữ liệu và xử lý nghiệp vụ.",
    "Ưu điểm quan trọng nhất là hệ thống minh bạch hóa quy trình cấm/chọn. Mỗi lượt thao tác được ghi nhận, trạng thái được hiển thị cho các vai trò liên quan và kết quả Cost có dữ liệu đối chiếu. Điều này giúp giảm tranh luận trong quá trình tổ chức trận đấu.":
    "Ưu điểm quan trọng nhất là hệ thống minh bạch hóa quy trình cấm/chọn. Mỗi lượt thao tác được ghi nhận, trạng thái được hiển thị cho các vai trò liên quan và kết quả Cost có dữ liệu đối chiếu. Cơ chế lưu vết này làm giảm tranh luận trong quá trình tổ chức trận đấu.",
    "Thiết kế dữ liệu không chỉ phục vụ một phòng đơn mà còn chuẩn bị cho giải đấu, thông báo, hồ sơ người chơi và hoạt động cộng đồng. Đây là điểm có ý nghĩa nếu sản phẩm tiếp tục được phát triển sau học phần.":
    "Thiết kế dữ liệu phục vụ phòng đấu đơn, đồng thời chuẩn bị cho giải đấu, thông báo, hồ sơ người chơi và hoạt động cộng đồng. Phạm vi dữ liệu này phù hợp với định hướng phát triển sản phẩm sau học phần.",
    "Hệ thống chưa thể tự động xác minh tuyệt đối toàn bộ thông tin nhân vật, cung mệnh và vũ khí trong game. Một phần dữ liệu vẫn cần người chơi khai báo hoặc trọng tài kiểm tra, đặc biệt khi nguồn dữ liệu ngoài không đầy đủ hoặc người chơi không công khai hồ sơ.":
    "Hệ thống chưa thể tự động xác minh tuyệt đối toàn bộ thông tin nhân vật, cung mệnh và vũ khí trong trò chơi. Một phần dữ liệu vẫn cần người chơi khai báo hoặc trọng tài kiểm tra, đặc biệt khi nguồn dữ liệu ngoài không đầy đủ hoặc người chơi không công khai hồ sơ.",
    "Phần thống kê cần được thiết kế theo hướng truy vấn từ dữ liệu đã xác nhận, tránh dựa vào dữ liệu tạm thời trên client. Khi số lượng trận tăng, dữ liệu thống kê có thể trở thành giá trị riêng của hệ thống, giúp sản phẩm không chỉ phục vụ thao tác trong trận mà còn hỗ trợ phân tích sau giải.":
    "Phần thống kê cần được thiết kế theo hướng truy vấn từ dữ liệu đã xác nhận, tránh dựa vào dữ liệu tạm thời trên client. Khi số lượng trận tăng, dữ liệu thống kê có thể trở thành giá trị riêng của hệ thống, phục vụ thao tác trong trận và hỗ trợ phân tích sau giải.",
    "Thông qua quá trình thực hiện, đề tài cho thấy khả năng vận dụng các công nghệ phát triển web hiện đại vào một bài toán có trạng thái nghiệp vụ rõ ràng. Kết quả đạt được không chỉ là giao diện thao tác, mà còn là mô hình dữ liệu, kiến trúc phân tầng, cơ chế realtime và định hướng mở rộng. Những hạn chế hiện tại là cơ sở để tiếp tục hoàn thiện sản phẩm trong tương lai.":
    "Hệ thống vận dụng các công nghệ phát triển web hiện đại vào bài toán có trạng thái nghiệp vụ rõ ràng. Kết quả đạt được gồm giao diện thao tác, mô hình dữ liệu, kiến trúc phân tầng, cơ chế realtime và định hướng mở rộng. Những hạn chế hiện tại là cơ sở để tiếp tục hoàn thiện sản phẩm trong tương lai.",
    "Khi nghiệp vụ được mô tả rõ bằng actor, use case, trạng thái và bảng dữ liệu, quá trình hiện thực trở nên có định hướng hơn. Mỗi component hoặc API đều có lý do tồn tại và gắn với một bước trong vòng đời trận đấu. Điều này giúp báo cáo cũng có mạch trình bày chặt chẽ hơn, tránh liệt kê chức năng một cách rời rạc.":
    "Khi nghiệp vụ được mô tả rõ bằng actor, use case, trạng thái và bảng dữ liệu, quá trình hiện thực có định hướng cụ thể hơn. Mỗi component hoặc API đều gắn với một bước trong vòng đời trận đấu. Báo cáo vì vậy trình bày chức năng theo mạch nghiệp vụ, thay vì liệt kê rời rạc theo màn hình.",
    "Đối với một dự án có nhiều trạng thái và nhiều vai trò, tổ chức mã nguồn có ảnh hưởng lớn đến khả năng sửa lỗi. Nếu logic kiểm tra lượt nằm trong nhiều component khác nhau, việc thay đổi luật draft sẽ rất khó kiểm soát. Ngược lại, khi quy tắc được đưa vào policy/service, thay đổi luật có thể được thực hiện tập trung hơn và kiểm thử dễ hơn.":
    "Đối với một dự án có nhiều trạng thái và nhiều vai trò, tổ chức mã nguồn ảnh hưởng trực tiếp đến khả năng sửa lỗi. Nếu logic kiểm tra lượt nằm trong nhiều component khác nhau, việc thay đổi luật draft sẽ khó kiểm soát. Khi quy tắc được đưa vào policy/service, thay đổi luật có thể được thực hiện tập trung hơn và kiểm thử thuận lợi hơn.",
    "Đề tài cũng giúp sinh viên tiếp cận mô hình phát triển web hiện đại, nơi frontend, backend, cơ sở dữ liệu và triển khai không còn tách rời hoàn toàn. Next.js cho phép xây dựng giao diện và API trong cùng một dự án; Prisma giúp liên kết TypeScript với PostgreSQL; Supabase bổ sung realtime; Docker hỗ trợ đóng gói môi trường chạy. Việc phối hợp các công nghệ này giúp sinh viên hiểu rõ hơn cách một ứng dụng web vận hành từ giao diện đến dữ liệu.":
    "Quá trình thực hiện đề tài tạo điều kiện tiếp cận mô hình phát triển web hiện đại, trong đó frontend, backend, cơ sở dữ liệu và triển khai được phối hợp trong cùng một quy trình. Next.js cho phép xây dựng giao diện và API trong cùng một dự án; Prisma liên kết TypeScript với PostgreSQL; Supabase bổ sung realtime; Docker hỗ trợ đóng gói môi trường chạy. Việc phối hợp các công nghệ này làm rõ cách một ứng dụng web vận hành từ giao diện đến dữ liệu.",
    "Từ góc độ sản phẩm, đề tài cho thấy một phần mềm tốt không chỉ cần có chức năng đúng mà còn cần phù hợp với ngữ cảnh sử dụng. Trong phòng cấm/chọn, người dùng cần thao tác nhanh, nhìn thấy trạng thái rõ ràng và tin tưởng kết quả. Vì vậy, các quyết định về bố cục, màu sắc, thông báo lỗi, realtime và kết quả đều phải phục vụ mục tiêu này.":
    "Từ góc độ sản phẩm, phần mềm cần có chức năng đúng và phù hợp với ngữ cảnh sử dụng. Trong phòng cấm/chọn, người dùng cần thao tác nhanh, nhận biết trạng thái rõ ràng và tin tưởng kết quả. Vì vậy, các quyết định về bố cục, màu sắc, thông báo lỗi, realtime và kết quả đều phải phục vụ mục tiêu này.",
    "Nhìn tổng thể, đề tài có giá trị ở việc kết hợp một nhu cầu cụ thể của cộng đồng với các công nghệ phát triển web hiện đại. Sản phẩm có tính ứng dụng rõ ràng, dữ liệu có mô hình quan hệ, giao diện có định hướng thi đấu và kiến trúc có phân tầng. Những yếu tố này giúp báo cáo không chỉ mô tả một website đã làm, mà còn trình bày được tư duy phân tích, thiết kế và đánh giá hệ thống.":
    "Nhìn tổng thể, hệ thống kết hợp nhu cầu cụ thể của cộng đồng với các công nghệ phát triển web hiện đại. Sản phẩm có mô hình dữ liệu quan hệ, giao diện định hướng thi đấu và kiến trúc phân tầng. Báo cáo trình bày quá trình phân tích, thiết kế, hiện thực và đánh giá hệ thống thay vì chỉ mô tả các màn hình đã xây dựng.",
    "Mã nguồn cũng thể hiện": "Mã nguồn được tổ chức theo",
    "cho thấy": "ghi nhận",
    "thể hiện": "biểu thị",
    "giúp cải thiện trải nghiệm": "cải thiện trải nghiệm",
    "Cấu hình app, wiring service": "Cấu hình ứng dụng, wiring service",
}


STYLE_GUARD_REGEX_REPLACEMENTS = [
    (r"\bgame\b", "trò chơi"),
    (r"\bdemo\b", "bản minh họa"),
    (r"\btool\b", "công cụ"),
    (r"\bbấm nút\b", "thực hiện thao tác"),
    (r"\bchạy được\b", "có thể vận hành"),
    (r"\blàm được\b", "hiện thực được"),
    (r"\bdễ hơn\b", "thuận lợi hơn"),
    (r"\bgọn hơn\b", "tối ưu hơn"),
    (r"\brối\b", "thiếu rõ ràng"),
    (r"\bgiúp giảm\b", "làm giảm"),
    (r"\bgiúp người dùng\b", "hỗ trợ người dùng"),
    (r"\bgiúp người chơi\b", "hỗ trợ người chơi"),
    (r"\bgiúp trọng tài\b", "hỗ trợ trọng tài"),
    (r"\bgiúp hệ thống\b", "hỗ trợ hệ thống"),
    (r"\bgiúp mã nguồn\b", "hỗ trợ mã nguồn"),
    (r"\bgiúp quá trình\b", "hỗ trợ quá trình"),
    (r"\bgiúp\b", "hỗ trợ"),
]


def rewrite_text_for_style_guard(text: str) -> str:
    if not text:
        return text
    if "App Router" in text or "src/app" in text:
        protected = text
    else:
        protected = text.replace(" app ", " ứng dụng ")
    for old, new in STYLE_GUARD_REPLACEMENTS.items():
        protected = protected.replace(old, new)
    for pattern, new in STYLE_GUARD_REGEX_REPLACEMENTS:
        protected = re.sub(pattern, new, protected, flags=re.IGNORECASE)
    protected = protected.replace("không chỉ ", "")
    protected = protected.replace(", mà còn ", ", đồng thời ")
    protected = protected.replace(" mà còn ", " và ")
    protected = protected.replace("Điều này ", "Cơ chế này ")
    return protected


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.text == new_text:
        return
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def apply_vietnamese_style_guard(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        replace_paragraph_text(paragraph, rewrite_text_for_style_guard(paragraph.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_paragraph_text(paragraph, rewrite_text_for_style_guard(paragraph.text))


def build_doc():
    ensure_dirs()
    extract_media()
    diagrams = create_diagrams()
    doc = Document()
    set_styles(doc)
    set_update_fields(doc)
    footer_page_numbers(doc)
    add_cover(doc)
    add_front_matter(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc, diagrams)
    chapter_4(doc)
    chapter_5(doc)
    references_and_appendix(doc)
    apply_vietnamese_style_guard(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
