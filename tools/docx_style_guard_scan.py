from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document


PATTERNS = [
    r"\bđề tài cho thấy\b",
    r"\bđề tài cũng giúp\b",
    r"\bđề tài cũng thể hiện\b",
    r"\bđề tài mở ra\b",
    r"\bđề tài chứng minh rằng\b",
    r"\bbáo cáo cho thấy\b",
    r"\bbáo cáo thể hiện\b",
    r"\bphần này cho thấy\b",
    r"\bchương này cho thấy\b",
    r"\bnội dung này giúp\b",
    r"\bcách tiếp cận này giúp\b",
    r"\bcách tổ chức này giúp\b",
    r"\bngười đọc có thể thấy\b",
    r"\bngười đọc dễ thấy\b",
    r"\bcó thể thấy rằng\b",
    r"\bđiều này cho thấy\b",
    r"\bđiều này giúp\b",
    r"\bđiều này có nghĩa là\b",
    r"\bnhờ đó người đọc\b",
    r"\btừ đó người đọc\b",
    r"\bchúng ta có thể thấy\b",
    r"\bta thấy rằng\b",
    r"\bở đây ta thấy\b",
    r"\bnhư đã nói ở trên\b",
    r"\bphần này sẽ nói về\b",
    r"\btiếp theo em sẽ trình bày\b",
    r"\bsau đây là\b",
    r"\bnói chung\b",
    r"\bnhìn chung\b",
    r"\bcó thể nói là\b",
    r"\bnói cách khác\b",
    r"\bhiểu đơn giản là\b",
    r"\bnếu nhìn theo hướng\b",
    r"\bnếu xét theo góc độ\b",
    r"\bkhông chỉ\b",
    r"\bmà còn\b",
    r"\bđóng vai trò quan trọng\b",
    r"\bgóp phần nâng cao\b",
    r"\bgiúp cải thiện trải nghiệm\b",
    r"\btạo nền tảng vững chắc\b",
    r"\bmở ra cơ hội\b",
    r"\blà yếu tố then chốt\b",
    r"\bcó ý nghĩa quan trọng\b",
    r"\bđây là điểm khác biệt đáng kể\b",
    r"\btrong bối cảnh hiện nay\b",
    r"\bngày càng phát triển mạnh mẽ\b",
    r"\bvới sự phát triển của công nghệ\b",
    r"\btừ đó nâng cao hiệu quả\b",
    r"\bmang lại giá trị thực tiễn\b",
    r"\bcó tính ứng dụng cao\b",
    r"\bgame\b",
    r"\bbấm nút\b",
    r"\bchạy được\b",
    r"\blàm được\b",
    r"\blàm cho rõ hơn\b",
    r"\bdễ hơn\b",
    r"\bgọn hơn\b",
    r"\brối\b",
    r"\brất tốt\b",
    r"\bkhá tốt\b",
    r"\bkhá rõ\b",
    r"\bkhá nhiều\b",
    r"\bcực kỳ\b",
    r"\bxịn\b",
    r"\bmượt\b",
    r"\bổn áp\b",
    r"\bngon\b",
    r"\bsản phẩm thật\b",
    r"\bdemo\b",
    r"\btool\b",
    r"(?<![/A-Za-z])app(?!\s+Router|[A-Za-z])",
    r"\bngười dùng hiểu ngay\b",
    r"\bnhìn vào là biết\b",
    r"\bkhông bị loạn\b",
    r"\bkhông bị rối\b",
    r"\blàm chắc lõi\b",
    r"\bphần này hơi\b",
    r"\bcái này\b",
    r"\bthứ này\b",
]


def iter_paragraphs(doc: Document):
    for i, p in enumerate(doc.paragraphs, start=1):
        yield f"p#{i}", p.text
    for ti, table in enumerate(doc.tables, start=1):
        for ri, row in enumerate(table.rows, start=1):
            for ci, cell in enumerate(row.cells, start=1):
                for pi, p in enumerate(cell.paragraphs, start=1):
                    yield f"t#{ti}r{ri}c{ci}p{pi}", p.text


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: python docx_style_guard_scan.py FILE.docx")
        return 2
    path = Path(sys.argv[1])
    doc = Document(path)
    compiled = [(pat, re.compile(pat, flags=re.IGNORECASE)) for pat in PATTERNS]
    findings = []
    help_count = 0
    for loc, text in iter_paragraphs(doc):
        if not text.strip():
            continue
        help_count += len(re.findall(r"\bgiúp\b", text, flags=re.IGNORECASE))
        for raw, pattern in compiled:
            if pattern.search(text):
                findings.append((loc, raw, " ".join(text.split())[:240]))
    print(f"file={path}")
    print(f"findings={len(findings)}")
    print(f"giup_count={help_count}")
    for loc, raw, text in findings[:300]:
        print(f"{loc}\t{raw}\t{text}")
    if len(findings) > 300:
        print(f"... {len(findings) - 300} more")
    return 1 if findings else 0


if __name__ == "__main__":
    raise SystemExit(main())
