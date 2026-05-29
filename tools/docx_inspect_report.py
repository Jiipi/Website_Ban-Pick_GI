from __future__ import annotations

import json
import re
import sys
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document


HEADING_RE = re.compile(r"^\s*(\d+(?:\.\d+)*)\.?\s+(.+)")


def para_text(paragraph) -> str:
    return re.sub(r"\s+", " ", paragraph.text).strip()


def docx_zip_counts(path: Path) -> dict:
    with zipfile.ZipFile(path) as zf:
        names = zf.namelist()
    media = [n for n in names if n.startswith("word/media/")]
    charts = [n for n in names if n.startswith("word/charts/")]
    embeddings = [n for n in names if n.startswith("word/embeddings/")]
    return {
        "zip_entries": len(names),
        "media_files": len(media),
        "chart_files": len(charts),
        "embedded_files": len(embeddings),
        "media_exts": dict(Counter(Path(n).suffix.lower() for n in media)),
    }


def inspect_doc(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [para_text(p) for p in doc.paragraphs]
    paragraphs = [p for p in paragraphs if p]
    styles = Counter(p.style.name if p.style is not None else "" for p in doc.paragraphs)

    headings = []
    manual_numbered = []
    short_paras = 0
    for idx, p in enumerate(doc.paragraphs, start=1):
        text = para_text(p)
        if not text:
            continue
        style = p.style.name if p.style is not None else ""
        if style.startswith("Heading") or HEADING_RE.match(text):
            headings.append(
                {
                    "paragraph": idx,
                    "style": style,
                    "text": text[:180],
                }
            )
        m = HEADING_RE.match(text)
        if m and not style.startswith("Heading"):
            manual_numbered.append({"paragraph": idx, "style": style, "text": text[:180]})
        if len(text) <= 45:
            short_paras += 1

    table_summaries = []
    for table_idx, table in enumerate(doc.tables, start=1):
        row_lengths = []
        first_row = []
        for r_idx, row in enumerate(table.rows):
            row_text = [" ".join(cell.text.split()) for cell in row.cells]
            row_lengths.append(sum(len(x) for x in row_text))
            if r_idx == 0:
                first_row = row_text
        table_summaries.append(
            {
                "index": table_idx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "first_row": first_row,
                "avg_row_chars": round(sum(row_lengths) / len(row_lengths), 1) if row_lengths else 0,
                "very_short_rows": sum(1 for n in row_lengths if n < 60),
            }
        )

    return {
        "path": str(path),
        "size": path.stat().st_size,
        "paragraph_count": len(doc.paragraphs),
        "nonempty_paragraph_count": len(paragraphs),
        "word_count_approx": sum(len(p.split()) for p in paragraphs),
        "short_nonempty_paragraphs": short_paras,
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "sections": len(doc.sections),
        "top_styles": styles.most_common(20),
        "heading_like_count": len(headings),
        "manual_numbered_heading_count": len(manual_numbered),
        "headings": headings[:160],
        "manual_numbered_headings": manual_numbered[:80],
        "table_summaries": table_summaries[:80],
        **docx_zip_counts(path),
    }


def extract_outline(path: Path, out_path: Path) -> None:
    doc = Document(path)
    lines = []
    for i, p in enumerate(doc.paragraphs, start=1):
        text = para_text(p)
        if not text:
            continue
        style = p.style.name if p.style is not None else ""
        marker = ""
        if style.startswith("Heading"):
            marker = "HEADING"
        elif HEADING_RE.match(text):
            marker = "NUMBERED"
        elif len(text) <= 45:
            marker = "SHORT"
        else:
            marker = "TEXT"
        lines.append(f"{i:04d}\t{marker}\t{style}\t{text}")
    out_path.write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    if len(sys.argv) < 3:
        print("Usage: python docx_inspect_report.py OUT_DIR DOCX [DOCX...]")
        return 2

    out_dir = Path(sys.argv[1])
    out_dir.mkdir(parents=True, exist_ok=True)
    results = []
    for raw in sys.argv[2:]:
        path = Path(raw)
        result = inspect_doc(path)
        results.append(result)
        safe = path.stem.replace(" ", "_")
        (out_dir / f"{safe}.json").write_text(
            json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        extract_outline(path, out_dir / f"{safe}.outline.txt")
    (out_dir / "summary.json").write_text(
        json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(json.dumps(results, ensure_ascii=False, indent=2)[:6000])
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
